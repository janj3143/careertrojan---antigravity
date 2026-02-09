
# Enhanced Sidebar Integration
import sys
from pathlib import Path
shared_path = Path(__file__).parent.parent / "shared"
if str(shared_path) not in sys.path:
    sys.path.insert(0, str(shared_path))

try:
    from enhanced_sidebar import render_enhanced_sidebar, inject_sidebar_css
    ENHANCED_SIDEBAR_AVAILABLE = True
except ImportError:
    ENHANCED_SIDEBAR_AVAILABLE = False


# Activate Enhanced Sidebar
if ENHANCED_SIDEBAR_AVAILABLE:
    inject_sidebar_css()
    render_enhanced_sidebar()

"""
Email Integration Module for IntelliCV Admin Portal
================================================

This module provides comprehensive email integration capabilities:
- IMAP/POP3 email account connections
- Document extraction from email attachments
- Email metadata processing and storage
- Integration with existing CV/document parsing pipeline

Supported Email Providers:
- Gmail (IMAP)
- Outlook/Hotmail (IMAP)
- Yahoo Mail (IMAP)
- Custom IMAP servers
- Exchange servers

Document Types Supported:
- PDF files
- Microsoft Word (.doc, .docx)
- Text files (.txt)
- RTF files
- OpenDocument Text (.odt)
"""

import imaplib
import poplib
import email
import email.header
import email.message
import email.utils
import sqlite3
import json
import os
import hashlib
import csv
import re
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
import logging
from dataclasses import dataclass
import tempfile
import mimetypes

# Document processing imports
try:
    import PyPDF2
    from docx import Document
    import zipfile
    HAS_DOC_PROCESSING = True
except ImportError:
    HAS_DOC_PROCESSING = False
    logging.warning("Document processing libraries not available. Install PyPDF2 and python-docx for full functionality.")


@dataclass
class EmailAccount:
    """Email account configuration"""
    name: str
    email_address: str
    provider: str
    imap_server: str
    imap_port: int
    username: str
    password: str
    use_ssl: bool = True
    is_active: bool = True


@dataclass
class EmailMessage:
    """Email message metadata"""
    message_id: str
    subject: str
    sender: str
    recipients: List[str]
    date_received: datetime
    body_text: str
    body_html: str
    has_attachments: bool
    attachment_count: int
    folder: str
    account_id: str


@dataclass
class ExtractedDocument:
    """Extracted document from email attachment"""
    document_id: str
    email_message_id: str
    filename: str
    file_type: str
    file_size: int
    content_text: str
    content_hash: str
    extraction_date: datetime
    metadata: Dict[str, Any]


class EmailIntegrationManager:
    """Main class for managing email integration and document extraction"""
    
    def __init__(self, db_path: str = "db/intellicv_data.db"):
        self.db_path = db_path
        self.logger = logging.getLogger(__name__)
        self._setup_database()
        
        # Common email provider configurations
        self.provider_configs = {
            "gmail": {
                "imap_server": "imap.gmail.com",
                "imap_port": 993,
                "pop_server": "pop.gmail.com",
                "pop_port": 995,
                "use_ssl": True
            },
            "outlook": {
                "imap_server": "outlook.office365.com",
                "imap_port": 993,
                "pop_server": "outlook.office365.com",
                "pop_port": 995,
                "use_ssl": True
            },
            "yahoo": {
                "imap_server": "imap.mail.yahoo.com",
                "imap_port": 993,
                "pop_server": "pop.mail.yahoo.com",
                "pop_port": 995,
                "use_ssl": True
            }
        }
    
    def _setup_database(self):
        """Setup database tables for email integration"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Email accounts table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS email_accounts (
                    id TEXT PRIMARY KEY,
                    name TEXT NOT NULL,
                    email_address TEXT NOT NULL UNIQUE,
                    provider TEXT NOT NULL,
                    imap_server TEXT,
                    imap_port INTEGER,
                    username TEXT,
                    encrypted_password TEXT,
                    use_ssl BOOLEAN DEFAULT TRUE,
                    is_active BOOLEAN DEFAULT TRUE,
                    last_sync TIMESTAMP,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            
            # Email messages table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS email_messages (
                    id TEXT PRIMARY KEY,
                    message_id TEXT NOT NULL,
                    account_id TEXT NOT NULL,
                    subject TEXT,
                    sender TEXT,
                    recipients TEXT, -- JSON array
                    date_received TIMESTAMP,
                    body_text TEXT,
                    body_html TEXT,
                    has_attachments BOOLEAN DEFAULT FALSE,
                    attachment_count INTEGER DEFAULT 0,
                    folder TEXT DEFAULT 'INBOX',
                    processed BOOLEAN DEFAULT FALSE,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (account_id) REFERENCES email_accounts (id)
                )
            ''')
            
            # Extracted documents table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS extracted_documents (
                    id TEXT PRIMARY KEY,
                    email_message_id TEXT NOT NULL,
                    filename TEXT NOT NULL,
                    file_type TEXT,
                    file_size INTEGER,
                    content_text TEXT,
                    content_hash TEXT UNIQUE,
                    extraction_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    metadata TEXT, -- JSON
                    is_resume BOOLEAN DEFAULT FALSE,
                    is_processed BOOLEAN DEFAULT FALSE,
                    processing_status TEXT DEFAULT 'pending',
                    processing_results TEXT, -- JSON
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (email_message_id) REFERENCES email_messages (id)
                )
            ''')
            
            # Email sync logs table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS email_sync_logs (
                    id TEXT PRIMARY KEY,
                    account_id TEXT NOT NULL,
                    sync_start TIMESTAMP,
                    sync_end TIMESTAMP,
                    messages_processed INTEGER DEFAULT 0,
                    documents_extracted INTEGER DEFAULT 0,
                    errors_count INTEGER DEFAULT 0,
                    status TEXT DEFAULT 'running',
                    error_details TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (account_id) REFERENCES email_accounts (id)
                )
            ''')
            
            conn.commit()
            self.logger.info("Email integration database schema created successfully")
    
    def add_email_account(self, account: EmailAccount) -> str:
        """Add a new email account configuration"""
        account_id = hashlib.md5(f"{account.email_address}_{datetime.now()}".encode()).hexdigest()
        
        # Simple password encryption (in production, use proper encryption)
        encrypted_password = self._simple_encrypt(account.password)
        
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO email_accounts 
                (id, name, email_address, provider, imap_server, imap_port, 
                 username, encrypted_password, use_ssl, is_active)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                account_id, account.name, account.email_address, account.provider,
                account.imap_server, account.imap_port, account.username,
                encrypted_password, account.use_ssl, account.is_active
            ))
            conn.commit()
        
        self.logger.info(f"Email account added: {account.email_address}")
        return account_id
    
    def get_email_accounts(self) -> List[Dict[str, Any]]:
        """Get all configured email accounts"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                SELECT id, name, email_address, provider, is_active, last_sync
                FROM email_accounts 
                ORDER BY name
            ''')
            
            accounts = []
            for row in cursor.fetchall():
                accounts.append({
                    'id': row[0],
                    'name': row[1],
                    'email_address': row[2],
                    'provider': row[3],
                    'is_active': bool(row[4]),
                    'last_sync': row[5]
                })
            
            return accounts
    
    def test_email_connection(self, email_address: str, password: str, provider: str) -> Dict[str, Any]:
        """Test email connection without storing the account
        
        Args:
            email_address: Email address to test
            password: App password or OAuth token
            provider: Email provider (gmail, outlook, yahoo)
            
        Returns:
            Dict with success status and connection details
        """
        try:
            if provider not in self.provider_configs:
                return {
                    "success": False,
                    "error": f"Unsupported provider: {provider}",
                    "details": "Provider not in configuration"
                }
            
            config = self.provider_configs[provider]
            server = config["imap_server"]
            port = config["imap_port"]
            
            # Connect to IMAP server
            mail = imaplib.IMAP4_SSL(server, port)
            
            # Login with credentials
            mail.login(email_address, password)
            self.logger.info(f"Successfully tested connection to {email_address}")
            
            # Get inbox info
            mail.select('INBOX')
            status, messages = mail.search(None, 'ALL')
            
            total_emails = 0
            if status == 'OK':
                message_ids = messages[0].split()
                total_emails = len(message_ids)
            
            # Close connection
            mail.close()
            mail.logout()
            
            return {
                "success": True,
                "message": f"Connection successful to {provider}",
                "total_emails_in_inbox": total_emails,
                "server": server,
                "port": port
            }
            
        except imaplib.IMAP4.error as e:
            error_msg = f"IMAP error: {str(e)}"
            self.logger.error(f"IMAP connection test failed for {email_address}: {error_msg}")
            
            # Enhanced error handling for authentication failures
            if "AUTHENTICATIONFAILED" in str(e) or "Invalid credentials" in str(e):
                troubleshooting_guide = {
                    "gmail": [
                        "✅ Enable 2-Factor Authentication on your Google account",
                        "🔑 Generate App Password: Google Account → Security → App passwords", 
                        "📧 Use the 16-character App Password (NOT your regular Gmail password)",
                        "🔗 Link: https://myaccount.google.com/apppasswords"
                    ],
                    "yahoo": [
                        "✅ Enable 2-Factor Authentication on Yahoo account",
                        "🔑 Generate App Password: Account Security → Generate app password",
                        "📧 Select 'Other app' and enter 'IMAP' as name", 
                        "🔗 Link: https://login.yahoo.com/account/security"
                    ],
                    "outlook": [
                        "✅ Enable 2-Factor Authentication on Microsoft account",
                        "🔑 Generate App Password: Security → Advanced security → App passwords",
                        "📧 Create password for 'Email' category",
                        "🔗 Link: https://account.microsoft.com/security"
                    ]
                }
                
                return {
                    "success": False,
                    "error": error_msg,
                    "error_type": "AUTHENTICATION_FAILED",
                    "details": "Authentication failed - App Password required",
                    "troubleshooting": troubleshooting_guide.get(provider, [
                        "Enable 2-Factor Authentication on your email account",
                        "Generate an App Password for IMAP access",
                        "Use App Password instead of regular password"
                    ])
                }
            else:
                return {
                    "success": False,
                    "error": error_msg,
                    "details": "IMAP connection failed"
                }
            
        except Exception as e:
            error_msg = f"Connection error: {str(e)}"
            self.logger.error(f"Email connection test failed for {email_address}: {error_msg}")
            return {
                "success": False,
                "error": error_msg,
                "details": "Unable to connect to email server"
            }
    
    def connect_to_email(self, account_id: str) -> Optional[imaplib.IMAP4_SSL]:
        """Connect to email account using IMAP"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute('''
                    SELECT email_address, imap_server, imap_port, username, 
                           encrypted_password, use_ssl
                    FROM email_accounts WHERE id = ? AND is_active = TRUE
                ''', (account_id,))
                
                account_data = cursor.fetchone()
                if not account_data:
                    self.logger.error(f"Account not found or inactive: {account_id}")
                    return None
                
                email_addr, server, port, username, enc_password, use_ssl = account_data
                password = self._simple_decrypt(enc_password)
                
                # Connect to IMAP server
                if use_ssl:
                    mail = imaplib.IMAP4_SSL(server, port)
                else:
                    mail = imaplib.IMAP4(server, port)
                
                mail.login(username or email_addr, password)
                self.logger.info(f"Successfully connected to {email_addr}")
                return mail
                
        except Exception as e:
            self.logger.error(f"Failed to connect to email account {account_id}: {str(e)}")
            return None
    
    def scan_emails_for_documents(self, account_id: str, days_back: int = 30, 
                                folder: str = "INBOX", historical_scan: bool = False,
                                batch_size: int = 100) -> Dict[str, Any]:
        """
        Scan emails for documents and extract them
        
        Args:
            account_id: Email account ID
            days_back: Days to scan back (ignored if historical_scan=True)
            folder: Email folder to scan
            historical_scan: If True, scan ALL emails back to 2011
            batch_size: Number of emails to process in each batch
        """
        mail = self.connect_to_email(account_id)
        if not mail:
            return {"error": "Failed to connect to email account"}
        
        try:
            # Start sync log
            sync_log_id = self._start_sync_log(account_id)
            
            mail.select(folder)
            
            # Search criteria based on scan type
            if historical_scan:
                # Scan ALL emails back to 2011 for archive processing
                since_date = "01-Jan-2011"
                search_criteria = f'(SINCE {since_date})'
                self.logger.info(f"Starting historical email archive scan back to {since_date}")
            else:
                # Standard recent email scan
                since_date = (datetime.now() - timedelta(days=days_back)).strftime("%d-%b-%Y")
                search_criteria = f'(SINCE {since_date})'
                self.logger.info(f"Starting recent email scan from {since_date}")
            
            typ, data = mail.search(None, search_criteria)
            email_ids = data[0].split()
            
            total_emails = len(email_ids)
            processed_count = 0
            extracted_docs = 0
            resumes_found = 0
            errors = []
            
            self.logger.info(f"Found {total_emails} emails to process")
            
            # Process emails in batches for large archives
            for i in range(0, min(len(email_ids), batch_size * 10 if historical_scan else batch_size)):
                email_id = email_ids[-(i+1)]  # Start from most recent
                
                try:
                    # Fetch email
                    typ, msg_data = mail.fetch(email_id, '(RFC822)')
                    email_body = msg_data[0][1]
                    email_message = email.message_from_bytes(email_body)
                    
                    # Process email and extract documents
                    message_info = self._process_email_message(
                        email_message, account_id, folder
                    )
                    
                    if message_info and message_info.has_attachments:
                        docs_extracted = self._extract_documents_from_email(
                            email_message, message_info.message_id
                        )
                        extracted_docs += docs_extracted
                        
                        # Count resumes specifically
                        if docs_extracted > 0:
                            resumes_in_email = self._count_resumes_in_email(message_info.message_id)
                            resumes_found += resumes_in_email
                    
                    processed_count += 1
                    
                    # Progress logging for large batches
                    if processed_count % 50 == 0:
                        self.logger.info(f"Processed {processed_count}/{min(total_emails, batch_size * 10 if historical_scan else batch_size)} emails...")
                    
                except Exception as e:
                    error_msg = f"Error processing email {email_id}: {str(e)}"
                    errors.append(error_msg)
                    self.logger.error(error_msg)
                
                # Break if we've processed enough for a standard scan
                if not historical_scan and processed_count >= batch_size:
                    break
            
            # Update sync log
            self._complete_sync_log(sync_log_id, processed_count, extracted_docs, len(errors))
            
            # Update last sync time
            self._update_last_sync(account_id)
            
            mail.logout()
            
            return {
                "success": True,
                "processed_emails": processed_count,
                "total_emails_found": total_emails,
                "extracted_documents": extracted_docs,
                "resumes_found": resumes_found,
                "errors": errors,
                "sync_log_id": sync_log_id,
                "historical_scan": historical_scan
            }
            
        except Exception as e:
            self.logger.error(f"Email scan failed: {str(e)}")
            return {"error": str(e)}
    
    def scan_email_archive_for_resumes(self, account_id: str, 
                                     target_years: List[int] = None) -> Dict[str, Any]:
        """
        Specialized method to scan email archives specifically for resumes/CVs
        Optimized for historical data going back to 2011
        
        Args:
            account_id: Email account ID
            target_years: Specific years to scan (if None, scans all from 2011)
        """
        mail = self.connect_to_email(account_id)
        if not mail:
            return {"error": "Failed to connect to email account"}
        
        try:
            if target_years is None:
                current_year = datetime.now().year
                target_years = list(range(2011, current_year + 1))
            
            self.logger.info(f"Starting resume archive scan for years: {target_years}")
            
            total_resumes = 0
            total_processed = 0
            year_results = {}
            
            for year in target_years:
                self.logger.info(f"Scanning year {year} for resumes...")
                
                # Scan each year separately for better control
                mail.select("INBOX")
                
                # Search for emails in this year with attachments
                since_date = f"01-Jan-{year}"
                before_date = f"31-Dec-{year}"
                
                # Look for emails with common resume-related terms and attachments
                search_criteria = f'(SINCE {since_date} BEFORE {before_date})'
                
                typ, data = mail.search(None, search_criteria)
                email_ids = data[0].split()
                
                year_resumes = 0
                year_processed = 0
                
                for email_id in email_ids[:200]:  # Limit per year for performance
                    try:
                        typ, msg_data = mail.fetch(email_id, '(RFC822)')
                        email_body = msg_data[0][1]
                        email_message = email.message_from_bytes(email_body)
                        
                        # Quick check for resume-related content
                        if self._email_contains_resume_indicators(email_message):
                            message_info = self._process_email_message(
                                email_message, account_id, "INBOX"
                            )
                            
                            if message_info and message_info.has_attachments:
                                docs_extracted = self._extract_documents_from_email(
                                    email_message, message_info.message_id
                                )
                                
                                if docs_extracted > 0:
                                    resumes_in_email = self._count_resumes_in_email(message_info.message_id)
                                    year_resumes += resumes_in_email
                        
                        year_processed += 1
                        
                    except Exception as e:
                        self.logger.error(f"Error processing email in year {year}: {str(e)}")
                
                year_results[year] = {
                    "processed": year_processed,
                    "resumes_found": year_resumes
                }
                
                total_resumes += year_resumes
                total_processed += year_processed
                
                self.logger.info(f"Year {year}: {year_processed} emails processed, {year_resumes} resumes found")
            
            mail.logout()
            
            return {
                "success": True,
                "total_processed": total_processed,
                "total_resumes_found": total_resumes,
                "year_breakdown": year_results,
                "years_scanned": target_years
            }
            
        except Exception as e:
            self.logger.error(f"Archive resume scan failed: {str(e)}")
            return {"error": str(e)}
    
    def export_resumes_to_ai_enrichment(self, limit: int = None) -> Dict[str, Any]:
        """
        Export extracted resumes to AI enrichment JSON format
        Integrates with existing AI enrichment pipeline
        """
        try:
            # Get all extracted resume documents
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                
                query = '''
                    SELECT ed.id, ed.filename, ed.content_text, ed.metadata,
                           ed.extraction_date, em.sender, em.date_received,
                           ea.email_address
                    FROM extracted_documents ed
                    JOIN email_messages em ON ed.email_message_id = em.id
                    JOIN email_accounts ea ON em.account_id = ea.id
                    WHERE ed.is_resume = TRUE
                    ORDER BY ed.extraction_date DESC
                '''
                
                if limit:
                    query += f' LIMIT {limit}'
                
                cursor.execute(query)
                resume_data = cursor.fetchall()
            
            if not resume_data:
                return {"error": "No resume documents found"}
            
            # Create AI enrichment export directory
            ai_export_path = Path(self.db_path).parent / "ai_enrichment_export"
            ai_export_path.mkdir(exist_ok=True)
            
            # Process each resume for AI enrichment
            enrichment_profiles = []
            processed_count = 0
            
            for row in resume_data:
                doc_id, filename, content_text, metadata_json, extraction_date, sender, email_date, account_email = row
                
                try:
                    metadata = json.loads(metadata_json) if metadata_json else {}
                    
                    # Create AI enrichment profile format
                    profile = {
                        "profile_id": doc_id,
                        "source": "email_extraction",
                        "original_filename": filename,
                        "extraction_metadata": {
                            "extracted_date": extraction_date,
                            "email_sender": sender,
                            "email_date": email_date,
                            "account": account_email,
                            "document_metadata": metadata
                        },
                        "raw_content": content_text,
                        "processing_status": "ready_for_enrichment",
                        "content_type": "cv_document",
                        "ai_processing": {
                            "status": "pending",
                            "created_at": datetime.now().isoformat(),
                            "priority": "high"  # Email-extracted resumes get high priority
                        }
                    }
                    
                    enrichment_profiles.append(profile)
                    processed_count += 1
                    
                    # Save individual profile file
                    profile_file = ai_export_path / f"resume_profile_{doc_id[:8]}.json"
                    with open(profile_file, 'w', encoding='utf-8') as f:
                        json.dump(profile, f, indent=2, ensure_ascii=False)
                    
                except Exception as e:
                    self.logger.error(f"Error processing resume {doc_id}: {str(e)}")
            
            # Save master export file
            master_export = {
                "export_metadata": {
                    "export_date": datetime.now().isoformat(),
                    "total_profiles": processed_count,
                    "source_system": "email_integration",
                    "export_type": "resume_batch"
                },
                "profiles": enrichment_profiles
            }
            
            master_file = ai_export_path / f"resume_batch_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            with open(master_file, 'w', encoding='utf-8') as f:
                json.dump(master_export, f, indent=2, ensure_ascii=False)
            
            # Update processing status in database
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                processed_ids = [profile["profile_id"] for profile in enrichment_profiles]
                placeholders = ','.join('?' * len(processed_ids))
                cursor.execute(f'''
                    UPDATE extracted_documents 
                    SET processing_status = 'exported_to_ai',
                        processing_results = ?
                    WHERE id IN ({placeholders})
                ''', [json.dumps({"export_date": datetime.now().isoformat(), "export_file": str(master_file)})] + processed_ids)
                conn.commit()
            
            self.logger.info(f"Exported {processed_count} resumes to AI enrichment pipeline")
            
            return {
                "success": True,
                "profiles_exported": processed_count,
                "export_file": str(master_file),
                "export_directory": str(ai_export_path),
                "ai_processing_ready": True
            }
            
        except Exception as e:
            self.logger.error(f"AI enrichment export failed: {str(e)}")
            return {"error": str(e)}
    
    def _process_email_message(self, email_message: email.message.Message, 
                             account_id: str, folder: str) -> Optional[EmailMessage]:
        """Process and store email message metadata"""
        try:
            # Extract email metadata
            message_id = email_message.get('Message-ID', '')
            subject = self._decode_header(email_message.get('Subject', ''))
            sender = self._decode_header(email_message.get('From', ''))
            recipients = self._decode_header(email_message.get('To', ''))
            
            # Parse date
            date_str = email_message.get('Date', '')
            try:
                date_received = email.utils.parsedate_to_datetime(date_str)
            except:
                date_received = datetime.now()
            
            # Extract body text
            body_text = ""
            body_html = ""
            has_attachments = False
            attachment_count = 0
            
            if email_message.is_multipart():
                for part in email_message.walk():
                    content_type = part.get_content_type()
                    content_disposition = str(part.get("Content-Disposition"))
                    
                    if "attachment" in content_disposition:
                        has_attachments = True
                        attachment_count += 1
                    elif content_type == "text/plain" and not body_text:
                        body_text = part.get_payload(decode=True).decode('utf-8', errors='ignore')
                    elif content_type == "text/html" and not body_html:
                        body_html = part.get_payload(decode=True).decode('utf-8', errors='ignore')
            else:
                body_text = email_message.get_payload(decode=True).decode('utf-8', errors='ignore')
            
            # Store in database
            msg_id = hashlib.md5(f"{message_id}_{account_id}".encode()).hexdigest()
            
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute('''
                    INSERT OR REPLACE INTO email_messages
                    (id, message_id, account_id, subject, sender, recipients,
                     date_received, body_text, body_html, has_attachments,
                     attachment_count, folder)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    msg_id, message_id, account_id, subject, sender, recipients,
                    date_received, body_text, body_html, has_attachments,
                    attachment_count, folder
                ))
                conn.commit()
            
            return EmailMessage(
                message_id=msg_id,
                subject=subject,
                sender=sender,
                recipients=[recipients],
                date_received=date_received,
                body_text=body_text,
                body_html=body_html,
                has_attachments=has_attachments,
                attachment_count=attachment_count,
                folder=folder,
                account_id=account_id
            )
            
        except Exception as e:
            self.logger.error(f"Error processing email message: {str(e)}")
            return None
    
    def _extract_documents_from_email(self, email_message: email.message.Message, 
                                    email_msg_id: str) -> int:
        """Extract documents from email attachments"""
        if not HAS_DOC_PROCESSING:
            self.logger.warning("Document processing libraries not available")
            return 0
        
        extracted_count = 0
        
        try:
            for part in email_message.walk():
                content_disposition = str(part.get("Content-Disposition"))
                
                if "attachment" in content_disposition:
                    filename = part.get_filename()
                    if filename:
                        filename = self._decode_header(filename)
                        
                        # Check if it's a document we can process
                        if self._is_processable_document(filename):
                            file_data = part.get_payload(decode=True)
                            
                            if file_data:
                                doc_extracted = self._process_document_attachment(
                                    filename, file_data, email_msg_id
                                )
                                if doc_extracted:
                                    extracted_count += 1
            
            return extracted_count
            
        except Exception as e:
            self.logger.error(f"Error extracting documents: {str(e)}")
            return 0
    
    def _process_document_attachment(self, filename: str, file_data: bytes, 
                                   email_msg_id: str) -> bool:
        """Process individual document attachment"""
        try:
            # Generate document ID
            content_hash = hashlib.sha256(file_data).hexdigest()
            doc_id = hashlib.md5(f"{filename}_{email_msg_id}_{content_hash}".encode()).hexdigest()
            
            # Extract text content based on file type
            file_ext = Path(filename).suffix.lower()
            content_text = ""
            metadata = {}
            
            if file_ext == '.pdf':
                content_text, metadata = self._extract_pdf_content(file_data)
            elif file_ext in ['.doc', '.docx']:
                content_text, metadata = self._extract_word_content(file_data)
            elif file_ext == '.txt':
                content_text = file_data.decode('utf-8', errors='ignore')
            
            # Check if document already exists
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute('SELECT id FROM extracted_documents WHERE content_hash = ?', (content_hash,))
                if cursor.fetchone():
                    self.logger.info(f"Document already exists: {filename}")
                    return False
                
                # Store extracted document
                cursor.execute('''
                    INSERT INTO extracted_documents
                    (id, email_message_id, filename, file_type, file_size,
                     content_text, content_hash, metadata, is_resume)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    doc_id, email_msg_id, filename, file_ext,
                    len(file_data), content_text, content_hash,
                    json.dumps(metadata), self._is_likely_resume(content_text)
                ))
                conn.commit()
            
            self.logger.info(f"Document extracted: {filename}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error processing document {filename}: {str(e)}")
            return False
    
    def _extract_pdf_content(self, file_data: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text content from PDF"""
        try:
            with tempfile.NamedTemporaryFile() as temp_file:
                temp_file.write(file_data)
                temp_file.flush()
                
                with open(temp_file.name, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    text_content = ""
                    
                    for page in pdf_reader.pages:
                        text_content += page.extract_text() + "\n"
                    
                    metadata = {
                        "page_count": len(pdf_reader.pages),
                        "document_info": dict(pdf_reader.metadata) if pdf_reader.metadata else {}
                    }
                    
                    return text_content.strip(), metadata
                    
        except Exception as e:
            self.logger.error(f"Error extracting PDF content: {str(e)}")
            return "", {"error": str(e)}
    
    def _extract_word_content(self, file_data: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text content from Word document"""
        try:
            with tempfile.NamedTemporaryFile(suffix='.docx') as temp_file:
                temp_file.write(file_data)
                temp_file.flush()
                
                doc = Document(temp_file.name)
                text_content = ""
                
                for paragraph in doc.paragraphs:
                    text_content += paragraph.text + "\n"
                
                metadata = {
                    "paragraph_count": len(doc.paragraphs),
                    "core_properties": {
                        "author": doc.core_properties.author,
                        "created": str(doc.core_properties.created),
                        "modified": str(doc.core_properties.modified),
                        "title": doc.core_properties.title
                    }
                }
                
                return text_content.strip(), metadata
                
        except Exception as e:
            self.logger.error(f"Error extracting Word content: {str(e)}")
            return "", {"error": str(e)}
    
    def _is_processable_document(self, filename: str) -> bool:
        """Check if file is a processable document type"""
        supported_extensions = {'.pdf', '.doc', '.docx', '.txt', '.rtf', '.odt'}
        return Path(filename).suffix.lower() in supported_extensions
    
    def _is_likely_resume(self, content: str) -> bool:
        """Simple heuristic to detect if document is likely a resume/CV"""
        resume_keywords = [
            'experience', 'education', 'skills', 'resume', 'cv', 'curriculum vitae',
            'employment', 'qualifications', 'objective', 'summary', 'career',
            'professional', 'achievements', 'references', 'contact', 'phone',
            'email', 'address', 'linkedin', 'university', 'degree', 'certification'
        ]
        
        content_lower = content.lower()
        keyword_count = sum(1 for keyword in resume_keywords if keyword in content_lower)
        
        # If document contains 5 or more resume keywords, likely a resume
        return keyword_count >= 5
    
    def _is_likely_job_description(self, content: str) -> bool:
        """Detect if document is likely a job description/posting"""
        job_keywords = [
            'job description', 'position', 'responsibilities', 'requirements',
            'qualifications', 'candidate', 'applicant', 'hiring', 'employment',
            'salary', 'benefits', 'we are looking for', 'seeking', 'opportunity',
            'role', 'duties', 'company', 'team', 'department', 'reports to',
            'job title', 'location', 'full time', 'part time', 'contract',
            'apply', 'application', 'interview', 'hr', 'human resources'
        ]
        
        content_lower = content.lower()
        keyword_count = sum(1 for keyword in job_keywords if keyword in content_lower)
        
        # If document contains 4 or more job keywords, likely a job description
        return keyword_count >= 4
    
    def _extract_email_addresses_from_content(self, content: str) -> List[str]:
        """Extract email addresses from document content"""
        import re
        
        # Comprehensive email regex pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, content)
        
        # Clean and deduplicate emails
        cleaned_emails = []
        for email in emails:
            email = email.lower().strip()
            if email and email not in cleaned_emails:
                # Filter out common false positives
                if not any(exclude in email for exclude in ['example.com', 'test.com', 'domain.com']):
                    cleaned_emails.append(email)
        
        return cleaned_emails
    
    def _classify_document_type(self, content: str, filename: str) -> str:
        """Classify document as resume, job_description, or other"""
        is_resume = self._is_likely_resume(content)
        is_job_desc = self._is_likely_job_description(content)
        
        # Check filename for additional clues
        filename_lower = filename.lower()
        
        if 'resume' in filename_lower or 'cv' in filename_lower:
            return 'resume'
        elif any(term in filename_lower for term in ['job', 'position', 'role', 'posting']):
            return 'job_description'
        elif is_resume and not is_job_desc:
            return 'resume'
        elif is_job_desc and not is_resume:
            return 'job_description'
        elif is_resume and is_job_desc:
            # If both, filename or other clues determine
            if 'apply' in content.lower() or 'hiring' in content.lower():
                return 'job_description'
            else:
                return 'resume'
        else:
            return 'other'
    
    def get_extracted_documents(self, limit: int = 100) -> List[Dict[str, Any]]:
        """Get list of extracted documents"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                SELECT ed.id, ed.filename, ed.file_type, ed.file_size,
                       ed.extraction_date, ed.is_resume, ed.is_processed,
                       em.subject, em.sender, em.date_received,
                       ea.email_address, ea.name
                FROM extracted_documents ed
                JOIN email_messages em ON ed.email_message_id = em.id
                JOIN email_accounts ea ON em.account_id = ea.id
                ORDER BY ed.extraction_date DESC
                LIMIT ?
            ''', (limit,))
            
            documents = []
            for row in cursor.fetchall():
                documents.append({
                    'id': row[0],
                    'filename': row[1],
                    'file_type': row[2],
                    'file_size': row[3],
                    'extraction_date': row[4],
                    'is_resume': bool(row[5]),
                    'is_processed': bool(row[6]),
                    'email_subject': row[7],
                    'email_sender': row[8],
                    'email_date': row[9],
                    'account_email': row[10],
                    'account_name': row[11]
                })
            
            return documents
    
    def get_document_content(self, document_id: str) -> Optional[Dict[str, Any]]:
        """Get full document content and metadata"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                SELECT filename, file_type, content_text, metadata,
                       extraction_date, is_resume, processing_results
                FROM extracted_documents 
                WHERE id = ?
            ''', (document_id,))
            
            row = cursor.fetchone()
            if row:
                return {
                    'filename': row[0],
                    'file_type': row[1],
                    'content_text': row[2],
                    'metadata': json.loads(row[3]) if row[3] else {},
                    'extraction_date': row[4],
                    'is_resume': bool(row[5]),
                    'processing_results': json.loads(row[6]) if row[6] else None
                }
            
            return None
    
    # Utility methods
    def _decode_header(self, header_value: str) -> str:
        """Decode email header values"""
        if not header_value:
            return ""
        
        try:
            decoded_header = email.header.decode_header(header_value)
            decoded_string = ""
            
            for part, encoding in decoded_header:
                if isinstance(part, bytes):
                    decoded_string += part.decode(encoding or 'utf-8', errors='ignore')
                else:
                    decoded_string += part
            
            return decoded_string
        except:
            return str(header_value)
    
    def _simple_encrypt(self, password: str) -> str:
        """Simple password encryption (use proper encryption in production)"""
        import base64
        return base64.b64encode(password.encode()).decode()
    
    def _simple_decrypt(self, encrypted_password: str) -> str:
        """Simple password decryption"""
        import base64
        return base64.b64decode(encrypted_password.encode()).decode()
    
    def _start_sync_log(self, account_id: str) -> str:
        """Start a sync operation log"""
        log_id = hashlib.md5(f"{account_id}_{datetime.now()}".encode()).hexdigest()
        
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO email_sync_logs (id, account_id, sync_start, status)
                VALUES (?, ?, ?, ?)
            ''', (log_id, account_id, datetime.now(), 'running'))
            conn.commit()
        
        return log_id
    
    def _complete_sync_log(self, log_id: str, processed: int, extracted: int, errors: int):
        """Complete a sync operation log"""
        status = 'completed' if errors == 0 else 'completed_with_errors'
        
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                UPDATE email_sync_logs 
                SET sync_end = ?, messages_processed = ?, documents_extracted = ?,
                    errors_count = ?, status = ?
                WHERE id = ?
            ''', (datetime.now(), processed, extracted, errors, status, log_id))
            conn.commit()
    
    def _update_last_sync(self, account_id: str):
        """Update last sync timestamp for account"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                UPDATE email_accounts SET last_sync = ?, updated_at = ?
                WHERE id = ?
            ''', (datetime.now(), datetime.now(), account_id))
            conn.commit()
    
    def _count_resumes_in_email(self, email_message_id: str) -> int:
        """Count how many resumes were extracted from a specific email"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                SELECT COUNT(*) FROM extracted_documents 
                WHERE email_message_id = ? AND is_resume = TRUE
            ''', (email_message_id,))
            return cursor.fetchone()[0]
    
    def _email_contains_resume_indicators(self, email_message: email.message.Message) -> bool:
        """
        Quick check if email likely contains resume-related content
        Used for efficient filtering during archive scans
        """
        # Check subject line for resume indicators
        subject = self._decode_header(email_message.get('Subject', '')).lower()
        resume_keywords = [
            'resume', 'cv', 'curriculum vitae', 'application', 'position',
            'job', 'career', 'employment', 'hire', 'candidate', 'profile',
            'qualifications', 'experience', 'skills', 'background'
        ]
        
        subject_score = sum(1 for keyword in resume_keywords if keyword in subject)
        
        # Check sender domain patterns (common job sites, recruiters)
        sender = self._decode_header(email_message.get('From', '')).lower()
        job_domains = [
            'linkedin', 'indeed', 'monster', 'careerbuilder', 'glassdoor',
            'ziprecruiter', 'workday', 'jobvite', 'recruit', 'hr', 'talent'
        ]
        
        domain_score = sum(1 for domain in job_domains if domain in sender)
        
        # Check if email has attachments (basic requirement)
        has_attachments = False
        if email_message.is_multipart():
            for part in email_message.walk():
                if "attachment" in str(part.get("Content-Disposition", "")):
                    has_attachments = True
                    break
        
        # Score-based decision (at least 1 keyword OR job domain + attachments)
        return has_attachments and (subject_score > 0 or domain_score > 0)
    
    def get_ai_enrichment_integration_status(self) -> Dict[str, Any]:
        """Get status of AI enrichment integration and processing"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                
                # Count resumes by processing status
                cursor.execute('''
                    SELECT processing_status, COUNT(*) 
                    FROM extracted_documents 
                    WHERE is_resume = TRUE 
                    GROUP BY processing_status
                ''')
                status_counts = dict(cursor.fetchall())
                
                # Get recent exports
                cursor.execute('''
                    SELECT COUNT(*) FROM extracted_documents 
                    WHERE is_resume = TRUE 
                    AND processing_status = 'exported_to_ai'
                    AND extraction_date >= date('now', '-7 days')
                ''')
                recent_exports = cursor.fetchone()[0]
                
                # Get total resume count
                cursor.execute('''
                    SELECT COUNT(*) FROM extracted_documents WHERE is_resume = TRUE
                ''')
                total_resumes = cursor.fetchone()[0]
                
                return {
                    "total_resumes": total_resumes,
                    "status_breakdown": status_counts,
                    "recent_exports": recent_exports,
                    "ready_for_processing": status_counts.get('pending', 0),
                    "ai_exported": status_counts.get('exported_to_ai', 0),
                    "processing_pipeline_active": True
                }
        
        except Exception as e:
            self.logger.error(f"Error getting AI integration status: {str(e)}")
            return {"error": str(e)}
    
    def extract_email_contacts_from_archive(self, account_id: str) -> Dict[str, Any]:
        """
        Extract email addresses from email contacts and document content
        Creates comprehensive email database for marketing purposes
        """
        try:
            all_contacts = set()
            document_emails = set()
            
            # Get all extracted documents
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute('''
                    SELECT ed.content_text, em.sender, em.recipients, em.date_received
                    FROM extracted_documents ed
                    JOIN email_messages em ON ed.email_message_id = em.id
                    JOIN email_accounts ea ON em.account_id = ea.id
                    WHERE ea.id = ?
                ''', (account_id,))
                
                documents = cursor.fetchall()
            
            contact_database = []
            
            for content_text, sender, recipients, date_received in documents:
                # Extract from sender
                if sender:
                    sender_emails = self._extract_email_addresses_from_content(sender)
                    all_contacts.update(sender_emails)
                
                # Extract from recipients
                if recipients:
                    recipient_emails = self._extract_email_addresses_from_content(recipients)
                    all_contacts.update(recipient_emails)
                
                # Extract from document content
                if content_text:
                    doc_emails = self._extract_email_addresses_from_content(content_text)
                    document_emails.update(doc_emails)
                    
                    # Create contact entries with context
                    for email in doc_emails:
                        contact_database.append({
                            "email": email,
                            "source": "document_content",
                            "context": "extracted_from_attachment",
                            "date_found": date_received,
                            "sender_context": sender
                        })
            
            # Add email contacts with sender context
            for email in all_contacts:
                contact_database.append({
                    "email": email,
                    "source": "email_communication",
                    "context": "email_sender_or_recipient",
                    "date_found": datetime.now().isoformat(),
                    "sender_context": "email_archive"
                })
            
            # Deduplicate and enrich contact database
            unique_contacts = {}
            for contact in contact_database:
                email = contact["email"]
                if email not in unique_contacts:
                    unique_contacts[email] = contact
                else:
                    # Merge contexts if duplicate email
                    existing = unique_contacts[email]
                    if contact["source"] != existing["source"]:
                        existing["context"] += f", {contact['context']}"
            
            final_contacts = list(unique_contacts.values())
            
            return {
                "success": True,
                "total_unique_emails": len(final_contacts),
                "contact_database": final_contacts,
                "document_emails": len(document_emails),
                "communication_emails": len(all_contacts)
            }
            
        except Exception as e:
            self.logger.error(f"Error extracting email contacts: {str(e)}")
            return {"error": str(e)}
    
    def export_to_ai_enrichment_folder(self, include_resumes: bool = True, 
                                     include_job_descriptions: bool = True,
                                     include_email_contacts: bool = True) -> Dict[str, Any]:
        """
        Export all extracted data to Data_forAi_Enrichment_linked_Admin_portal_final folder
        This matches the documented system requirements for AI enrichment integration
        """
        try:
            # Path to AI enrichment folder (as documented)
            ai_enrichment_path = Path("c:/IntelliCV/admin_portal_final/Data_forAi_Enrichment_linked_Admin_portal_final")
            
            # Create subdirectories
            email_data_path = ai_enrichment_path / "email_extracted_data"
            email_data_path.mkdir(parents=True, exist_ok=True)
            
            resumes_path = email_data_path / "resumes"
            jobs_path = email_data_path / "job_descriptions"
            contacts_path = email_data_path / "email_contacts"
            
            resumes_path.mkdir(exist_ok=True)
            jobs_path.mkdir(exist_ok=True)
            contacts_path.mkdir(exist_ok=True)
            
            export_summary = {
                "export_date": datetime.now().isoformat(),
                "exported_items": {},
                "file_locations": {}
            }
            
            # Export resumes if requested
            if include_resumes:
                with sqlite3.connect(self.db_path) as conn:
                    cursor = conn.cursor()
                    cursor.execute('''
                        SELECT ed.id, ed.filename, ed.content_text, ed.metadata,
                               em.sender, em.date_received, ea.email_address
                        FROM extracted_documents ed
                        JOIN email_messages em ON ed.email_message_id = em.id
                        JOIN email_accounts ea ON em.account_id = ea.id
                        WHERE ed.is_resume = TRUE
                    ''')
                    
                    resumes = cursor.fetchall()
                    resume_data = []
                    
                    for resume in resumes:
                        doc_id, filename, content, metadata_json, sender, date_received, account_email = resume
                        
                        # Extract emails from resume content
                        extracted_emails = self._extract_email_addresses_from_content(content)
                        
                        resume_record = {
                            "document_id": doc_id,
                            "original_filename": filename,
                            "content": content,
                            "metadata": json.loads(metadata_json) if metadata_json else {},
                            "extraction_info": {
                                "email_sender": sender,
                                "email_date": date_received,
                                "account": account_email,
                                "extracted_emails": extracted_emails
                            },
                            "document_type": "resume",
                            "processed_date": datetime.now().isoformat()
                        }
                        
                        resume_data.append(resume_record)
                        
                        # Save individual resume file
                        resume_file = resumes_path / f"resume_{doc_id[:8]}.json"
                        with open(resume_file, 'w', encoding='utf-8') as f:
                            json.dump(resume_record, f, indent=2, ensure_ascii=False)
                    
                    # Save master resumes file
                    master_resumes_file = resumes_path / "all_resumes.json"
                    with open(master_resumes_file, 'w', encoding='utf-8') as f:
                        json.dump(resume_data, f, indent=2, ensure_ascii=False)
                    
                    export_summary["exported_items"]["resumes"] = len(resume_data)
                    export_summary["file_locations"]["resumes"] = str(resumes_path)
            
            # Export job descriptions if requested
            if include_job_descriptions:
                with sqlite3.connect(self.db_path) as conn:
                    cursor = conn.cursor()
                    cursor.execute('''
                        SELECT ed.id, ed.filename, ed.content_text, ed.metadata,
                               em.sender, em.date_received, ea.email_address
                        FROM extracted_documents ed
                        JOIN email_messages em ON ed.email_message_id = em.id
                        JOIN email_accounts ea ON em.account_id = ea.id
                        WHERE ed.content_text IS NOT NULL
                    ''')
                    
                    documents = cursor.fetchall()
                    job_descriptions = []
                    
                    for doc in documents:
                        doc_id, filename, content, metadata_json, sender, date_received, account_email = doc
                        
                        # Check if it's a job description
                        if self._is_likely_job_description(content):
                            extracted_emails = self._extract_email_addresses_from_content(content)
                            
                            job_record = {
                                "document_id": doc_id,
                                "original_filename": filename,
                                "content": content,
                                "metadata": json.loads(metadata_json) if metadata_json else {},
                                "extraction_info": {
                                    "email_sender": sender,
                                    "email_date": date_received,
                                    "account": account_email,
                                    "extracted_emails": extracted_emails
                                },
                                "document_type": "job_description",
                                "processed_date": datetime.now().isoformat()
                            }
                            
                            job_descriptions.append(job_record)
                            
                            # Save individual job description file
                            job_file = jobs_path / f"job_{doc_id[:8]}.json"
                            with open(job_file, 'w', encoding='utf-8') as f:
                                json.dump(job_record, f, indent=2, ensure_ascii=False)
                    
                    # Save master job descriptions file
                    if job_descriptions:
                        master_jobs_file = jobs_path / "all_job_descriptions.json"
                        with open(master_jobs_file, 'w', encoding='utf-8') as f:
                            json.dump(job_descriptions, f, indent=2, ensure_ascii=False)
                    
                    export_summary["exported_items"]["job_descriptions"] = len(job_descriptions)
                    export_summary["file_locations"]["job_descriptions"] = str(jobs_path)
            
            # Export email contacts if requested
            if include_email_contacts:
                # Get all email accounts for this extraction
                accounts = self.get_email_accounts()
                all_contacts = []
                
                for account in accounts:
                    contact_result = self.extract_email_contacts_from_archive(account['id'])
                    if contact_result.get('success'):
                        account_contacts = contact_result['contact_database']
                        for contact in account_contacts:
                            contact['source_account'] = account['email_address']
                        all_contacts.extend(account_contacts)
                
                if all_contacts:
                    # Save master contacts file (JSON)
                    contacts_json_file = contacts_path / "email_contacts.json"
                    with open(contacts_json_file, 'w', encoding='utf-8') as f:
                        json.dump(all_contacts, f, indent=2, ensure_ascii=False)
                    
                    # Save CSV format for easy import
                    contacts_csv_file = contacts_path / "email_contacts.csv"
                    with open(contacts_csv_file, 'w', newline='', encoding='utf-8') as f:
                        if all_contacts:
                            writer = csv.DictWriter(f, fieldnames=all_contacts[0].keys())
                            writer.writeheader()
                            writer.writerows(all_contacts)
                
                export_summary["exported_items"]["email_contacts"] = len(all_contacts)
                export_summary["file_locations"]["email_contacts"] = str(contacts_path)
            
            # Save export summary
            summary_file = email_data_path / f"export_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            with open(summary_file, 'w', encoding='utf-8') as f:
                json.dump(export_summary, f, indent=2, ensure_ascii=False)
            
            self.logger.info(f"Data exported to AI enrichment folder: {ai_enrichment_path}")
            
            return {
                "success": True,
                "export_path": str(ai_enrichment_path),
                "summary": export_summary,
                "summary_file": str(summary_file)
            }
            
        except Exception as e:
            self.logger.error(f"Error exporting to AI enrichment folder: {str(e)}")
            return {"error": str(e)}


# Initialize email integration manager
def get_email_manager() -> EmailIntegrationManager:
    """Get or create email integration manager instance"""
    return EmailIntegrationManager()


if __name__ == "__main__":
    # Test the email system
    manager = EmailIntegrationManager()
    print("Email Integration Manager initialized successfully!")
    print(f"Database path: {manager.db_path}")
    print("Ready for email integration!")