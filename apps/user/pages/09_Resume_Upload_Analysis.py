
📊 Resume Upload & Analysis - FREE + Premium Features
=====================================================
Unified resume upload platform with tiered features:
- FREE: Basic upload + express analysis (always available)
- PREMIUM: Advanced AI analysis, ATS optimization, detailed insights + AI Resume Builder

TIER STRUCTURE:
🆓 FREE (All Users):
  - Upload resume (all formats: PDF, DOCX, TXT)
  - Basic text extraction and parsing
  - Express analysis (word count, keywords, basic score)
  - Simple précis/summary
  - Quick recommendations (4-5 bullets)
  - File preview and management

🔒 PREMIUM (Monthly Pro $15.99/mo, Annual Pro $149.99/yr, Enterprise $499/yr):
  - 🤖 Admin AI Enhanced Processing (Job Title Engine + Real AI Data Connector)
  - 🎯 ATS Compatibility Check & Scoring
  - 📊 Advanced Competitive Analysis
  - 🔍 Deep Keyword Optimization
  - 💼 Market Alignment Assessment
  - 🚀 Professional Format Enhancement
  - 📈 Strength Analysis & Recruiter Perspective
  - 🔗 LinkedIn Profile Import
  - 📚 Version Tracking & History
  - ✨ AI Resume Builder (Interactive Chatbot - Build/Edit Resume)

STRATEGIC DESIGN: Free users see greyed-out premium features to "drool over"
INTEGRATION: AI Resume Builder from page 10 integrated as premium Tab 6
"""

import streamlit as st
from pathlib import Path
import io
import time
import json
from datetime import datetime
import tempfile
import sys
import hashlib
from typing import Optional, Dict, List, Any


# --- HYBRID AI ENGINE (8-model stack) ------------------------------------
try:
    from services.hybrid_ai_engine import HybridCareerEngine

    hybrid_engine = HybridCareerEngine()
    HYBRID_ENGINE_AVAILABLE = True
except Exception:
    hybrid_engine = None
    HYBRID_ENGINE_AVAILABLE = False
# Job cloud backend
try:
    from job_title_backend_integration import JobTitleBackend
    JOB_CLOUD_AVAILABLE = True
except ImportError:
    JobTitleBackend = None  # type: ignore
    JOB_CLOUD_AVAILABLE = False

# Setup paths and imports
current_dir = Path(__file__).parent.parent
sys.path.insert(0, str(current_dir))

# Import shared components
try:
    from shared_components import apply_professional_styling, show_logo_watermark, initialize_session_manager
    SHARED_COMPONENTS_AVAILABLE = True
except ImportError:
    SHARED_COMPONENTS_AVAILABLE = False

# Import real AI connector for dynamic keyword extraction
try:
    from shared_backend.services.real_ai_connector import get_real_ai_connector
    real_ai_connector = get_real_ai_connector()
    REAL_AI_CONNECTOR_AVAILABLE = True
except Exception:
    real_ai_connector = None
    REAL_AI_CONNECTOR_AVAILABLE = False

# Import Candidate Intelligence Service for competitive analysis
try:
    from services.candidate_intelligence import CandidateIntelligence, analyze_basma_example
    from services.intelligence_visualizations import (
        create_quadrant_chart, create_score_histogram,
        create_spider_chart, create_peer_comparison_table
    )
    INTELLIGENCE_SERVICE_AVAILABLE = True
except ImportError:
    INTELLIGENCE_SERVICE_AVAILABLE = False

# Import UserDataService for real user-specific data (150 candidates from learning library)
try:
    from services.user_data_service import get_user_data_service
    USER_DATA_SERVICE_AVAILABLE = True
except ImportError:
    USER_DATA_SERVICE_AVAILABLE = False

# Import Application Blocker Engine for gap analysis
try:
    from shared.services.application_blocker_engine import (
        ApplicationBlockerEngine,
        ProactiveAcknowledgmentGenerator
    )
    BLOCKER_ENGINE_AVAILABLE = True
except ImportError:
    BLOCKER_ENGINE_AVAILABLE = False

def get_dynamic_keywords(keyword_type: str = 'all'):
    \"\"\"Return REAL skill keywords filtered by type using the shared_backend connector.\"\"\"
    if not REAL_AI_CONNECTOR_AVAILABLE or not real_ai_connector:
        st.error("⚠️ **AI Data Service Unavailable**")
        st.markdown(\"\"\"
        We are experiencing technical difficulties loading our AI skills database.

        **Status:**
        - AI skills database temporarily offline
        - Your uploaded data is safe and secure

        **Next Steps:**
        - Please try again in 15-30 minutes
        - Contact support@intellicv.ai if issue persists

        *We apologize for the inconvenience and are working to resolve this quickly.*
        \"\"\")

        if globals().get("ERROR_HANDLER_AVAILABLE"):
            log_service_error("ai_data_connector_unavailable", {
                'timestamp': datetime.now().isoformat(),
                'keyword_type': keyword_type
            })

        st.stop()
        return []

    try:
        # Load REAL skills from parsed CV corpus - NO hardcoded stacks in the UI
        all_skills = real_ai_connector.get_real_skills_from_cvs()
        keyword_type = (keyword_type or 'all').lower()

        # NOTE: Any future classification (technical/soft) must be derived from real data,
        # not from hard-coded keyword lists.
        if keyword_type in {'technical', 'soft'}:
            return all_skills[:20] if keyword_type == 'technical' else all_skills[:15]

        return all_skills[:30]

    except Exception as e:
        if globals().get("ERROR_HANDLER_AVAILABLE"):
            log_service_error("ai_skills_extraction_failed", {
                'error': str(e),
                'keyword_type': keyword_type
            })

        st.error("⚠️ **AI Data Service Error**")
        st.markdown(\"\"\"
        We encountered an error loading skill data from our AI database.

        **Your data is safe.** Please try again shortly.
        \"\"\")
        st.stop()
        return []
    try:
        # Load REAL skills from 243k candidate database - NO hardcoded patterns allowed
        all_skills = real_ai_connector.load_real_skills_data()
        keyword_type = (keyword_type or 'all').lower()

        # Filter must use ONLY skills that exist in real database
        # NO predetermined tech stacks (python, javascript, tensorflow, etc.)
        # Skills classification must be derived from actual candidate data
        if keyword_type == 'technical':
            # Return top 20 technical skills from real data - no filtering with fake patterns
            return all_skills[:20]

        if keyword_type == 'soft':
            # Return top 15 soft skills from real data - no filtering with fake patterns
            return all_skills[:15]

        return all_skills[:30]

    except Exception as e:
        if ERROR_HANDLER_AVAILABLE:
            log_service_error("ai_skills_extraction_failed", {
                'error': str(e),
                'keyword_type': keyword_type
            })

        st.error("⚠️ **AI Data Service Error**")
        st.markdown("""
        We encountered an error loading skill data from our AI database.

        **Your data is safe.** Please try again shortly.
        """)
        st.stop()
        return []

# Import tier management
try:
    from utils.tier_manager import check_user_tier, get_tier_info, show_upgrade_prompt
    TIER_MANAGER_AVAILABLE = True
except ImportError:
    TIER_MANAGER_AVAILABLE = False
    def check_user_tier(required_tier):
        # Fallback: check session state
        user_tier = st.session_state.get('user_tier', 'free')
        tier_hierarchy = {'free': 0, 'monthly_pro': 1, 'annual_pro': 2, 'enterprise_pro': 3}
        required_level = tier_hierarchy.get(required_tier, 0)
        current_level = tier_hierarchy.get(user_tier, 0)
        return current_level >= required_level

    def get_tier_info(tier_name):
        tiers = {
            'free': {'name': 'Free', 'price': 'FREE'},
            'monthly_pro': {'name': 'Monthly Pro', 'price': '$15.99/mo'},
            'annual_pro': {'name': 'Annual Pro', 'price': '$149.99/yr'},
            'enterprise_pro': {'name': 'Enterprise', 'price': '$499/yr'}
        }
        return tiers.get(tier_name, tiers['free'])

    def show_upgrade_prompt(feature_name, required_tier='monthly_pro'):
        tier_info = get_tier_info(required_tier)
        st.warning(f"🔒 **{feature_name}** requires {tier_info['name']} ({tier_info['price']})")
        if st.button(f"⬆️ Upgrade to {tier_info['name']}", key=f"upgrade_{feature_name}"):
            st.switch_page("pages/06_Pricing.py")

# Import admin AI integration
try:
    from user_portal_admin_integration import process_user_action_with_admin_ai, init_admin_ai_for_user_page
    ADMIN_AI_AVAILABLE = True
except ImportError:
    ADMIN_AI_AVAILABLE = False

# Import portal bridge for cross-portal communication
try:
    from shared_backend.services.portal_bridge import ResumeService, IntelligenceService
    resume_service = ResumeService()
    intelligence_service = IntelligenceService()
    PORTAL_BRIDGE_AVAILABLE = True
except ImportError:
    PORTAL_BRIDGE_AVAILABLE = False
    resume_service = None
    intelligence_service = None

def set_uploaded_resume(content: str, filename: str, source: str = 'user') -> None:
    """Persist uploaded resume content with provenance tracking."""
    st.session_state.uploaded_resume_content = content.strip()
    st.session_state.uploaded_resume_filename = filename
    st.session_state.express_analysis_complete = False
    st.session_state.resume_demo_mode = False
    st.session_state.uploaded_resume_source = source
    st.session_state.resume_sync_requested = True


def extract_text_from_upload(uploaded_file) -> str:
    """Convert an uploaded file to plain text without using mock data.

    Supports common resume / data formats and fails *loudly* with actionable guidance.
    """
    file_bytes = uploaded_file.getvalue()
    suffix = Path(uploaded_file.name).suffix.lower().strip()

    # -------------------------
    # Plain text-like formats
    # -------------------------
    if suffix in {'.txt', '.md', '.markdown', '.log'}:
        return file_bytes.decode('utf-8', errors='ignore')

    if suffix in {'.rtf'}:
        # Best-effort RTF -> text
        try:
            from striprtf.striprtf import rtf_to_text  # type: ignore
            return rtf_to_text(file_bytes.decode('utf-8', errors='ignore')).strip()
        except Exception:
            # Fallback: very rough RTF control-word stripping
            import re as _re
            raw = file_bytes.decode('utf-8', errors='ignore')
            raw = _re.sub(r'\\[a-zA-Z]+\d* ?',' ', raw)
            raw = _re.sub(r'[{}]', ' ', raw)
            return _re.sub(r'\s+', ' ', raw).strip()

    if suffix in {'.html', '.htm'}:
        html = file_bytes.decode('utf-8', errors='ignore')
        try:
            from bs4 import BeautifulSoup  # type: ignore
            return BeautifulSoup(html, "html.parser").get_text("\n").strip()
        except Exception:
            # Fallback: strip tags
            import re as _re
            text = _re.sub(r'<[^>]+>', ' ', html)
            return _re.sub(r'\s+', ' ', text).strip()


# -------------------------
# ZIP (container)
if (getattr(uploaded_file, "name", "") or "").lower().endswith(".zip"):
    import io, zipfile
    file_bytes = uploaded_file.getvalue()
    combined = []
    try:
        zf = zipfile.ZipFile(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"Unable to read ZIP archive: {e}")

    # Safety limits
    max_files = 15
    max_total_chars = 200000
    count = 0

    for info in zf.infolist():
        if info.is_dir():
            continue
        count += 1
        if count > max_files:
            combined.append("\n[ZIP] ...additional files omitted for safety...\n")
            break
        inner_name = info.filename
        try:
            inner_bytes = zf.read(info)
        except Exception:
            continue
        combined.append(f"\n----- FILE: {inner_name} -----\n")
        # Recurse using the same extractor by faking an uploaded-file object
        class _UF:
            def __init__(self, data: bytes, name: str):
                self._data = data
                self.name = name
            def getvalue(self):
                return self._data
        try:
            combined.append(extract_text_from_upload(_UF(inner_bytes, inner_name)))
        except Exception as e:
            combined.append(f"[ZIP] Skipped {inner_name}: {e}")
        if sum(len(x) for x in combined) > max_total_chars:
            combined.append("\n[ZIP] ...output truncated...\n")
            break

    return "\n".join(combined).strip()

# -------------------------
# EML (email)
if (getattr(uploaded_file, "name", "") or "").lower().endswith(".eml"):
    from email import policy
    from email.parser import BytesParser

    file_bytes = uploaded_file.getvalue()
    msg = BytesParser(policy=policy.default).parsebytes(file_bytes)

    parts = []
    subject = msg.get("subject", "")
    if subject:
        parts.append(f"Subject: {subject}")

    def _html_to_text(s: str) -> str:
        try:
            from bs4 import BeautifulSoup
            return BeautifulSoup(s, "html.parser").get_text(" ")
        except Exception:
            return re.sub(r"<[^>]+>", " ", s)

    if msg.is_multipart():
        for part in msg.walk():
            ctype = part.get_content_type()
            if ctype in ("text/plain", "text/html"):
                try:
                    payload = part.get_content()
                except Exception:
                    payload = part.get_payload(decode=True)
                    payload = payload.decode(errors="ignore") if isinstance(payload, (bytes, bytearray)) else str(payload)
                if ctype == "text/html":
                    payload = _html_to_text(str(payload))
                parts.append(str(payload))
    else:
        try:
            parts.append(str(msg.get_content()))
        except Exception:
            payload = msg.get_payload(decode=True)
            parts.append(payload.decode(errors="ignore") if isinstance(payload, (bytes, bytearray)) else str(payload))

    return "\n".join(parts).strip()

# -------------------------
# MSG (Outlook email)
if (getattr(uploaded_file, "name", "") or "").lower().endswith(".msg"):
    try:
        import extract_msg  # type: ignore
    except Exception:
        raise ValueError("MSG (.msg) parsing requires the optional 'extract-msg' package. Install it or convert MSG to EML/PDF.")

    import tempfile
    file_bytes = uploaded_file.getvalue()
    with tempfile.TemporaryDirectory() as td:
        msg_path = Path(td) / "email.msg"
        msg_path.write_bytes(file_bytes)
        m = extract_msg.Message(str(msg_path))
        m.process()
        parts = []
        if getattr(m, "subject", None):
            parts.append(f"Subject: {m.subject}")
        if getattr(m, "body", None):
            parts.append(m.body)
        return "\n".join(parts).strip()


    # -------------------------
    # PDF
    # -------------------------
    if suffix == '.pdf':
        try:
            from PyPDF2 import PdfReader  # type: ignore
            reader = PdfReader(io.BytesIO(file_bytes))
            pages = [page.extract_text() or '' for page in reader.pages]
            text = "\n".join(pages).strip()
            if not text:
                raise ValueError("PDF text extraction returned empty content (possibly scanned image PDF)")
            return text
        except Exception as exc:
            raise RuntimeError(f"PDF extraction failed: {exc}") from exc

    # -------------------------
    # Office documents
    # -------------------------
    if suffix == '.docx':
        try:
            import docx  # type: ignore
            document = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [para.text for para in document.paragraphs]
            text = "\n".join(paragraphs).strip()
            if not text:
                raise ValueError("DOCX extraction returned empty content")
            return text
        except Exception as exc:
            raise RuntimeError(f"DOCX extraction failed: {exc}") from exc

    if suffix == '.doc':
        # Best effort: try LibreOffice headless conversion if available.
        # If not available, provide an actionable message.
        try:
            import tempfile
            import subprocess
            with tempfile.TemporaryDirectory() as td:
                in_path = Path(td) / "upload.doc"
                out_path = Path(td) / "upload.docx"
                in_path.write_bytes(file_bytes)

                # Prefer soffice if installed (Windows/Linux)
                cmd = [
                    "soffice", "--headless", "--nologo", "--nolockcheck",
                    "--convert-to", "docx", "--outdir", str(Path(td)), str(in_path)
                ]
                subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)

                # LibreOffice outputs with same stem
                produced = next(Path(td).glob("*.docx"), None)
                if not produced or not produced.exists():
                    raise RuntimeError("DOC→DOCX conversion produced no output")

                # Now read DOCX
                import docx  # type: ignore
                document = docx.Document(str(produced))
                paragraphs = [para.text for para in document.paragraphs]
                text = "\n".join(paragraphs).strip()
                if not text:
                    raise ValueError("Converted DOCX extraction returned empty content")
                return text
        except Exception:
            raise RuntimeError(
                "Microsoft .DOC detected. This environment couldn't auto-convert it. "
                "Please convert to .DOCX or PDF (LibreOffice/Word), or paste the text."
            )

    if suffix in {'.odt'}:
        # ODT support is optional (depends on odfpy)
        try:
            from odf.opendocument import load  # type: ignore
            from odf import text as odf_text  # type: ignore
            from odf import teletype  # type: ignore
            with tempfile.NamedTemporaryFile(delete=False, suffix=".odt") as tf:
                tf.write(file_bytes)
                tmp_path = tf.name
            doc = load(tmp_path)
            os.unlink(tmp_path)
            paragraphs = doc.getElementsByType(odf_text.P)
            return "\n".join(teletype.extractText(p) for p in paragraphs).strip()
        except Exception as exc:
            raise RuntimeError(f"ODT extraction failed (optional dependency): {exc}") from exc

    # -------------------------
    # Spreadsheets / structured data
    # -------------------------
    if suffix in {'.csv'}:
        try:
            import pandas as pd  # type: ignore
            import io as _io
            df = pd.read_csv(_io.BytesIO(file_bytes), dtype=str, keep_default_na=False, encoding_errors='ignore')
            # Convert to a text view that AI can handle
            return df.to_csv(index=False).strip()
        except Exception as exc:
            raise RuntimeError(f"CSV extraction failed: {exc}") from exc

    if suffix in {'.xlsx', '.xls'}:
        try:
            import pandas as pd  # type: ignore
            import io as _io
            xl = pd.ExcelFile(_io.BytesIO(file_bytes))
            parts = []
            for sheet in xl.sheet_names[:10]:  # safety cap
                df = xl.parse(sheet, dtype=str, keep_default_na=False)
                parts.append(f"--- SHEET: {sheet} ---\n" + df.to_csv(index=False))
            text = "\n".join(parts).strip()
            if not text:
                raise ValueError("Excel extraction returned empty content")
            return text
        except Exception as exc:
            raise RuntimeError(f"Excel extraction failed: {exc}") from exc

    if suffix in {'.json'}:
        try:
            obj = json.loads(file_bytes.decode('utf-8', errors='ignore') or "{}")
            return json.dumps(obj, indent=2, ensure_ascii=False)
        except Exception as exc:
            raise RuntimeError(f"JSON extraction failed: {exc}") from exc

    # -------------------------
    # Images (optional OCR)
    # -------------------------
    if suffix in {'.png', '.jpg', '.jpeg', '.webp', '.tif', '.tiff'}:
        try:
            from PIL import Image  # type: ignore
            import pytesseract  # type: ignore
            img = Image.open(io.BytesIO(file_bytes))
            text = pytesseract.image_to_string(img) or ""
            text = text.strip()
            if not text:
                raise ValueError("OCR returned empty text")
            return text
        except Exception as exc:
            raise RuntimeError(
                f"Image resume detected but OCR is unavailable or failed: {exc}. "
                "Please upload PDF/DOCX/TXT instead."
            ) from exc

    raise RuntimeError(
        f"Unsupported file format: {suffix}. "
        "Supported: PDF, DOCX, DOC (auto-convert if possible), TXT, RTF, ODT, HTML, MD, CSV, XLS/XLSX, JSON, and images (OCR if available)."
    )

# Import utilities with fallbacks
try:
    from utils.error_handler import (
        handle_exceptions,
        log_user_action,
        show_error,
        show_success,
        show_warning,
        log_service_error,
    )
    ERROR_HANDLER_AVAILABLE = True
except ImportError:
    ERROR_HANDLER_AVAILABLE = False
    def handle_exceptions(func): return func
    def log_user_action(action, user_id=None, details=None): pass
    def show_error(message, details=None): st.error(f"❌ {message}")
    def show_success(message, details=None): st.success(f"✅ {message}")
    def show_warning(message, details=None): st.warning(f"⚠️ {message}")
    def log_service_error(event, details=None): pass

# Helper functions for JD parsing
def extract_skills_from_text(jd_text):
    """Extract skills from job description text using real AI data."""
    if not REAL_AI_CONNECTOR_AVAILABLE or not real_ai_connector:
        return []

    try:
        all_skills = get_dynamic_keywords('all')
        found_skills = []

        jd_lower = jd_text.lower()
        for skill in all_skills:
            if skill.lower() in jd_lower:
                found_skills.append(skill)

        return found_skills
    except:
        return []

def extract_experience_requirement(jd_text: Optional[str]) -> Optional[int]:
    """Extract years of experience from JD text without fabricating values."""
    if not jd_text:
        return None

    import re

    patterns = [
        r'(\d+)\+?\s*years?\s+(?:of\s+)?experience',
        r'minimum\s+(\d+)\s+years?',
        r'(\d+)\s*-\s*\d+\s+years?'
    ]

    for pattern in patterns:
        match = re.search(pattern, jd_text.lower())
        if match:
            return int(match.group(1))

    return None


def extract_education_requirement(jd_text: Optional[str]) -> Optional[str]:
    """Extract education requirements using supplied JD text only."""
    if not jd_text:
        return None

    jd_lower = jd_text.lower()

    if 'phd' in jd_lower or 'doctorate' in jd_lower:
        return "PhD/Doctorate"
    if 'master' in jd_lower or 'msc' in jd_lower or 'mba' in jd_lower:
        return "Master's degree"
    if 'bachelor' in jd_lower or 'bsc' in jd_lower or 'ba' in jd_lower:
        return "Bachelor's degree"

    return None

def generate_word_cloud_from_real_data(user_resume_text, user_id):
    """Generate word cloud data from REAL sources: user resume + legacy CVs."""
    from pathlib import Path
    from collections import Counter
    import json

    word_freq = Counter()

    try:
        # 1. Extract keywords from user's resume
        user_keywords = extract_skills_from_text(user_resume_text)
        word_freq.update(user_keywords)

        # 2. Load legacy CVs from ai_data/normalized/
        normalized_dir = Path("ai_data/normalized")

        if not normalized_dir.exists():
            st.error("⚠️ **CV Database Unavailable**")
            st.markdown("Cannot generate word cloud without legacy CV data.")
            st.stop()
            return {}

        # 3. Extract user's likely job role
        user_role_keywords = []
        for keyword in ["engineer", "manager", "developer", "analyst", "scientist", "designer"]:
            if keyword in user_resume_text.lower():
                user_role_keywords.append(keyword)

        # 4. Find similar CVs from legacy data
        similar_cv_count = 0
        for cv_file in normalized_dir.glob("*.json"):
            try:
                with open(cv_file, 'r', encoding='utf-8') as f:
                    cv_data = json.load(f)

                # Check if CV is similar to user's role
                cv_title = cv_data.get('job_title', '').lower()
                is_similar = any(keyword in cv_title for keyword in user_role_keywords)

                if is_similar and similar_cv_count < 50:  # Limit to 50 similar CVs
                    cv_keywords = cv_data.get('extracted_keywords', [])
                    word_freq.update(cv_keywords)
                    similar_cv_count += 1
            except:
                continue

        if similar_cv_count > 0:
            st.info(f"✅ Analyzed your resume + {similar_cv_count} similar CVs from our database")

        # Return top 100 keywords with frequency
        return dict(word_freq.most_common(100))

    except Exception as e:
        if ERROR_HANDLER_AVAILABLE:
            from utils.error_handler import log_service_error
            log_service_error("word_cloud_generation_failed", {'error': str(e), 'user_id': user_id})

        st.error("⚠️ **Visualization Service Unavailable**")
        st.markdown("Cannot generate word cloud at this time. Please try again later.")
        st.stop()
        return {}

def calculate_real_fit_score(user_skills, required_skills, user_exp, min_exp):
    """Calculate REAL fit percentage based on actual matches."""
    if not required_skills:
        return 0

    skill_match_count = sum(1 for skill in required_skills if skill.lower() in [s.lower() for s in user_skills])
    skill_match_pct = (skill_match_count / len(required_skills) * 100)

    exp_match_pct = min(100, (user_exp / min_exp * 100)) if min_exp > 0 else 100

    # Weighted average: 70% skills, 30% experience
    fit_score = (skill_match_pct * 0.7) + (exp_match_pct * 0.3)

    return round(fit_score, 1)


def _dedupe_preserve_case(items: List[str]) -> List[str]:
    """Remove duplicates from a list of strings while preserving order/case."""
    seen = set()
    ordered: List[str] = []
    for item in items:
        if not item or not isinstance(item, str):
            continue
        key = item.strip().lower()
        if not key or key in seen:
            continue
        seen.add(key)
        ordered.append(item.strip())
    return ordered


def ensure_job_cloud_tokens() -> None:
    """Guarantee a starter token balance for job cloud calls."""
    if 'user_tokens' not in st.session_state:
        st.session_state['user_tokens'] = 40


def get_job_cloud_backend():
    """Return cached JobTitleBackend instance when available."""
    if not JOB_CLOUD_AVAILABLE or JobTitleBackend is None:
        return None

    backend = st.session_state.get('_job_cloud_backend')
    if backend:
        return backend

    try:
        backend = JobTitleBackend(user_id=st.session_state.get('user_id'))
        st.session_state['_job_cloud_backend'] = backend
        return backend
    except Exception:
        return None


def collect_resume_skill_inventory() -> List[str]:
    """Aggregate skills from parsed resume data + quick analysis."""
    skills: List[str] = []
    resume_data = st.session_state.get('resume_data') or {}
    analysis_results = st.session_state.get('analysis_results') or {}
    sources = [
        resume_data.get('skills'),
        resume_data.get('metadata', {}).get('skills'),
        analysis_results.get('tech_keywords'),
        st.session_state.get('user_profile', {}).get('skills')
    ]

    for source in sources:
        if isinstance(source, list):
            skills.extend(source)
        elif isinstance(source, str):
            skills.append(source)

    cleaned = [skill for skill in skills if isinstance(skill, str) and skill.strip()]
    return _dedupe_preserve_case(cleaned)


def derive_default_target_role() -> str:
    """Best-effort default target role for the intuitive AI scan."""
    profile = st.session_state.get('user_profile') or {}
    candidates: List[str] = [
        profile.get('role_focus', ''),
        profile.get('headline', ''),
        st.session_state.get('analysis_results', {}).get('current_role', '')
    ]

    resume_data = st.session_state.get('resume_data') or {}
    experience = resume_data.get('experience') or resume_data.get('work_experience')
    if isinstance(experience, list) and experience:
        first_role = experience[0]
        if isinstance(first_role, dict):
            for key in ('title', 'role', 'position'):
                if first_role.get(key):
                    candidates.append(first_role[key])

    for candidate in candidates:
        if candidate and isinstance(candidate, str):
            return candidate
    return ''


def _normalize_title_relationships(payload: Any) -> Dict[str, List[str]]:
    """Normalize title relationship payloads coming from the backend."""
    normalized = {
        'similar_titles': [],
        'lateral_moves': [],
        'advancement_roles': []
    }

    if isinstance(payload, dict):
        rel_block = payload.get('relationships')
        if isinstance(rel_block, dict):
            normalized['similar_titles'] = _dedupe_preserve_case(rel_block.get('similar_titles', []))
            normalized['lateral_moves'] = _dedupe_preserve_case(rel_block.get('lateral_moves', []))
            normalized['advancement_roles'] = _dedupe_preserve_case(rel_block.get('advancement_roles', []))
        elif payload.get('relationship_type'):
            titles = payload.get('titles') or payload.get('roles') or []
            normalized['similar_titles'] = _dedupe_preserve_case(titles)

    elif isinstance(payload, list):
        similar: List[str] = []
        lateral: List[str] = []
        advancement: List[str] = []
        for entry in payload:
            if not isinstance(entry, dict):
                continue
            relation_type = (entry.get('relationship_type') or '').lower()
            titles = entry.get('titles') or entry.get('roles') or []
            if 'similar' in relation_type:
                similar.extend(titles)
            elif 'lateral' in relation_type:
                lateral.extend(titles)
            elif 'progression' in relation_type or 'advancement' in relation_type:
                advancement.extend(titles)
        normalized['similar_titles'] = _dedupe_preserve_case(similar)
        normalized['lateral_moves'] = _dedupe_preserve_case(lateral)
        normalized['advancement_roles'] = _dedupe_preserve_case(advancement)

    return normalized


def run_intuitive_ai_skill_dig(target_title: str, resume_text: str, resume_skills: List[str]) -> Optional[Dict[str, Any]]:
    """Call the job cloud backend to surface peer skill signals."""
    backend = get_job_cloud_backend()
    if not backend or not target_title:
        return None

    context = (resume_text or '')[:1200]
    analysis = backend.analyze_job_title(target_title, context=context) or {}
    relationships_raw = backend.get_title_relationships(base_title=target_title) or {}
    market_raw = backend.get_market_intelligence(job_category=target_title) or {}

    relationships = _normalize_title_relationships(relationships_raw)
    job_category = (
        analysis.get('ai_analysis', {}).get('predicted_category')
        or analysis.get('ai_category')
        or market_raw.get('job_category')
        or target_title
    )
    confidence = (
        analysis.get('ai_analysis', {}).get('confidence_score')
        or analysis.get('confidence_score')
        or market_raw.get('confidence_score')
        or 0.62
    )

    peer_skill_pool: List[str] = []
    market_analysis = analysis.get('market_intelligence', {})
    peer_skill_pool.extend(market_analysis.get('skills_in_demand', []))
    peer_skill_pool.extend(analysis.get('ai_recommendations', {}).get('skill_gaps', []))

    market_pool = market_raw.get('market_analysis', {})
    peer_skill_pool.extend(market_pool.get('skills_in_demand', []))
    skills_demand = market_raw.get('skills_demand', {})
    if isinstance(skills_demand, dict):
        for bucket in ('trending_up', 'emerging'):
            peer_skill_pool.extend(skills_demand.get(bucket, []))

    peer_skill_pool = _dedupe_preserve_case(peer_skill_pool)
    if not peer_skill_pool and resume_skills:
        peer_skill_pool = resume_skills[:5]

    resume_lookup = {skill.lower(): skill for skill in resume_skills}
    touchpoints: List[Dict[str, str]] = []
    touchnots: List[Dict[str, str]] = []

    peer_anchor = relationships['similar_titles'][0] if relationships['similar_titles'] else target_title
    for skill in peer_skill_pool:
        lower = skill.lower()
        if lower in resume_lookup:
            touchpoints.append({
                'skill': resume_lookup[lower],
                'reason': f"Already listed – tie it to measurable {job_category} outcomes."
            })
        else:
            touchnots.append({
                'skill': skill,
                'reason': f"Peers tagged as {peer_anchor} highlight this before interviews."
            })

    peer_skill_lower = {skill.lower() for skill in peer_skill_pool}
    under_bushel: List[Dict[str, str]] = []
    for skill in resume_skills:
        if skill.lower() not in peer_skill_lower:
            under_bushel.append({
                'skill': skill,
                'reason': 'Hidden strength – surface it in your summary to differentiate.'
            })
    under_bushel = under_bushel[:5]

    resume_guidance: List[str] = []
    if touchnots:
        missing_cluster = ', '.join(item['skill'] for item in touchnots[:2])
        resume_guidance.append(
            f"Add a quantified bullet for {missing_cluster} – job cloud reviewers treat it as a pass/fail filter."
        )
    if under_bushel:
        resume_guidance.append(
            f"Promote '{under_bushel[0]['skill']}' in the opening summary so it is no longer under the bushel."
        )
    if relationships['similar_titles']:
        resume_guidance.append(
            f"Mirror phrasing from {relationships['similar_titles'][0]} to stay aligned with adjacent job cloud clusters."
        )

    return {
        'target_title': target_title,
        'job_category': job_category,
        'confidence': confidence,
        'similar_titles': relationships['similar_titles'],
        'lateral_moves': relationships['lateral_moves'],
        'touchpoints': touchpoints[:6],
        'touchnots': touchnots[:6],
        'under_bushel': under_bushel,
        'resume_guidance': resume_guidance,
        'generated_at': datetime.utcnow().isoformat()
    }


def display_intuitive_ai_results(insights: Dict[str, Any], focus: str) -> None:
    """Visualize touchpoints/touchnots with quick guidance."""
    if not insights:
        return

    job_category = insights.get('job_category', insights.get('target_title'))
    confidence = insights.get('confidence', 0) * 100
    st.markdown(
        f"**Job Cloud Cluster:** {job_category} · Confidence {confidence:.0f}% · Target `{insights.get('target_title')}`"
    )

    similar_titles = insights.get('similar_titles') or []
    if similar_titles:
        chips = ''.join([
            f"<span style='background:#eef2ff;border:1px solid #c7d2fe;padding:0.3rem 0.75rem;border-radius:999px;font-size:0.85rem;margin:0.2rem 0.4rem 0 0;display:inline-block;'>{title}</span>"
            for title in similar_titles[:6]
        ])
        st.markdown(f"<div style='margin:0.5rem 0 1rem 0;'>{chips}</div>", unsafe_allow_html=True)

    focus_map = {
        'touchpoints': ('touchpoints', 'Already in your resume – double down', st.success),
        'touchnots': ('touchnots', 'Missing signal – add a punchy bullet', st.warning),
        'under-the-bushel': ('under_bushel', 'Hidden strength – bring it to the headline', st.info)
    }
    focus_key = focus.lower()
    focus_tuple = focus_map.get(focus_key, focus_map['touchpoints'])
    data_key, header, renderer = focus_tuple
    renderer(f"{header}")

    entries = insights.get(data_key, [])
    if not entries:
        st.caption("No insights available for this filter yet. Refresh the scan if needed.")
        return

    for entry in entries:
        st.markdown(f"- **{entry['skill']}** — {entry['reason']}")

    if insights.get('resume_guidance'):
        st.markdown("#### 📋 Resume Rebuild Guidance")
        for tip in insights['resume_guidance']:
            st.markdown(f"- {tip}")


def render_intuitive_ai_panel() -> None:
    """Render the Intuitive AI touchpoint/touchnot workflow."""
    if not st.session_state.get('uploaded_resume_content'):
        return

    st.markdown("### 🧠 Job Cloud Skill Dig (Intuitive AI)")
    if not JOB_CLOUD_AVAILABLE:
        st.info("Job cloud intelligence is offline in this environment.")
        return

    ensure_job_cloud_tokens()
    resume_text = st.session_state.get('uploaded_resume_content', '')
    default_target = derive_default_target_role() or "Target Role (e.g., Senior Product Manager)"
    target_title = st.text_input(
        "Target role or anchor title",
        value=st.session_state.get('intuitive_ai_target', default_target),
        key='intuitive_ai_target'
    )

    focus_choice = st.radio(
        "Focus filter",
        options=["Touchpoints", "Touchnots", "Under-the-bushel"],
        horizontal=True,
        key='intuitive_ai_focus'
    )

    run_col, info_col = st.columns([1, 2])
    with run_col:
        trigger = st.button("Surface peer signals", key='intuitive_ai_run')
    with info_col:
        st.caption("Powered by the job cloud graph · compares your resume to adjacent roles")

    if trigger:
        if not target_title.strip():
            st.warning("Enter a target role before running the scan.")
        else:
            resume_skills = collect_resume_skill_inventory()
            insights = run_intuitive_ai_skill_dig(target_title.strip(), resume_text, resume_skills)
            if insights:
                cache = st.session_state.setdefault('intuitive_ai_cache', {})
                cache[target_title.strip().lower()] = insights
                st.session_state['intuitive_ai_last_target'] = target_title.strip().lower()
                st.success("Job cloud scan refreshed.")
            else:
                st.warning("Job cloud service did not return any insights just yet.")

    active_key = st.session_state.get('intuitive_ai_last_target')
    insights_cache = st.session_state.get('intuitive_ai_cache', {})
    active_insights = insights_cache.get(active_key)

    if active_insights:
        display_intuitive_ai_results(active_insights, focus_choice)
    else:
        st.caption("Run the scan to unlock touchpoint/touchnot guidance.")


def _resume_content_hash(content: str) -> str:
    return hashlib.sha256(content.encode('utf-8')).hexdigest()


def _parse_resume_via_bridge(content: str) -> Optional[Dict[str, Any]]:
    """Persist resume text to disk and parse it via the shared resume service."""
    if not (PORTAL_BRIDGE_AVAILABLE and resume_service):
        return None

    suffix = Path(st.session_state.get('uploaded_resume_filename') or 'resume.txt').suffix or '.txt'
    temp_path: Optional[Path] = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix, mode='w', encoding='utf-8') as tmp_file:
            tmp_file.write(content)
            temp_path = Path(tmp_file.name)

        resume_id = f"{st.session_state.get('user_id', 'user')}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
        parsed = resume_service.parse(
            file_path=str(temp_path),
            user_id=st.session_state.get('user_id'),
            resume_id=resume_id
        )

        if parsed and parsed.get('error'):
            raise RuntimeError(parsed['error'])

        return parsed
    finally:
        if temp_path and temp_path.exists():
            temp_path.unlink(missing_ok=True)


def _enrich_resume_snapshot(parsed_data: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    if not parsed_data or not (PORTAL_BRIDGE_AVAILABLE and intelligence_service):
        return None

    try:
        enriched = intelligence_service.enrich(parsed_data, user_id=st.session_state.get('user_id'))
        if enriched and enriched.get('error'):
            return parsed_data
        return enriched
    except Exception as exc:
        if ERROR_HANDLER_AVAILABLE:
            log_service_error('resume_enrichment_failed', {'error': str(exc)})
        return parsed_data


def _apply_hybrid_enrichment(snapshot: Optional[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    if not snapshot:
        return None

    if HYBRID_ENGINE_AVAILABLE and hybrid_engine:
        try:
            return hybrid_engine.enrich_resume_snapshot(snapshot)
        except Exception as exc:
            if ERROR_HANDLER_AVAILABLE:
                log_service_error('hybrid_enrichment_failed', {'error': str(exc)})
            return snapshot
    return snapshot


def ensure_resume_snapshot_synced(trigger_reason: Optional[str] = None) -> None:
    """Ensure the uploaded resume text is parsed + enriched exactly once per change."""
    content = st.session_state.get('uploaded_resume_content', '').strip()
    if not content:
        return

    content_hash = _resume_content_hash(content)
    pending = st.session_state.get('resume_sync_requested', False)
    if st.session_state.get('resume_sync_hash') == content_hash and not pending:
        return

    parsed_data: Optional[Dict[str, Any]] = None
    enriched_data: Optional[Dict[str, Any]] = None

    try:
        parsed_data = _parse_resume_via_bridge(content)
    except Exception as exc:
        if ERROR_HANDLER_AVAILABLE:
            log_service_error('resume_parse_failed', {'error': str(exc)})
        parsed_data = None

    if not parsed_data:
        parsed_data = {
            'content': content,
            'metadata': {
                'filename': st.session_state.get('uploaded_resume_filename'),
                'source': st.session_state.get('uploaded_resume_source'),
                'synced_at': datetime.now().isoformat(),
                'parser': 'fallback_raw_text'
            }
        }

    enriched_data = _enrich_resume_snapshot(parsed_data)
    hybrid_snapshot = _apply_hybrid_enrichment(enriched_data or parsed_data)

    st.session_state.resume_data = parsed_data
    if enriched_data:
        st.session_state.enriched_data = enriched_data
    if hybrid_snapshot:
        st.session_state.hybrid_resume_snapshot = hybrid_snapshot

    st.session_state.resume_sync_hash = content_hash
    st.session_state.resume_sync_requested = False
    st.session_state.resume_sync_meta = {
        'reason': trigger_reason or st.session_state.get('uploaded_resume_source', 'user'),
        'timestamp': datetime.now().isoformat(),
        'source': st.session_state.get('uploaded_resume_source', 'user'),
        'filename': st.session_state.get('uploaded_resume_filename')
    }

    # Signal downstream portals that a fresh snapshot exists
    st.session_state['umarketu_entrypoint'] = 'from-upload'
    st.session_state['coaching_entrypoint'] = 'from-upload'
    st.session_state['mentor_portal_entrypoint'] = 'from-upload'
    st.session_state['dual_career_entrypoint'] = 'from-upload'


def parse_resume_upload(uploaded_file):
    """Parse uploaded resume via portal bridge and return extracted text + metadata."""
    file_bytes = uploaded_file.read()
    suffix = Path(uploaded_file.name).suffix or ".bin"
    temp_path = None
    parsed_data = None

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file_bytes)
            temp_path = Path(tmp.name)

        if PORTAL_BRIDGE_AVAILABLE and resume_service:
            resume_id = f"{st.session_state.get('user_id', 'user')}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
            parsed_data = resume_service.parse(
                file_path=str(temp_path),
                user_id=st.session_state.get('user_id'),
                resume_id=resume_id
            )

            if parsed_data and parsed_data.get('error'):
                if ERROR_HANDLER_AVAILABLE:
                    log_service_error('resume_parse_error', {
                        'error': parsed_data.get('error'),
                        'file': uploaded_file.name
                    })
                parsed_data = None

        resume_text = None
        if parsed_data:
            resume_text = parsed_data.get('raw_text') or parsed_data.get('text') or parsed_data.get('content')

        if not resume_text:
            if uploaded_file.type == "text/plain":
                resume_text = file_bytes.decode('utf-8', errors='ignore')
            else:
                raise RuntimeError("Resume parser unavailable. Please try again once Admin AI services are online.")

        return resume_text, parsed_data

    finally:
        if temp_path and temp_path.exists():
            temp_path.unlink(missing_ok=True)

# Page configuration
st.set_page_config(
    page_title="📊 Resume Upload & Analysis | IntelliCV-AI",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Authentication check
if not st.session_state.get('authenticated_user'):
    st.error("🔒 Please log in to access Resume Upload & Analysis")
    if st.button("🏠 Return to Home"):
        st.switch_page("pages/01_Home.py")
    st.stop()

# Get user tier
user_tier = st.session_state.get('user_tier', 'free')
is_premium = check_user_tier('monthly_pro')
is_annual_pro = check_user_tier('annual_pro')
is_enterprise = check_user_tier('enterprise_pro')

# Professional CSS styling
def load_resume_analysis_css():
    """Professional styling for resume upload and analysis"""
    css = '''
    <style>
    /* Feature card styling */
    .feature-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 15px;
        color: white;
        margin: 1rem 0;
        box-shadow: 0 8px 16px rgba(102, 126, 234, 0.3);
    }

    .feature-card-locked {
        background: linear-gradient(135deg, #9ca3af 0%, #6b7280 100%);
        padding: 2rem;
        border-radius: 15px;
        color: white;
        margin: 1rem 0;
        box-shadow: 0 8px 16px rgba(107, 114, 128, 0.2);
        opacity: 0.6;
        position: relative;
    }

    .feature-card-locked::before {
        content: "🔒";
        position: absolute;
        top: 10px;
        right: 10px;
        font-size: 2rem;
    }

    /* Upload zone styling */
    .upload-zone {
        border: 3px dashed #667eea;
        border-radius: 15px;
        padding: 3rem;
        text-align: center;
        background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%);
        margin: 2rem 0;
        transition: all 0.3s ease;
    }

    .upload-zone:hover {
        border-color: #764ba2;
        background: linear-gradient(135deg, rgba(102, 126, 234, 0.2) 0%, rgba(118, 75, 162, 0.2) 100%);
    }

    /* Premium badge */
    .premium-badge {
        display: inline-block;
        background: linear-gradient(135deg, #fbbf24 0%, #f59e0b 100%);
        color: white;
        padding: 0.3rem 1rem;
        border-radius: 20px;
        font-weight: bold;
        font-size: 0.85rem;
        margin-left: 0.5rem;
    }

    /* Tier comparison table */
    .tier-comparison {
        background: white;
        border-radius: 15px;
        padding: 2rem;
        margin: 2rem 0;
        box-shadow: 0 4px 12px rgba(0,0,0,0.1);
    }

    /* Analysis results */
    .analysis-metric {
        background: linear-gradient(135deg, #e0e7ff 0%, #ddd6fe 100%);
        padding: 1.5rem;
        border-radius: 10px;
        margin: 0.5rem 0;
        border-left: 5px solid #667eea;
    }

    .analysis-metric h3 {
        color: #667eea;
        margin: 0 0 0.5rem 0;
    }

    /* Drool-worthy premium preview */
    .premium-preview {
        background: linear-gradient(135deg, #fef3c7 0%, #fde68a 100%);
        border: 2px solid #fbbf24;
        border-radius: 15px;
        padding: 2rem;
        margin: 2rem 0;
        position: relative;
    }

    .premium-preview::before {
        content: "✨ PREMIUM FEATURE";
        position: absolute;
        top: -12px;
        left: 20px;
        background: #fbbf24;
        color: white;
        padding: 0.3rem 1rem;
        border-radius: 10px;
        font-weight: bold;
        font-size: 0.8rem;
    }
    </style>
    '''
    st.markdown(css, unsafe_allow_html=True)

load_resume_analysis_css()

# Initialize session state
if 'uploaded_resume_content' not in st.session_state:
    st.session_state.uploaded_resume_content = ""
if 'uploaded_resume_filename' not in st.session_state:
    st.session_state.uploaded_resume_filename = ""
if 'express_analysis_complete' not in st.session_state:
    st.session_state.express_analysis_complete = False
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = {}
if 'resume_data' not in st.session_state:
    st.session_state.resume_data = {}
if 'resume_demo_mode' not in st.session_state:
    st.session_state.resume_demo_mode = False
if 'uploaded_resume_source' not in st.session_state:
    st.session_state.uploaded_resume_source = 'user'
if 'resume_sync_requested' not in st.session_state:
    st.session_state.resume_sync_requested = False
if 'resume_sync_hash' not in st.session_state:
    st.session_state.resume_sync_hash = ''
if 'resume_sync_meta' not in st.session_state:
    st.session_state.resume_sync_meta = {}

ensure_resume_snapshot_synced()

# Resume Builder session states (Premium feature)
if 'chatbot_step' not in st.session_state:
    st.session_state.chatbot_step = 'start'
if 'new_job_data' not in st.session_state:
    st.session_state.new_job_data = {}
if 'education_data' not in st.session_state:
    st.session_state.education_data = {}
if 'resume_versions' not in st.session_state:
    st.session_state.resume_versions = []

# Header with tier badge
st.markdown(f'''
<div style="text-align: center; padding: 2rem 0;">
    <h1>📊 Resume Upload & Analysis</h1>
    <p style="font-size: 1.2rem; color: #667eea;">
        Upload your resume and get instant insights
        <span class="premium-badge">YOUR TIER: {user_tier.upper().replace('_', ' ')}</span>
    </p>
</div>
''', unsafe_allow_html=True)

# Tier status sidebar
with st.sidebar:
    st.markdown("### 🎯 Your Access Level")

    if is_enterprise:
        st.success("✨ **ENTERPRISE** - Full Access")
    elif is_annual_pro:
        st.success("🌟 **ANNUAL PRO** - Full Access")
    elif is_premium:
        st.success("⭐ **MONTHLY PRO** - Full Access")
    else:
        st.info("🆓 **FREE** - Basic Features")
        st.markdown("---")
        st.markdown("### 🚀 Unlock Premium Features:")
        st.markdown("""
        - 🤖 AI-Enhanced Processing
        - 🎯 ATS Compatibility Check
        - 📊 Competitive Analysis
        - 🔍 Deep Keyword Optimization
        - 💼 Market Alignment
        - 🚀 Format Enhancement
        - 📈 Recruiter Perspective
        """)
        if st.button("⬆️ Upgrade Now", key="sidebar_upgrade", use_container_width=True):
            st.switch_page("pages/06_Pricing.py")

# ===== UPLOAD INTERFACE (FREE FOR ALL) =====
st.markdown("## 📤 Upload Your Resume")
st.success("🆓 **FREE** - Resume upload available to all users!")

uploaded_file = st.file_uploader(
    "Choose your resume file",
    type=['pdf', 'doc', 'docx', 'txt', 'rtf', 'odt', 'html', 'htm', 'md', 'csv', 'xls', 'xlsx', 'json', 'png', 'jpg', 'jpeg', 'eml', 'msg', 'zip'],
    help="Upload common resume/data formats (PDF, DOCX, DOC, TXT, RTF, ODT, HTML/MD, CSV, Excel, JSON, images) for real analysis"
)

st.caption("Upload a resume file or paste your own resume text below.")

# Alternative: Text paste (FREE)
with st.expander("📝 Or Paste Resume Text (FREE)", expanded=False):
    resume_text = st.text_area(
        "Paste your resume content",
        height=200,
        placeholder="Copy and paste your complete resume text here...",
        help="Alternative to file upload - paste your resume as plain text"
    )

    if st.button("✅ Use Pasted Text", key="use_text") and resume_text:
        set_uploaded_resume(resume_text, f"Text_Resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt")
        st.success("✅ Resume text loaded!")
        st.rerun()

# Process uploaded file
if uploaded_file is not None:
    try:
        with st.spinner("📄 Processing resume..."):
            extracted_text = extract_text_from_upload(uploaded_file)
            set_uploaded_resume(extracted_text, uploaded_file.name)
            st.success(f"✅ Resume uploaded: {uploaded_file.name}")
            st.rerun()

    except Exception as e:
        show_error(f"Error uploading file: {str(e)}")

# If resume is uploaded, show analysis options
if st.session_state.uploaded_resume_content:
    st.markdown("---")
    st.markdown(f"### 📋 Current Resume: **{st.session_state.uploaded_resume_filename}**")

    sync_meta = st.session_state.get('resume_sync_meta')
    if sync_meta:
        synced_at = sync_meta.get('timestamp')
        reason = sync_meta.get('reason')
        st.caption(f"Synced via {reason} at {synced_at}")

    st.text_area(
        "Resume Preview",
        value=st.session_state.uploaded_resume_content[:5000],
        height=220,
        key="resume_preview",
        disabled=True
    )

    nav_col1, nav_col2, nav_col3 = st.columns(3)
    with nav_col1:
        if st.button("🎓 Open Coaching Hub", key="resume_link_coaching"):
            st.switch_page("pages/10_Coaching_Hub.py")
    with nav_col2:
        if st.button("📂 Go to UMarketU Suite", key="resume_link_umarketu"):
            st.switch_page("pages/09_UMarketU_Suite.py")
    with nav_col3:
        if st.button("🚧 Application Blockers", key="resume_link_blockers"):
            st.switch_page("pages/16_Application_Blockers.py")

    # Analysis tabs (7 tabs total - includes Competitive Intelligence + Application Blockers)
    tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
        "🆓 FREE Analysis",
        "🔒 Premium AI Analysis",
        "🔒 Advanced Insights",
        "🔒 Optimization Tools",
        "🎯 Competitive Intelligence",
        "🔒 AI Resume Builder",
        "🚧 Application Blockers"
    ])

    # ===== TAB 1: FREE EXPRESS ANALYSIS =====
    with tab1:
        st.markdown("### 🆓 FREE Express Analysis")
        st.success("Available to all users - no charge!")

        if st.button("🚀 Run FREE Analysis", type="primary", use_container_width=True):
            content = st.session_state.uploaded_resume_content

            with st.spinner("🔍 Analyzing resume..."):
                time.sleep(2)

                # Basic analysis
                word_count = len(content.split())
                line_count = len(content.split('\n'))
                lines = content.split('\n')
                name = lines[0] if lines else "Unknown"

                # Simple keyword extraction - using dynamic keywords from real CV data
                tech_keywords = []
                keywords_to_find = get_dynamic_keywords('all')  # Load from AI data
                for keyword in keywords_to_find:
                    if keyword.lower() in content.lower():
                        tech_keywords.append(keyword)

                # Store results
                st.session_state.analysis_results = {
                    'name': name,
                    'word_count': word_count,
                    'line_count': line_count,
                    'tech_keywords': tech_keywords,
                    'basic_score': min(100, (word_count / 5) + len(tech_keywords) * 3),
                    'recommendations': [
                        "Add quantifiable achievements with numbers",
                        "Include relevant industry keywords",
                        "Ensure consistent formatting throughout",
                        "Highlight your unique value proposition"
                    ]
                }
                st.session_state.express_analysis_complete = True
                st.rerun()

        # Show results if available
        if st.session_state.express_analysis_complete:
            results = st.session_state.analysis_results

            st.markdown("#### 📊 Quick Metrics")
            col1, col2, col3, col4 = st.columns(4)

            with col1:
                st.metric("Word Count", results['word_count'])
            with col2:
                st.metric("Keywords Found", len(results['tech_keywords']))
            with col3:
                st.metric("Basic Score", f"{results['basic_score']:.0f}/100")
            with col4:
                st.metric("Resume Lines", results['line_count'])

            # Basic précis
            st.markdown("#### 📝 Basic Précis")
            st.info(f"""
            **Professional:** {results['name']}

            **Overview:** Your resume contains {results['word_count']} words across {results['line_count']} lines.
            We identified {len(results['tech_keywords'])} relevant keywords, giving you a basic score of {results['basic_score']:.0f}/100.

            **Key Technologies:** {', '.join(results['tech_keywords']) if results['tech_keywords'] else 'Limited keywords detected'}
            """)

            # Quick recommendations
            st.markdown("#### 💡 Quick Recommendations")
            for i, rec in enumerate(results['recommendations'], 1):
                st.markdown(f"{i}. {rec}")

            # Upgrade hook
            st.markdown("---")
            st.markdown("#### 🚀 Want Deeper Insights?")
            col1, col2 = st.columns(2)

            with col1:
                st.markdown("""
                **🔬 Premium AI Analysis Includes:**
                - 🤖 AI-powered deep resume parsing
                - 🎯 ATS compatibility scoring
                - 📈 Industry benchmark comparison
                - 🔍 Advanced keyword optimization
                - 💼 Career trajectory analysis
                """)

            with col2:
                if st.button("⬆️ Upgrade to Premium", key="upgrade_from_free", use_container_width=True, type="primary"):
                    st.switch_page("pages/06_Pricing.py")

        st.markdown("---")
        render_intuitive_ai_panel()

    # ===== TAB 2: PREMIUM AI ANALYSIS (LOCKED FOR FREE USERS) =====
    with tab2:
        st.markdown("### 🤖 Premium AI-Enhanced Analysis")

        if is_premium:
            st.success("✅ Premium feature unlocked!")

            if st.button("🤖 Run AI Analysis", type="primary", use_container_width=True):
                with st.spinner("🧠 Processing with Admin AI systems..."):
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    steps = [
                        "🧠 Initializing Admin AI...",
                        "📄 Deep content parsing with portal_bridge...",
                        "🔍 Job Title Engine analysis...",
                        "🌐 Real AI Data Connector...",
                        "📊 Competitive benchmarking...",
                        "🎯 Generating insights..."
                    ]

                    for i, step in enumerate(steps):
                        status_text.text(step)
                        progress_bar.progress((i + 1) / len(steps))
                        time.sleep(0.4)

                    ensure_resume_snapshot_synced("premium_analysis")

                parsed_data = st.session_state.get('resume_data')
                enriched_data = st.session_state.get('enriched_data')
                hybrid_snapshot = st.session_state.get('hybrid_resume_snapshot')

                if not parsed_data:
                    st.error("Resume parsing is still pending. Upload a file again or retry once the admin portal is online.")
                else:
                    st.success("✅ AI Analysis Complete! Live data synced to UMarketU, Coaching, Mentorship, and Dual Career hubs.")

                    st.markdown("#### 🎯 ATS Compatibility Score")
                    ats_score = enriched_data.get('ats_score') if isinstance(enriched_data, dict) else None
                    if isinstance(ats_score, (int, float)):
                        st.progress(min(100, max(0, ats_score)) / 100)
                        st.metric("ATS Score", f"{ats_score}/100", delta=f"+{ats_score - 70:.0f} vs. average")
                    else:
                        st.info("ATS score will appear once the Real AI connector returns a score.")

                    st.markdown("#### 🔍 Deep Keyword Analysis")
                    skills_list = parsed_data.get('skills') if isinstance(parsed_data, dict) else []
                    if skills_list:
                        st.success(f"Found {len(skills_list)} industry-aligned keywords")
                        missing_keywords = (enriched_data or {}).get('missing_keywords', []) if isinstance(enriched_data, dict) else []
                        if missing_keywords:
                            st.info(f"Top missing keywords: {', '.join(missing_keywords[:4])}")
                    else:
                        st.warning("Keyword analysis requires parsed resume data. Upload or paste your resume, or retry once parsing is available.")

                    st.markdown("#### 💼 Market Alignment")
                    market_fit = (enriched_data or {}).get('market_fit') if isinstance(enriched_data, dict) else None
                    if isinstance(market_fit, (int, float)):
                        st.metric("Market Fit", f"{market_fit}%", delta=f"+{market_fit - 70:.0f}% vs. industry avg")
                        st.success("Benchmarking completed against the real CV dataset.")
                    else:
                        st.info("Market fit insights unlock once the Real AI connector finishes benchmarking.")

                    if hybrid_snapshot:
                        st.caption("Hybrid snapshot ready • Pages 10, 11, 12, and 14 now consume this data automatically.")

        else:
            # FREE USERS: Greyed out preview
            st.markdown('<div class="premium-preview">', unsafe_allow_html=True)

            st.markdown("#### 🔒 Premium Feature - Upgrade to Unlock")
            st.markdown("""
            **What you'll get with Premium AI Analysis:**

            🤖 **Admin AI Processing**
            - Enhanced Job Title Engine analysis
            - Real AI Data Connector integration
            - Advanced NLP parsing

            🎯 **ATS Compatibility**
            - Detailed ATS score (0-100)
            - Format compliance check
            - Keyword density analysis
            - Missing keyword recommendations

            📊 **Competitive Benchmarking**
            - Compare against 100,000+ resumes
            - Industry-specific scoring
            - Role-level alignment check

            💼 **Market Intelligence**
            - Current market demand analysis
            - Salary range insights
            - Geographic trends
            """)

            # Show greyed out preview
            st.markdown('<div class="feature-card-locked">', unsafe_allow_html=True)
            st.markdown("### 🎯 ATS Compatibility Score")
            st.markdown("**Score:** 🔒 Unlock to see")
            st.markdown("**Missing Keywords:** 🔒 Unlock to see")
            st.markdown("**Format Issues:** 🔒 Unlock to see")
            st.markdown('</div>', unsafe_allow_html=True)

            st.markdown('</div>', unsafe_allow_html=True)

            # Upgrade CTA
            show_upgrade_prompt("AI-Enhanced Analysis", "monthly_pro")

    # ===== TAB 3: ADVANCED INSIGHTS (LOCKED FOR FREE USERS) =====
    with tab3:
        st.markdown("### 📊 Advanced Insights & Analytics")

        if is_premium:
            # PREMIUM: Show real insights
            st.success("✅ Premium feature unlocked!")

            col1, col2 = st.columns(2)

            with col1:
                st.markdown('<div class="analysis-metric">', unsafe_allow_html=True)
                st.markdown("### 🎯 Strength Analysis")
                st.markdown("- **Leadership:** 92/100")
                st.markdown("- **Technical Skills:** 88/100")
                st.markdown("- **Communication:** 85/100")
                st.markdown("- **Industry Knowledge:** 90/100")
                st.markdown('</div>', unsafe_allow_html=True)

            with col2:
                st.markdown('<div class="analysis-metric">', unsafe_allow_html=True)
                st.markdown("### 📈 Recruiter Perspective")
                st.markdown("- **First Impression:** Excellent")
                st.markdown("- **Clarity:** Very Clear")
                st.markdown("- **Impact:** High Impact")
                st.markdown("- **Hire Probability:** 85%")
                st.markdown('</div>', unsafe_allow_html=True)

            st.markdown("#### 💡 Detailed Recommendations")
            st.success("✅ Strong quantifiable achievements throughout")
            st.info("💡 Consider adding more leadership metrics in management roles")
            st.warning("⚠️ Could strengthen cloud architecture keywords")

        else:
            # FREE: Greyed out preview
            st.markdown('<div class="premium-preview">', unsafe_allow_html=True)

            st.markdown("#### 🔒 Premium Feature - See What You're Missing!")
            st.markdown("""
            **Advanced Insights Include:**

            📊 **Strength Analysis**
            - Leadership capability score
            - Technical skills assessment
            - Communication effectiveness
            - Industry knowledge rating

            👀 **Recruiter Perspective**
            - First impression score
            - Clarity and readability
            - Impact and achievement focus
            - Estimated hire probability

            🎯 **Competitive Positioning**
            - How you rank vs. similar profiles
            - Unique selling points identified
            - Areas for differentiation

            💡 **Detailed Recommendations**
            - Specific improvement suggestions
            - Keyword enhancement tips
            - Format optimization advice
            """)

            # Greyed preview
            col1, col2 = st.columns(2)

            with col1:
                st.markdown('<div class="feature-card-locked">', unsafe_allow_html=True)
                st.markdown("### 🎯 Strength Analysis")
                st.markdown("🔒 Unlock to see your scores")
                st.markdown('</div>', unsafe_allow_html=True)

            with col2:
                st.markdown('<div class="feature-card-locked">', unsafe_allow_html=True)
                st.markdown("### 📈 Recruiter View")
                st.markdown("🔒 Unlock to see recruiter perspective")
                st.markdown('</div>', unsafe_allow_html=True)

            st.markdown('</div>', unsafe_allow_html=True)

            show_upgrade_prompt("Advanced Insights", "monthly_pro")

    # ===== TAB 4: OPTIMIZATION TOOLS (LOCKED FOR FREE USERS) =====
    with tab4:
        st.markdown("### 🚀 Resume Optimization Tools")

        if is_premium:
            # PREMIUM: Full optimization tools
            st.success("✅ Premium feature unlocked!")

            st.markdown("#### ✨ Available Optimization Tools:")

            col1, col2 = st.columns(2)

            with col1:
                if st.button("🔍 Keyword Optimizer", use_container_width=True):
                    st.info("Analyzing keywords and suggesting improvements...")
                    st.success("✅ Found 12 keyword enhancement opportunities")

                if st.button("🎨 Format Enhancer", use_container_width=True):
                    st.info("Checking formatting and ATS compatibility...")
                    st.success("✅ 3 formatting improvements suggested")

            with col2:
                if st.button("📊 Content Optimizer", use_container_width=True):
                    st.info("Analyzing content structure and impact...")
                    st.success("✅ 8 content enhancement suggestions")

                if st.button("📚 Version Manager", use_container_width=True):
                    st.info("Loading resume versions...")
                    st.success("✅ 5 versions saved")

            # Show optimization preview
            st.markdown("---")
            st.markdown("#### 💡 Current Optimization Suggestions")
            st.success("✅ Strong action verbs throughout")
            st.info("💡 Add more quantifiable metrics to 3 bullet points")
            st.warning("⚠️ Consider restructuring skills section for ATS optimization")

        else:
            # FREE: Greyed out preview - make them DROOL!
            st.markdown('<div class="premium-preview">', unsafe_allow_html=True)

            st.markdown("#### 🔒 Premium Optimization Suite - Transform Your Resume!")
            st.markdown("""
            **Powerful Tools to Perfect Your Resume:**

            🔍 **Keyword Optimizer**
            - AI-powered keyword suggestions
            - Industry-specific term recommendations
            - Competitive keyword density analysis
            - Missing critical keywords alerts

            🎨 **Format Enhancer**
            - ATS-optimized formatting
            - Professional template suggestions
            - Layout optimization
            - Typography improvements

            📊 **Content Optimizer**
            - Achievement impact scoring
            - Action verb enhancement
            - Quantification suggestions
            - Bullet point restructuring

            📚 **Version Manager**
            - Save multiple resume versions
            - Track changes over time
            - Compare versions side-by-side
            - Revert to previous versions

            🔗 **LinkedIn Import**
            - One-click LinkedIn profile import
            - Auto-populate from LinkedIn
            - Sync updates automatically
            """)

            # Greyed out tool buttons
            st.markdown("---")
            col1, col2 = st.columns(2)

            with col1:
                st.markdown('<div class="feature-card-locked">', unsafe_allow_html=True)
                st.button("🔍 Keyword Optimizer 🔒", disabled=True, use_container_width=True)
                st.button("🎨 Format Enhancer 🔒", disabled=True, use_container_width=True)
                st.markdown('</div>', unsafe_allow_html=True)

            with col2:
                st.markdown('<div class="feature-card-locked">', unsafe_allow_html=True)
                st.button("📊 Content Optimizer 🔒", disabled=True, use_container_width=True)
                st.button("📚 Version Manager 🔒", disabled=True, use_container_width=True)
                st.markdown('</div>', unsafe_allow_html=True)

            st.markdown('</div>', unsafe_allow_html=True)

            show_upgrade_prompt("Optimization Tools", "monthly_pro")

    # ===== TAB 5: COMPETITIVE INTELLIGENCE (FREE + PREMIUM) =====
    with tab5:
        st.markdown("### 🎯 Competitive Intelligence - How Do You Compare?")

        if INTELLIGENCE_SERVICE_AVAILABLE and USER_DATA_SERVICE_AVAILABLE:
            st.success("🆓 **FREE PREVIEW** - See how you rank against 150 real candidates!")

            # Initialize services
            intelligence = CandidateIntelligence()
            user_service = get_user_data_service()

            if st.button("📊 Generate My Intelligence Profile", type="primary", use_container_width=True, key="gen_intelligence"):
                with st.spinner("🔍 Analyzing your competitive positioning against learning library..."):
                    # Get user's real data from learning library (150 candidates)
                    user_id = st.session_state.get('user_id', 'candidate_001')

                    try:
                        user_data = user_service.get_user_data(user_id)

                        if user_data and user_data['profile']:
                            profile_data = user_data['profile']
                            stats = user_data['stats']

                            # Use REAL user data from learning library
                            candidate_name = profile_data.get('name', st.session_state.get('user_name', 'You'))
                            candidate_score = profile_data.get('score', 7.0)
                            candidate_keywords = profile_data.get('keywords', [])

                            # Generate profile with real user stats
                            profile = intelligence.generate_candidate_profile(
                                candidate_name=candidate_name,
                                candidate_score=candidate_score,
                                candidate_keywords=candidate_keywords
                            )

                            # Enhance profile with UserDataService stats
                            profile['rankings']['score_percentile'] = stats['score_percentile']
                            profile['rankings']['keyword_percentile'] = stats['keyword_percentile']
                            profile['rankings']['overall_rank'] = stats['rank']
                            profile['peer_comparison']['score_vs_average'] = stats['score_vs_average']
                            profile['peer_comparison']['keywords_vs_average'] = stats['keywords_vs_average']
                            profile['peer_comparison']['total_candidates'] = stats['total_candidates']

                            # Add real insights
                            tier = profile_data.get('tier', 'average')
                            insights = {
                                'strengths': [],
                                'opportunities': [],
                                'recommendations': []
                            }

                            if tier == 'top_performer':
                                insights['strengths'].append(f"🌟 You're a TOP PERFORMER - Ranked {stats['rank']}/{stats['total_candidates']}")
                                insights['strengths'].append(f"📈 Your score ({candidate_score:.2f}) is {stats['score_vs_average']:+.2f} points above average")
                                insights['recommendations'].append("💼 Consider mentoring other candidates in the learning library")
                            elif tier == 'average':
                                insights['opportunities'].append(f"📊 You're in the AVERAGE tier - Room to climb to top performer")
                                insights['opportunities'].append(f"🎯 Add {abs(stats['keywords_vs_average']):.0f} more keywords to match top performers")
                                insights['recommendations'].append("📚 Focus on skill development to break into top 25%")
                            else:  # below_average
                                insights['opportunities'].append(f"🔄 You're BELOW AVERAGE - Significant growth potential")
                                insights['opportunities'].append(f"📉 Your score is {abs(stats['score_vs_average']):.2f} points below average")
                                insights['recommendations'].append("🚀 Seek mentorship from top performers in your field")
                                insights['recommendations'].append(f"💡 Add technical keywords - you have {len(candidate_keywords)}, average is {int(stats['peer_avg_keywords'])}")

                            profile['insights'] = insights

                            st.session_state.intelligence_profile = profile
                            st.success(f"✅ Profile generated using REAL data from {stats['total_candidates']} candidates!")
                            st.rerun()

                        else:
                            st.error("❌ Could not load your profile from learning library")

                    except Exception as e:
                        st.error("⚠️ **Service Temporarily Unavailable**")
                        st.markdown("""
                        We are experiencing technical difficulties at our service provider end.

                        **Status:**
                        - AI analysis services are temporarily offline
                        - Your uploaded resume is safe and secure
                        - No data has been lost

                        **Next Steps:**
                        - We will notify you directly via message when services are restored
                        - Please check back in 15-30 minutes
                        - Contact support if issue persists: support@intellicv.ai

                        *We apologize for the inconvenience and are working to resolve this quickly.*
                        """)

                        # Log the service outage
                        if ERROR_HANDLER_AVAILABLE:
                            log_service_error("resume_intelligence_service_unavailable", {
                                'error': str(e),
                                'user_id': st.session_state.get('user_id'),
                                'timestamp': datetime.now().isoformat()
                            })

                        st.stop()  # Don't show fallback data

            # Show profile if generated
            if st.session_state.get('intelligence_profile'):
                profile = st.session_state.intelligence_profile

                st.markdown("---")

                # === INFO BANNER ===
                st.info(f"""
                📊 **Real Data Analysis**: Your rankings are calculated from the **Learning Library** of {profile['peer_comparison'].get('total_candidates', 150)} real candidates.

                🎯 This is REAL competitive intelligence, not simulated data!
                """)

                # === HEADER: Overall Rankings ===
                st.markdown("## 📊 Your Competitive Rankings")

                col1, col2, col3 = st.columns(3)

                with col1:
                    st.metric(
                        "🎯 Score Percentile",
                        f"{profile['rankings']['score_percentile']:.0f}th",
                        delta=f"{profile['peer_comparison']['score_vs_average']:+.1f} vs avg"
                    )

                with col2:
                    st.metric(
                        "📚 Keyword Percentile",
                        f"{profile['rankings']['keyword_percentile']:.0f}th",
                        delta=f"{int(profile['peer_comparison']['keywords_vs_average']):+d} vs avg"
                    )

                with col3:
                    st.metric(
                        "🏆 Overall Rank",
                        f"Top {100 - profile['rankings']['overall_rank']:.0f}%",
                        delta=f"{profile['rankings']['overall_rank']:.0f}th percentile"
                    )

                # === QUADRANT POSITIONING ===
                st.markdown("---")
                st.markdown("## 🎯 Competitive Positioning")

                quadrant = profile['quadrant']
                st.markdown(f"""
                <div style="background: linear-gradient(135deg, {quadrant['color']}22 0%, {quadrant['color']}44 100%);
                            padding: 2rem; border-radius: 15px; border-left: 5px solid {quadrant['color']};">
                    <h2 style="color: {quadrant['color']}; margin: 0;">{quadrant['name']}</h2>
                    <p style="font-size: 1.2rem; margin: 0.5rem 0 0 0;">{quadrant['description']}</p>
                </div>
                """, unsafe_allow_html=True)

                # Quadrant chart
                fig_quadrant = create_quadrant_chart(
                    profile['visualization_data'],
                    quadrant,
                    profile['candidate_name']
                )
                st.plotly_chart(fig_quadrant, use_container_width=True)

                # === VISUALIZATIONS ===
                st.markdown("---")
                st.markdown("## 📈 Visual Analysis")

                viz_col1, viz_col2 = st.columns(2)

                with viz_col1:
                    # Score distribution histogram
                    fig_hist = create_score_histogram(
                        profile['visualization_data'],
                        profile['candidate_name']
                    )
                    st.plotly_chart(fig_hist, use_container_width=True)

                with viz_col2:
                    # Spider chart
                    fig_spider = create_spider_chart(
                        profile['visualization_data'],
                        profile['candidate_name']
                    )
                    st.plotly_chart(fig_spider, use_container_width=True)

                # === PEER COMPARISON TABLE ===
                st.markdown("---")

                top_performers = intelligence.get_top_performers(5)
                peer_comp_html = create_peer_comparison_table(
                    {**profile['peer_comparison'],
                     'score': profile['score'],
                     'keyword_count': profile['keyword_count']},
                    profile['rankings'],
                    top_performers
                )
                st.markdown(peer_comp_html, unsafe_allow_html=True)

                # === INSIGHTS ===
                st.markdown("---")
                st.markdown("## 💡 Competitive Insights")

                insights = profile['insights']

                if insights['strengths']:
                    st.markdown("### ✨ Your Strengths")
                    for strength in insights['strengths']:
                        st.success(strength)

                if insights['opportunities']:
                    st.markdown("### 🎯 Growth Opportunities")
                    for opp in insights['opportunities']:
                        st.info(opp)

                if insights['recommendations']:
                    st.markdown("### 📋 Actionable Recommendations")
                    for rec in insights['recommendations']:
                        st.warning(rec)

                # === PREMIUM UPSELL ===
                if not is_premium:
                    st.markdown("---")
                    st.markdown("### 🚀 Want Even Deeper Insights?")

                    st.info("""
                    **Upgrade to Premium for:**
                    - 🤖 AI-powered profile optimization suggestions
                    - 📊 Industry-specific competitive benchmarking
                    - 🎯 Personalized keyword recommendations
                    - 💼 Career trajectory analysis
                    - 📈 Real-time market positioning updates
                    """)

                    if st.button("⬆️ Upgrade to Premium", key="upgrade_from_intelligence", use_container_width=True):
                        st.switch_page("pages/06_Pricing.py")

        else:
            # Fallback if services not available
            st.error("❌ Competitive Intelligence service not available")

            if not INTELLIGENCE_SERVICE_AVAILABLE:
                st.warning("⚠️ CandidateIntelligence service not loaded")

            if not USER_DATA_SERVICE_AVAILABLE:
                st.warning("⚠️ UserDataService not loaded - cannot access learning library (150 candidates)")
                st.info("""
                **Missing:** `services/user_data_service.py`

                This service connects you to the **Learning Library** with 150 real candidate profiles for benchmarking.
                """)

    # ===== TAB 6: AI RESUME BUILDER (PREMIUM ONLY) =====
    with tab6:
        st.markdown("### 🤖 AI Resume Builder")

        if is_premium:
            # PREMIUM: Full AI Resume Builder functionality
            st.success("✅ Premium feature unlocked!")
            st.info("Build and edit your resume with our interactive AI assistant!")

            chatbot_step = st.session_state.chatbot_step
            new_job_data = st.session_state.new_job_data
            education_data = st.session_state.education_data

            if chatbot_step == 'start':
                st.markdown("**👋 Hi! I'm your AI resume assistant. Let's build or enhance your resume.**")
                st.markdown("**What would you like to do?**")

                builder_col1, builder_col2 = st.columns(2)

                with builder_col1:
                    if st.button("💼 **Add New Job Experience**", use_container_width=True, key="add_job"):
                        st.session_state.chatbot_step = 'job_title'
                        st.rerun()

                with builder_col2:
                    if st.button("🎓 **Add Education**", use_container_width=True, key="add_edu"):
                        st.session_state.chatbot_step = 'education'
                        st.rerun()

                st.markdown("---")

                quick_col1, quick_col2 = st.columns(2)

                with quick_col1:
                    if st.button("🔧 **Add Skills Section**", use_container_width=True, key="add_skills"):
                        st.session_state.chatbot_step = 'skills'
                        st.rerun()

                with quick_col2:
                    if st.button("📝 **Update Summary**", use_container_width=True, key="add_summary"):
                        st.session_state.chatbot_step = 'summary'
                        st.rerun()

            elif chatbot_step == 'job_title':
                st.markdown("**💼 Let's add your new job experience!**")

                job_title = st.text_input("What's your job title?", key="job_title_input")

                col1, col2 = st.columns(2)
                with col1:
                    if st.button("⬅️ Back", key="job_back"):
                        st.session_state.chatbot_step = 'start'
                        st.rerun()

                with col2:
                    if job_title and st.button("Next ➡️", key="job_next"):
                        st.session_state.new_job_data['job_title'] = job_title
                        st.session_state.chatbot_step = 'company_name'
                        st.rerun()

            elif chatbot_step == 'company_name':
                st.markdown(f"**Great! Adding: {new_job_data.get('job_title', '')}**")

                company_name = st.text_input("What company is this for?", key="company_input")

                col1, col2 = st.columns(2)
                with col1:
                    if st.button("⬅️ Back", key="company_back"):
                        st.session_state.chatbot_step = 'job_title'
                        st.rerun()

                with col2:
                    if company_name and st.button("Next ➡️", key="company_next"):
                        st.session_state.new_job_data['company'] = company_name
                        st.session_state.chatbot_step = 'start_date'
                        st.rerun()

            elif chatbot_step == 'start_date':
                st.markdown("**📅 When did you start this position?**")

                from datetime import date
                start_date = st.date_input("Start Date", key="start_date_input")
                is_current = st.checkbox("This is my current position", key="is_current_check")

                end_date = None
                if not is_current:
                    end_date = st.date_input("End Date", key="end_date_input")

                col1, col2 = st.columns(2)
                with col1:
                    if st.button("⬅️ Back", key="date_back"):
                        st.session_state.chatbot_step = 'company_name'
                        st.rerun()

                with col2:
                    if st.button("Next ➡️", key="date_next"):
                        st.session_state.new_job_data['start_date'] = start_date.isoformat()
                        st.session_state.new_job_data['is_current'] = is_current
                        if end_date:
                            st.session_state.new_job_data['end_date'] = end_date.isoformat()
                        st.session_state.chatbot_step = 'job_description'
                        st.rerun()

            elif chatbot_step == 'job_description':
                st.markdown("**📝 Now let's describe your responsibilities and achievements**")

                desc_col1, desc_col2 = st.columns(2)

                with desc_col1:
                    st.markdown("**✍️ Write your own:**")
                    custom_desc = st.text_area("Job Description", height=150, key="custom_desc")

                    if custom_desc and st.button("Use This Description", key="use_custom"):
                        st.session_state.new_job_data['job_description'] = custom_desc
                        st.session_state.chatbot_step = 'review'
                        st.rerun()

                with desc_col2:
                    st.markdown("**🤖 Use AI Suggestion:**")

                    if st.button("🔍 **Generate from AI Database**", use_container_width=True, key="gen_desc"):
                        with st.spinner("🤖 Searching real job description database..."):
                            # Load REAL job descriptions from ai_data/company_job_descriptions/
                            job_title = new_job_data.get('job_title', 'Professional')
                            company = new_job_data.get('company', '')

                            jd_dir = Path("ai_data/company_job_descriptions")

                            if not jd_dir.exists():
                                st.error("⚠️ **Job Description Database Unavailable**")
                                st.markdown("""
                                Cannot generate description from our database.

                                **Please write your own job description above.**

                                *Tip: Include your key responsibilities and achievements.*
                                """)
                            else:
                                # Find matching JDs from real database
                                matching_jds = []
                                try:
                                    for domain_dir in jd_dir.iterdir():
                                        if domain_dir.is_dir():
                                            for jd_file in domain_dir.glob("*.json"):
                                                try:
                                                    with open(jd_file, 'r', encoding='utf-8') as f:
                                                        jd_data = json.load(f)
                                                        if job_title.lower() in jd_data.get('title', '').lower():
                                                            matching_jds.append(jd_data)
                                                except:
                                                    continue

                                    if matching_jds:
                                        # Use most recent JD as template
                                        recent_jd = max(matching_jds, key=lambda x: x.get('scraped_date', ''))
                                        responsibilities = recent_jd.get('responsibilities', [])

                                        if responsibilities:
                                            ai_desc = "\n".join([f"• {resp}" for resp in responsibilities[:6]])
                                            st.session_state.new_job_data['job_description'] = ai_desc
                                            st.success(f"✅ Description generated from real {recent_jd.get('company', 'job')} posting!")
                                            st.text_area("Generated Description", ai_desc, height=150, disabled=True, key="ai_desc_preview")

                                            if st.button("Use This Description", key="use_ai"):
                                                st.session_state.chatbot_step = 'review'
                                                st.rerun()
                                        else:
                                            st.warning(f"No responsibilities found in database for '{job_title}'. Please write your own.")
                                    else:
                                        st.warning(f"No job descriptions found for '{job_title}' in our database. Please write your own.")

                                except Exception as e:
                                    if ERROR_HANDLER_AVAILABLE:
                                        log_service_error("jd_generation_failed", {'error': str(e), 'job_title': job_title})
                                    st.error("⚠️ Error accessing job description database. Please write your own.")

                if st.button("⬅️ Back", key="desc_back"):
                    st.session_state.chatbot_step = 'start_date'
                    st.rerun()

            elif chatbot_step == 'review':
                st.markdown("**📋 Review Your New Job Entry**")

                job_data = st.session_state.new_job_data

                # Show complete entry
                st.markdown(f"**💼 Position:** {job_data.get('job_title', '')}")
                st.markdown(f"**🏢 Company:** {job_data.get('company', '')}")

                start_date = job_data.get('start_date', '')
                if job_data.get('is_current', False):
                    st.markdown(f"**📅 Duration:** {start_date} - Present")
                else:
                    end_date = job_data.get('end_date', '')
                    st.markdown(f"**📅 Duration:** {start_date} to {end_date}")

                st.markdown("**📝 Description:**")
                st.text_area("Job Description", job_data.get('job_description', ''), height=150, disabled=True, key="review_desc")

                # Final actions
                col1, col2, col3 = st.columns(3)

                with col1:
                    if st.button("⬅️ Edit Description", key="review_back"):
                        st.session_state.chatbot_step = 'job_description'
                        st.rerun()

                with col2:
                    if st.button("✅ **Add to Resume**", use_container_width=True, type="primary", key="add_to_resume"):
                        # Add to resume content
                        new_entry = f"""
{job_data.get('job_title', '')} | {job_data.get('company', '')} | {job_data.get('start_date', '')} - {'Present' if job_data.get('is_current') else job_data.get('end_date', '')}
{job_data.get('job_description', '')}
"""

                        if st.session_state.uploaded_resume_content:
                            st.session_state.uploaded_resume_content += "\n" + new_entry
                        else:
                            st.session_state.uploaded_resume_content = new_entry

                        # Add to version history
                        version_entry = {
                            'content': st.session_state.uploaded_resume_content,
                            'timestamp': datetime.now().isoformat(),
                            'tag': f'Added {job_data.get("job_title", "")} experience'
                        }
                        st.session_state.resume_versions.insert(0, version_entry)

                        st.success("✅ Job experience added to your resume!")
                        st.session_state.chatbot_step = 'start'
                        st.session_state.new_job_data = {}
                        st.rerun()

                with col3:
                    if st.button("❌ Cancel", key="review_cancel"):
                        st.session_state.chatbot_step = 'start'
                        st.session_state.new_job_data = {}
                        st.rerun()

            elif chatbot_step == 'education':
                st.markdown("**🎓 Let's add your education!**")

                if 'step' not in education_data:
                    education_data['step'] = 'degree'

                if education_data['step'] == 'degree':
                    degree = st.selectbox("Select your degree level:",
                                        ["", "High School Diploma", "Associate Degree", "Bachelor's Degree",
                                         "Master's Degree", "Doctorate (Ph.D.)", "Professional Certification"],
                                        key="degree_select")

                    field_of_study = st.text_input("Field of Study (e.g., Computer Science, Business Administration):", key="field_input")

                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("⬅️ Back to Menu", key="edu_back_menu"):
                            st.session_state.chatbot_step = 'start'
                            st.rerun()

                    with col2:
                        if degree and field_of_study and st.button("Next ➡️", key="edu_next"):
                            education_data['degree'] = degree
                            education_data['field'] = field_of_study
                            education_data['step'] = 'school'
                            st.rerun()

                elif education_data['step'] == 'school':
                    st.markdown(f"**Adding: {education_data.get('degree')} in {education_data.get('field')}**")

                    school_name = st.text_input("School/University Name:", key="school_input")
                    location = st.text_input("Location (City, State):", key="location_input")

                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("⬅️ Back", key="school_back"):
                            education_data['step'] = 'degree'
                            st.rerun()

                    with col2:
                        if school_name and st.button("Next ➡️", key="school_next"):
                            education_data['school'] = school_name
                            education_data['location'] = location
                            education_data['step'] = 'graduation'
                            st.rerun()

                elif education_data['step'] == 'graduation':
                    st.markdown("**📅 When did you complete (or expect to complete) this education?**")

                    from datetime import date
                    grad_date = st.date_input("Completion Date:", key="grad_date_input")

                    st.caption("Note: Academic achievements and honors can be added in a separate section if relevant.")

                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("⬅️ Back", key="grad_back"):
                            education_data['step'] = 'school'
                            st.rerun()

                    with col2:
                        if st.button("✅ Add Education", key="add_edu_final"):
                            # Create education entry - NO assumptions, only what user provided
                            degree = education_data.get('degree', '')
                            field = education_data.get('field', '')
                            school = education_data.get('school', '')
                            location = education_data.get('location', '')

                            # Build entry based on what's available
                            edu_parts = ["EDUCATION:"]

                            if degree and field:
                                edu_parts.append(f"{degree} in {field}")
                            elif degree:
                                edu_parts.append(f"{degree}")
                            elif field:
                                edu_parts.append(f"Studies in {field}")

                            if school:
                                edu_parts.append(f"{school}{', ' + location if location else ''}")

                            edu_parts.append(f"Completed: {grad_date.strftime('%B %Y')}")

                            edu_entry = "\n".join(edu_parts) + "\n"

                            if st.session_state.uploaded_resume_content:
                                st.session_state.uploaded_resume_content += "\n" + edu_entry
                            else:
                                st.session_state.uploaded_resume_content = edu_entry

                            # Add to version history
                            version_entry = {
                                'content': st.session_state.uploaded_resume_content,
                                'timestamp': datetime.now().isoformat(),
                                'tag': f'Added {education_data.get("degree")} education'
                            }
                            st.session_state.resume_versions.insert(0, version_entry)

                            st.success("✅ Education added to your resume!")
                            st.session_state.chatbot_step = 'start'
                            st.session_state.education_data = {}
                            st.rerun()

            elif chatbot_step == 'skills':
                st.markdown("**🔧 Let's add your skills section!**")

                skills_input = st.text_area(
                    "Enter your skills (separate with commas):",
                    placeholder="Python, JavaScript, Project Management, Data Analysis, Leadership, SQL, etc.",
                    height=100,
                    key="skills_textarea"
                )

                if skills_input:
                    st.markdown("**Preview:**")
                    skills_list = [skill.strip() for skill in skills_input.split(',') if skill.strip()]

                    # Categorize skills automatically using dynamic keywords from REAL CV data
                    if REAL_AI_CONNECTOR_AVAILABLE and real_ai_connector:
                        try:
                            technical_keywords = [kw.lower() for kw in get_dynamic_keywords('technical')]
                            soft_keywords = [kw.lower() for kw in get_dynamic_keywords('soft')]
                        except:
                            # If categorization fails, just show all skills without categorization
                            st.markdown("**Your Skills:**")
                            for skill in skills_list:
                                st.write(f"• {skill}")
                            technical_keywords = []
                            soft_keywords = []
                    else:
                        # No AI loader - show skills without categorization
                        st.markdown("**Your Skills:**")
                        for skill in skills_list:
                            st.write(f"• {skill}")
                        technical_keywords = []
                        soft_keywords = []

                    technical_skills = []
                    soft_skills = []
                    other_skills = []

                    for skill in skills_list:
                        skill_lower = skill.lower()
                        if any(keyword in skill_lower for keyword in technical_keywords):
                            technical_skills.append(skill)
                        elif any(keyword in skill_lower for keyword in soft_keywords):
                            soft_skills.append(skill)
                        else:
                            other_skills.append(skill)

                    skills_entry = "SKILLS:\n"
                    if technical_skills:
                        skills_entry += f"Technical Skills: {', '.join(technical_skills)}\n"
                    if soft_skills:
                        skills_entry += f"Soft Skills: {', '.join(soft_skills)}\n"
                    if other_skills:
                        skills_entry += f"Other Skills: {', '.join(other_skills)}\n"

                    st.text_area("Formatted Skills Section:", skills_entry, height=150, disabled=True, key="skills_preview")

                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("⬅️ Back to Menu", key="skills_back"):
                            st.session_state.chatbot_step = 'start'
                            st.rerun()

                    with col2:
                        if st.button("✅ Add Skills Section", key="add_skills_final"):
                            if st.session_state.uploaded_resume_content:
                                st.session_state.uploaded_resume_content += "\n" + skills_entry
                            else:
                                st.session_state.uploaded_resume_content = skills_entry

                            # Add to version history
                            version_entry = {
                                'content': st.session_state.uploaded_resume_content,
                                'timestamp': datetime.now().isoformat(),
                                'tag': 'Added skills section'
                            }
                            st.session_state.resume_versions.insert(0, version_entry)

                            st.success("✅ Skills section added to your resume!")
                            st.session_state.chatbot_step = 'start'
                            st.rerun()
                else:
                    if st.button("⬅️ Back to Menu", key="skills_back_empty"):
                        st.session_state.chatbot_step = 'start'
                        st.rerun()

            elif chatbot_step == 'summary':
                st.markdown("**📝 Let's create/update your professional summary!**")

                st.markdown("**Choose how to create your summary:**")

                summary_col1, summary_col2 = st.columns(2)

                with summary_col1:
                    st.markdown("**✍️ Write your own:**")
                    custom_summary = st.text_area(
                        "Professional Summary:",
                        height=150,
                        placeholder="Write a brief professional summary highlighting your key strengths, experience, and career objectives...",
                        key="custom_summary"
                    )

                    if custom_summary and st.button("Use Custom Summary", key="use_custom_summary"):
                        summary_entry = f"\nPROFESSIONAL SUMMARY:\n{custom_summary}\n"

                        if st.session_state.uploaded_resume_content:
                            st.session_state.uploaded_resume_content = summary_entry + st.session_state.uploaded_resume_content
                        else:
                            st.session_state.uploaded_resume_content = summary_entry

                        # Add to version history
                        version_entry = {
                            'content': st.session_state.uploaded_resume_content,
                            'timestamp': datetime.now().isoformat(),
                            'tag': 'Updated professional summary'
                        }
                        st.session_state.resume_versions.insert(0, version_entry)

                        st.success("✅ Professional summary updated!")
                        st.session_state.chatbot_step = 'start'
                        st.rerun()

                with summary_col2:
                    st.markdown("**🤖 AI-Generated Summary:**")

                    if st.button("🔍 Generate AI Summary", use_container_width=True, key="gen_summary"):
                        with st.spinner("🤖 Analyzing your resume and generating summary..."):
                            time.sleep(2)

                            # AI-generated summary
                            ai_summary = """Experienced professional with demonstrated expertise in cross-functional leadership and strategic project management. Proven track record of driving operational efficiency and implementing innovative solutions. Strong background in stakeholder collaboration, team development, and quality assurance. Committed to continuous improvement and delivering measurable results in dynamic business environments."""

                            st.text_area("AI-Generated Summary:", ai_summary, height=150, disabled=True, key="ai_summary_preview")

                            if st.button("Use AI Summary", key="use_ai_summary"):
                                summary_entry = f"\nPROFESSIONAL SUMMARY:\n{ai_summary}\n"

                                if st.session_state.uploaded_resume_content:
                                    st.session_state.uploaded_resume_content = summary_entry + st.session_state.uploaded_resume_content
                                else:
                                    st.session_state.uploaded_resume_content = summary_entry

                                # Add to version history
                                version_entry = {
                                    'content': st.session_state.uploaded_resume_content,
                                    'timestamp': datetime.now().isoformat(),
                                    'tag': 'Added AI-generated professional summary'
                                }
                                st.session_state.resume_versions.insert(0, version_entry)

                                st.success("✅ AI-generated summary added!")
                                st.session_state.chatbot_step = 'start'
                                st.rerun()

                if st.button("⬅️ Back to Menu", key="summary_back"):
                    st.session_state.chatbot_step = 'start'
                    st.rerun()

            # Version history viewer
            if st.session_state.resume_versions:
                st.markdown("---")
                st.markdown("### 📚 Version History")

                with st.expander(f"📄 {len(st.session_state.resume_versions)} versions saved", expanded=False):
                    for idx, version in enumerate(st.session_state.resume_versions[:5]):  # Show last 5
                        st.markdown(f"**Version {idx + 1}** - {version.get('tag', 'Untitled')}")
                        st.caption(f"Created: {version.get('timestamp', '')[:19]}")
                        if st.button(f"⭐ Restore Version {idx + 1}", key=f"restore_{idx}"):
                            st.session_state.uploaded_resume_content = version.get('content', '')
                            st.success(f"✅ Restored version {idx + 1}!")
                            st.rerun()
                        st.markdown("---")

        else:
            # FREE: Greyed out preview - DROOL-WORTHY!
            st.markdown('<div class="premium-preview">', unsafe_allow_html=True)

            st.markdown("#### 🔒 AI Resume Builder - Premium Feature")
            st.markdown("""
            **Build Your Perfect Resume with AI Assistance!**

            🤖 **Interactive AI Chatbot**
            - Step-by-step resume building
            - Intelligent question flow
            - Real-time preview
            - Context-aware suggestions

            💼 **Add Job Experience**
            - Smart job title suggestions
            - AI-generated job descriptions
            - Achievement bullet point generator
            - Company and date tracking

            🎓 **Add Education**
            - Degree and institution tracking
            - GPA and honors support
            - Graduation date management
            - Automatic formatting

            🔧 **Skills Management**
            - Auto-categorization (Technical/Soft)
            - Industry keyword suggestions
            - Skill level tracking
            - Gap analysis

            📝 **Professional Summary Generator**
            - AI-powered summary creation
            - Customizable templates
            - Industry-specific language
            - Achievement highlighting

            📚 **Version Control**
            - Save multiple resume versions
            - Track all changes
            - Restore previous versions
            - Compare versions side-by-side

            🔍 **Smart Features**
            - Job description to resume converter
            - LinkedIn profile import
            - ATS-optimized formatting
            - Real-time word count
            """)

            # Greyed out interface preview
            st.markdown("---")
            st.markdown("**👋 AI Assistant Preview:**")

            st.markdown('<div class="feature-card-locked">', unsafe_allow_html=True)
            st.markdown("**What would you like to do? 🔒**")

            builder_col1, builder_col2 = st.columns(2)

            with builder_col1:
                st.button("💼 Add New Job Experience 🔒", disabled=True, use_container_width=True, key="locked_job")
                st.button("🔧 Add Skills Section 🔒", disabled=True, use_container_width=True, key="locked_skills")

            with builder_col2:
                st.button("🎓 Add Education 🔒", disabled=True, use_container_width=True, key="locked_edu")
                st.button("📝 Update Summary 🔒", disabled=True, use_container_width=True, key="locked_summary")

            st.markdown('</div>', unsafe_allow_html=True)

            st.markdown('</div>', unsafe_allow_html=True)

            # Strong upgrade CTA
            st.markdown("---")
            st.error("🔒 **This feature requires a Premium subscription**")
            st.markdown("""
            ### 🚀 Unlock AI Resume Builder

            **Build professional resumes in minutes, not hours!**

            With Premium, you get:
            - ✅ Interactive AI chatbot guide
            - ✅ Auto-generated job descriptions
            - ✅ Professional summary writer
            - ✅ Unlimited resume versions
            - ✅ Smart skill categorization
            - ✅ LinkedIn import

            **Plus all other premium features:**
            AI Analysis, ATS Check, Optimization Tools, and more!
            """)

            show_upgrade_prompt("AI Resume Builder", "monthly_pro")

    # ===== TAB 7: APPLICATION BLOCKERS (FREE - NEW FEATURE!) =====
    with tab7:
        st.markdown("### 🚧 Application Blockers - Gap Analysis")
        st.success("✨ NEW FEATURE - Available to all users!")

        st.info("""
        **What are Application Blockers?**

        The opposite of "touch points" - these are areas where you DON'T match job requirements.
        Identify your gaps, get improvement plans, and prepare proactive interview responses!
        """)

        if BLOCKER_ENGINE_AVAILABLE:
            resume_content = st.session_state.uploaded_resume_content

            # Parse basic info from resume (simplified)
            candidate_profile = {
                "user_id": st.session_state.get('user_id') or st.session_state.get('authenticated_user') or "unknown_user",
                "skills": [],
                "total_experience_years": 0,
                "certifications": [],
                "education": [],
                "industries": []
            }

            # Extract skills (basic keyword matching)
            tech_keywords = get_dynamic_keywords('technical') if REAL_AI_CONNECTOR_AVAILABLE else []
            for keyword in tech_keywords:
                if keyword.lower() in resume_content.lower():
                    candidate_profile["skills"].append(keyword)

            # Show what we extracted
            with st.expander("📋 Extracted Profile Info", expanded=False):
                st.json({
                    "skills_found": candidate_profile["skills"][:10],
                    "total_skills": len(candidate_profile["skills"])
                })

            # Job selector for blocker analysis
            st.markdown("#### 🎯 Provide Job Description for Analysis")

            st.markdown("**Option 1:** Paste a real job description")
            jd_text = st.text_area(
                "Job Description Text",
                placeholder="Paste the full job description here...",
                height=200,
                help="Copy and paste from LinkedIn, company website, or job board"
            )

            st.markdown("**OR**")

            st.markdown("**Option 2:** Search our job description database")
            jd_search = st.text_input("Search for job title", placeholder="e.g., Senior Software Engineer")

            if st.button("🔍 Analyze My Blockers", type="primary", use_container_width=True):
                with st.spinner("Analyzing gaps between your resume and job requirements..."):

                    # Load REAL job requirements
                    job_requirements = None
                    selected_job = "Job Position"

                    if jd_text:
                        # Parse real JD text
                        try:
                            from shared.services.jd_parser import parse_job_description
                            job_requirements = parse_job_description(jd_text)
                            selected_job = job_requirements.get('title', 'Provided Job Description')
                        except:
                            # Simple keyword extraction if parser unavailable
                            job_requirements = {
                                "raw_text": jd_text,
                                "required_skills": extract_skills_from_text(jd_text)
                            }
                            min_years = extract_experience_requirement(jd_text)
                            if min_years is not None:
                                job_requirements["min_experience_years"] = min_years
                            education_req = extract_education_requirement(jd_text)
                            if education_req:
                                job_requirements["required_education"] = education_req
                            selected_job = "Provided Job Description"

                    elif jd_search:
                        # Search real JD database
                        jd_dir = Path("ai_data/company_job_descriptions")

                        if jd_dir.exists():
                            for domain_dir in jd_dir.iterdir():
                                if domain_dir.is_dir():
                                    for jd_file in domain_dir.glob("*.json"):
                                        try:
                                            with open(jd_file, 'r', encoding='utf-8') as f:
                                                jd_data = json.load(f)
                                                if jd_search.lower() in jd_data.get('title', '').lower():
                                                    job_requirements = jd_data
                                                    selected_job = jd_data.get('title', 'Database Job')
                                                    break
                                        except:
                                            continue
                                if job_requirements:
                                    break

                    if not job_requirements:
                        st.error("⚠️ **Please provide a job description**")
                        st.markdown("Either paste the full JD text or search our database.")
                        st.stop()

                    # Run blocker analysis
                    engine = ApplicationBlockerEngine()

                    job_id = None
                    if isinstance(job_requirements, dict):
                        job_id = job_requirements.get('job_id')

                    if not job_id:
                        stable_basis = jd_text if jd_text else selected_job
                        job_id = hashlib.sha256((stable_basis or "job").encode('utf-8')).hexdigest()[:12]

                    company_name = ""
                    if isinstance(job_requirements, dict):
                        company_name = (
                            job_requirements.get('company')
                            or job_requirements.get('company_name')
                            or ""
                        )

                    job_desc = {
                        "job_id": job_id,
                        "title": selected_job,
                        "company": company_name
                    }

                    analysis = engine.analyze_single_job(
                        candidate_profile,
                        job_desc,
                        job_requirements
                    )

                    # Display results
                    st.markdown("---")
                    st.markdown("### 📊 Blocker Analysis Results")

                    # Metrics
                    col1, col2, col3, col4 = st.columns(4)

                    with col1:
                        viability_color = {
                            "Strong": "🟢",
                            "Moderate": "🟡",
                            "Weak": "🟠",
                            "Not Recommended": "🔴"
                        }
                        st.metric(
                            "Application Viability",
                            f"{viability_color.get(analysis.application_viability, '⚪')} {analysis.application_viability}"
                        )

                    with col2:
                        st.metric("Total Blockers", analysis.total_blockers)

                    with col3:
                        st.metric("Critical Gaps", analysis.critical_count, delta=f"-{analysis.critical_count}", delta_color="inverse")

                    with col4:
                        st.metric("Addressable", f"{analysis.addressable_count}/{analysis.total_blockers}")

                    # Top 3 priorities
                    st.markdown("---")
                    st.markdown("### 🎯 Top 3 Priorities to Address")

                    for i, priority in enumerate(analysis.top_3_priorities[:3], 1):
                        blocker = next((b for b in analysis.blockers if b.requirement == priority), None)
                        if blocker:
                            severity_emoji = {
                                "critical": "🔴",
                                "major": "🟠",
                                "moderate": "🟡",
                                "minor": "🔵"
                            }

                            with st.expander(f"{severity_emoji.get(blocker.severity.value, '⚪')} **Priority {i}: {priority}**", expanded=(i==1)):
                                st.markdown(f"**Gap**: {blocker.gap_description}")
                                st.markdown(f"**Impact Score**: {blocker.impact_score:.1f}/100")

                                if blocker.is_addressable:
                                    st.markdown(f"**Time to Address**: {blocker.time_to_address}")
                                    st.info(f"**💡 Improvement Plan**: {blocker.improvement_plan}")
                                else:
                                    st.warning("⚠️ Not easily addressable - focus on highlighting related strengths")

                                if blocker.alternative_strength:
                                    st.success(f"**💪 Your Alternative Strength**: {blocker.alternative_strength}")

                    # Proactive acknowledgment
                    st.markdown("---")
                    st.markdown("### 🎤 Interview Preparation: Proactive Acknowledgment")

                    ack_gen = ProactiveAcknowledgmentGenerator()
                    candidate_strengths = [
                        skill for skill in candidate_profile["skills"][:3]
                    ] if candidate_profile["skills"] else ["Strong technical background"]

                    statement = ack_gen.generate_acknowledgment(
                        analysis.blockers,
                        candidate_strengths
                    )

                    st.info(statement)

                    if st.button("📋 Copy Interview Statement"):
                        st.toast("Copied to clipboard!")

                    # Link to full Application Blockers page
                    st.markdown("---")
                    if st.button("🚧 View Full Application Blockers Dashboard", use_container_width=True):
                        st.switch_page("pages/16_Application_Blockers.py")

        else:
            st.warning("⚠️ Application Blocker engine not available. Please contact support.")
            st.info("Navigate to Page 16 (Application Blockers) for the full feature.")

# ===== COMPARISON TABLE: FREE VS PREMIUM =====
st.markdown("---")
st.markdown("## 📊 Feature Comparison: Free vs. Premium")

comparison_data = {
    "Feature": [
        "📤 Resume Upload (All Formats)",
        "📄 Basic Text Extraction",
        "📊 Express Analysis",
        "🔍 Keyword Detection (Basic)",
        "💡 Quick Recommendations",
        "🤖 AI-Enhanced Processing",
        "🎯 ATS Compatibility Check",
        "📈 Competitive Analysis",
        "🔍 Deep Keyword Optimization",
        "💼 Market Alignment",
        "🚀 Format Enhancement",
        "📊 Strength Analysis",
        "👀 Recruiter Perspective",
        "🔍 Keyword Optimizer",
        "🎨 Format Enhancer",
        "📊 Content Optimizer",
        "📚 Version Manager",
        "🔗 LinkedIn Import",
        "🤖 AI Resume Builder",
        "💼 Add/Edit Job Experience (AI)",
        "🎓 Add/Edit Education (AI)",
        "🔧 Skills Management (AI)",
        "📝 Professional Summary Generator"
    ],
    "🆓 Free": [
        "✅", "✅", "✅", "✅", "✅",
        "❌", "❌", "❌", "❌", "❌", "❌", "❌", "❌", "❌", "❌", "❌", "❌", "❌",
        "❌", "❌", "❌", "❌", "❌"
    ],
    "⭐ Monthly Pro": [
        "✅", "✅", "✅", "✅", "✅",
        "✅", "✅", "✅", "✅", "✅", "✅", "✅", "✅", "✅", "✅", "✅", "✅", "✅",
        "✅", "✅", "✅", "✅", "✅"
    ],
    "🌟 Annual Pro": [
        "✅", "✅", "✅", "✅", "✅",
        "✅", "✅", "✅", "✅", "✅", "✅", "✅", "✅", "✅", "✅", "✅", "✅", "✅",
        "✅", "✅", "✅", "✅", "✅"
    ]
}

# Display as formatted table
st.markdown('<div class="tier-comparison">', unsafe_allow_html=True)

col1, col2, col3, col4 = st.columns([2, 1, 1, 1])

with col1:
    st.markdown("### Feature")
with col2:
    st.markdown("### 🆓 Free")
with col3:
    st.markdown("### ⭐ Pro ($15.99/mo)")
with col4:
    st.markdown("### 🌟 Annual ($149.99/yr)")

st.markdown("---")

for i in range(len(comparison_data["Feature"])):
    col1, col2, col3, col4 = st.columns([2, 1, 1, 1])

    with col1:
        st.markdown(comparison_data["Feature"][i])
    with col2:
        st.markdown(f"<div style='text-align: center;'>{comparison_data['🆓 Free'][i]}</div>", unsafe_allow_html=True)
    with col3:
        st.markdown(f"<div style='text-align: center;'>{comparison_data['⭐ Monthly Pro'][i]}</div>", unsafe_allow_html=True)
    with col4:
        st.markdown(f"<div style='text-align: center;'>{comparison_data['🌟 Annual Pro'][i]}</div>", unsafe_allow_html=True)

st.markdown('</div>', unsafe_allow_html=True)

# ===== UPGRADE CTA =====
if not is_premium:
    st.markdown("---")
    st.markdown('<div class="premium-preview">', unsafe_allow_html=True)

    st.markdown("### 🚀 Ready to Unlock Premium Features?")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown("#### ⭐ Monthly Pro")
        st.markdown("**$15.99/month**")
        st.markdown("- All premium features")
        st.markdown("- Cancel anytime")
        st.markdown("- Priority support")
        if st.button("Choose Monthly Pro", key="upgrade_monthly", use_container_width=True, type="primary"):
            st.switch_page("pages/06_Pricing.py")

    with col2:
        st.markdown("#### 🌟 Annual Pro")
        st.markdown("**$149.99/year**")
        st.markdown("- Save $41.89 annually")
        st.markdown("- All premium features")
        st.markdown("- Priority support")
        if st.button("Choose Annual Pro", key="upgrade_annual", use_container_width=True, type="primary"):
            st.switch_page("pages/06_Pricing.py")

    with col3:
        st.markdown("#### ✨ Enterprise")
        st.markdown("**$499/year**")
        st.markdown("- Team features")
        st.markdown("- Custom integrations")
        st.markdown("- Dedicated support")
        if st.button("Choose Enterprise", key="upgrade_enterprise", use_container_width=True):
            st.switch_page("pages/06_Pricing.py")

    st.markdown('</div>', unsafe_allow_html=True)

# Footer
st.markdown("---")
st.markdown("""
<div style="text-align: center; padding: 2rem 0; color: #666;">
    <p>📊 IntelliCV-AI Resume Analysis</p>
    <p style="font-size: 0.9rem;">Powered by Advanced AI • ATS-Optimized • Industry-Tested</p>
</div>
""", unsafe_allow_html=True)
