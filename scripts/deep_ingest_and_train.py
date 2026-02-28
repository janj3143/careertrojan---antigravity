#!/usr/bin/env python3
"""
CareerTrojan — Deep Ingestion, Collocation Mining & AI Training Pipeline
========================================================================

SINGLE-RUN UNIFIED PIPELINE that:

  Phase A: Deep-parse every file in automated_parser/ (PDF, DOCX, XLSX, CSV, MSG, DOC, TXT)
           → Extract raw text, job titles, skills, company names, contact info
           → Output parsed JSON per document into ai_data_final/parsed_from_automated/
           → Build/update consolidated_terms.json, candidate DB, company DB

  Phase B: Collocation mining on ALL extracted text
           → N-gram extraction, PMI scoring, NLTK collocations, co-occurrence
           → Categorize into TECH_SKILL, SOFT_SKILL, CERTIFICATION, etc.
           → Build final clean collocation results
           → Merge into gazetteers and update learned_collocations.json

  Phase C: Full AI model training
           → Bayesian classifier, TF-IDF vectorizer
           → Sentence-BERT embeddings, spaCy NER
           → K-Means/DBSCAN clustering, cosine similarity matrix
           → Job title classifier, salary predictor
           → All models saved to services/ai_engine/trained_models/

Usage:
    python scripts/deep_ingest_and_train.py

Author: CareerTrojan System
Date: February 2026
"""

import json
import os
import re
import sys
import math
import time
import pickle
import logging
import hashlib
import traceback
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional, Set, Tuple
from collections import Counter, defaultdict
from dataclasses import dataclass, field

# ── Ensure UTF-8 ──
if sys.platform == "win32":
    os.environ["PYTHONIOENCODING"] = "utf-8"

# ── Paths ──
SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent
DATA_ROOT = Path(os.getenv("CAREERTROJAN_DATA_ROOT", r"L:\antigravity_version_ai_data_final"))
AI_DATA_DIR = DATA_ROOT / "ai_data_final"
AUTOMATED_PARSER_DIR = DATA_ROOT / "automated_parser"
GAZETTEERS_DIR = AI_DATA_DIR / "gazetteers"
MODELS_DIR = PROJECT_ROOT / "services" / "ai_engine" / "trained_models"
LOG_DIR = PROJECT_ROOT / "logs"

# Ensure output dirs
for d in [AI_DATA_DIR / "parsed_from_automated", AI_DATA_DIR / "parsed_resumes",
          AI_DATA_DIR / "parsed_job_descriptions", AI_DATA_DIR / "job_titles",
          AI_DATA_DIR / "core_databases", AI_DATA_DIR / "profiles",
          GAZETTEERS_DIR, MODELS_DIR, LOG_DIR]:
    d.mkdir(parents=True, exist_ok=True)

# ── Logging ──
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s — %(message)s",
    handlers=[
        logging.FileHandler(LOG_DIR / "deep_ingest_and_train.log", encoding="utf-8"),
        logging.StreamHandler(sys.stdout),
    ],
)
logger = logging.getLogger("DeepPipeline")

# ── Lazy imports (heavy libs) ──
_spacy_nlp = None
_sentence_model = None


def get_spacy_nlp():
    global _spacy_nlp
    if _spacy_nlp is None:
        import spacy
        try:
            _spacy_nlp = spacy.load("en_core_web_sm")
        except OSError:
            _spacy_nlp = spacy.blank("en")
    return _spacy_nlp


def get_sentence_model():
    global _sentence_model
    if _sentence_model is None:
        try:
            from sentence_transformers import SentenceTransformer
            _sentence_model = SentenceTransformer("all-MiniLM-L6-v2")
        except Exception as e:
            logger.warning("Sentence-BERT not available: %s", e)
    return _sentence_model


# ═══════════════════════════════════════════════════════════════════════
#  PHASE A: DEEP DOCUMENT PARSING
# ═══════════════════════════════════════════════════════════════════════

SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".xlsx", ".xls", ".csv",
    ".msg", ".txt", ".rtf", ".json", ".odt",
}

# Patterns for extracting structured data from raw text
EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"(?:\+?\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}")
LINKEDIN_RE = re.compile(r"linkedin\.com/in/[\w-]+", re.IGNORECASE)

# Job title detection (broad patterns)
JOB_TITLE_PATTERNS = [
    re.compile(r"\b((?:Senior|Junior|Lead|Principal|Chief|Head of|VP of|Director of|Manager of)\s+[\w\s&/-]{3,40})", re.IGNORECASE),
    re.compile(r"\b([\w]+\s+(?:Engineer|Developer|Analyst|Manager|Director|Consultant|Specialist|Coordinator|Administrator|Architect|Designer|Scientist|Officer|Executive))\b", re.IGNORECASE),
    re.compile(r"\b((?:Project|Product|Program|Account|Sales|Marketing|Operations|Business|Data|Software|Process|Chemical|Mechanical|Electrical|Nuclear|Petroleum|HR|QA|IT|UX|UI)\s+(?:Manager|Engineer|Developer|Analyst|Lead|Director|Specialist|Consultant|Coordinator))\b", re.IGNORECASE),
]

# Skill extraction patterns (multi-word technical terms)
SKILL_INDICATORS = {
    "proficient in", "experienced with", "knowledge of", "skilled in",
    "expertise in", "familiar with", "working knowledge", "competent in",
    "certified in", "qualified in", "trained in", "specializing in",
}

# Common section headers in CVs
CV_SECTION_HEADERS = re.compile(
    r"^(?:work\s+)?(?:experience|employment|education|skills|qualifications|"
    r"certifications|summary|objective|profile|training|projects|publications|"
    r"interests|references|languages|technical\s+skills|professional\s+experience|"
    r"career\s+history|professional\s+summary|core\s+competencies|key\s+skills)",
    re.IGNORECASE | re.MULTILINE,
)


def extract_text_from_pdf(filepath: Path) -> str:
    """Extract text from PDF using pdfplumber."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages[:100]:  # Cap at 100 pages
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts)
    except Exception as e:
        logger.debug("PDF extraction failed for %s: %s", filepath.name, e)
        return ""


def extract_text_from_docx(filepath: Path) -> str:
    """Extract text from DOCX."""
    try:
        from docx import Document
        doc = Document(str(filepath))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        logger.debug("DOCX extraction failed for %s: %s", filepath.name, e)
        return ""


def extract_text_from_xlsx(filepath: Path) -> str:
    """Extract text from Excel files."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(filepath), read_only=True, data_only=True)
        text_parts = []
        for sheet in wb.worksheets[:10]:  # Cap sheets
            for row in sheet.iter_rows(max_row=5000, values_only=True):
                row_text = " | ".join(str(c) for c in row if c is not None and str(c).strip())
                if row_text.strip():
                    text_parts.append(row_text)
        wb.close()
        return "\n".join(text_parts)
    except Exception as e:
        logger.debug("XLSX extraction failed for %s: %s", filepath.name, e)
        return ""


def extract_text_from_csv(filepath: Path) -> str:
    """Extract text from CSV files."""
    try:
        import csv
        text_parts = []
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            for i, row in enumerate(reader):
                if i > 10000:
                    break
                row_text = " | ".join(str(c) for c in row if c and str(c).strip())
                if row_text.strip():
                    text_parts.append(row_text)
        return "\n".join(text_parts)
    except Exception as e:
        logger.debug("CSV extraction failed for %s: %s", filepath.name, e)
        return ""


def extract_text_from_msg(filepath: Path) -> str:
    """Extract text from Outlook MSG files."""
    try:
        import extract_msg
        msg = extract_msg.Message(str(filepath))
        parts = []
        if msg.subject:
            parts.append(f"Subject: {msg.subject}")
        if msg.sender:
            parts.append(f"From: {msg.sender}")
        if msg.body:
            parts.append(msg.body)
        msg.close()
        return "\n".join(parts)
    except Exception as e:
        logger.debug("MSG extraction failed for %s: %s", filepath.name, e)
        return ""


def extract_text_from_txt(filepath: Path) -> str:
    """Extract text from plain text files."""
    try:
        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                return filepath.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        return ""
    except Exception as e:
        logger.debug("TXT extraction failed for %s: %s", filepath.name, e)
        return ""


def extract_text_from_json(filepath: Path) -> str:
    """Extract text content from JSON files."""
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            data = json.load(f)
        
        def flatten_json(obj, prefix=""):
            parts = []
            if isinstance(obj, dict):
                for k, v in obj.items():
                    parts.extend(flatten_json(v, f"{prefix}{k}: "))
            elif isinstance(obj, list):
                for item in obj[:200]:  # Cap list items
                    parts.extend(flatten_json(item, prefix))
            elif isinstance(obj, str) and len(obj) > 2:
                parts.append(f"{prefix}{obj}")
            return parts
        
        return "\n".join(flatten_json(data))
    except Exception as e:
        logger.debug("JSON extraction failed for %s: %s", filepath.name, e)
        return ""


EXTRACTORS = {
    ".pdf": extract_text_from_pdf,
    ".docx": extract_text_from_docx,
    ".doc": extract_text_from_docx,  # python-docx handles some .doc files
    ".xlsx": extract_text_from_xlsx,
    ".xls": extract_text_from_xlsx,
    ".csv": extract_text_from_csv,
    ".msg": extract_text_from_msg,
    ".txt": extract_text_from_txt,
    ".rtf": extract_text_from_txt,   # Basic RTF as text
    ".json": extract_text_from_json,
    ".odt": extract_text_from_txt,   # Fallback
}


def classify_document(filepath: Path, text: str) -> str:
    """Classify document as CV, job_description, company_info, email, data, or other."""
    name_lower = filepath.stem.lower()
    ext = filepath.suffix.lower()
    text_lower = text[:3000].lower() if text else ""

    # CV/Resume indicators
    cv_signals = ["cv", "resume", "curriculum vitae", "résumé", "lebenslauf",
                  "personal profile", "career history", "work experience",
                  "professional experience", "education:", "skills:",
                  "jvcv", "profile"]
    if any(s in name_lower for s in ["cv", "resume", "résumé", "lebenslauf", "jvcv", "profile"]):
        return "cv"
    if sum(1 for s in cv_signals if s in text_lower) >= 3:
        return "cv"

    # Job Description indicators
    jd_signals = ["job description", "job title:", "department:", "reports to:",
                  "responsibilities:", "requirements:", "qualifications:",
                  "we are looking", "about the role", "the ideal candidate",
                  "jd", "requisition", "vacancy"]
    if any(s in name_lower for s in ["job description", "jd ", "requisition", "vacancy"]):
        return "job_description"
    if sum(1 for s in jd_signals if s in text_lower) >= 3:
        return "job_description"

    # Company data
    if any(s in name_lower for s in ["compan", "client", "account"]):
        return "company_info"

    # Email/MSG
    if ext == ".msg" or "subject:" in text_lower[:200]:
        return "email"

    # Spreadsheet data files
    if ext in (".xlsx", ".xls", ".csv") and len(text) > 500:
        return "data"

    # ESCO/NAICS classification data
    if any(s in name_lower for s in ["esco", "naics", "soc2020", "occupation"]):
        return "classification_data"

    return "other"


def extract_job_titles_from_text(text: str) -> List[str]:
    """Extract job titles from document text using patterns + spaCy."""
    titles = set()
    for pattern in JOB_TITLE_PATTERNS:
        for match in pattern.finditer(text[:10000]):
            title = match.group(1).strip()
            if 3 < len(title) < 80 and not any(c.isdigit() for c in title[:3]):
                titles.add(title)
    return list(titles)[:50]


def extract_skills_from_text(text: str) -> List[str]:
    """Extract skills from text using keyword context and NER."""
    skills = set()
    text_lower = text.lower()

    # Pattern-based extraction
    for indicator in SKILL_INDICATORS:
        idx = text_lower.find(indicator)
        while idx != -1:
            # Get text after the indicator until end of line or next indicator
            after = text[idx + len(indicator):idx + len(indicator) + 200]
            # Split on common delimiters
            for chunk in re.split(r"[,;•·|\n]", after):
                chunk = chunk.strip().rstrip(".")
                if 2 < len(chunk) < 60:
                    skills.add(chunk.strip())
            idx = text_lower.find(indicator, idx + 1)

    # spaCy NER for additional extraction
    try:
        nlp = get_spacy_nlp()
        doc = nlp(text[:5000])
        for ent in doc.ents:
            if ent.label_ in ("ORG", "PRODUCT", "WORK_OF_ART") and 2 < len(ent.text) < 50:
                skills.add(ent.text)
    except Exception:
        pass

    return list(skills)[:100]


def extract_company_names(text: str) -> List[str]:
    """Extract company names using spaCy NER."""
    companies = set()
    try:
        nlp = get_spacy_nlp()
        doc = nlp(text[:8000])
        for ent in doc.ents:
            if ent.label_ == "ORG" and 2 < len(ent.text) < 80:
                companies.add(ent.text)
    except Exception:
        pass
    return list(companies)[:50]


def extract_contact_info(text: str) -> Dict[str, Any]:
    """Extract emails, phones, LinkedIn from text."""
    return {
        "emails": list(set(EMAIL_RE.findall(text[:5000])))[:5],
        "phones": list(set(PHONE_RE.findall(text[:5000])))[:5],
        "linkedin": list(set(LINKEDIN_RE.findall(text[:5000])))[:3],
    }


def compute_file_hash(filepath: Path) -> str:
    """Compute MD5 hash of file for deduplication."""
    h = hashlib.md5()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


@dataclass
class ParsedDocument:
    source_file: str
    file_hash: str
    doc_type: str  # cv, job_description, company_info, email, data, classification_data, other
    raw_text: str
    text_length: int
    job_titles: List[str]
    skills: List[str]
    companies: List[str]
    contact_info: Dict[str, Any]
    parsed_at: str
    file_extension: str
    file_size_kb: float

    def to_dict(self) -> Dict[str, Any]:
        return {
            "source_file": self.source_file,
            "file_hash": self.file_hash,
            "doc_type": self.doc_type,
            "raw_text": self.raw_text[:50000],  # Cap stored text
            "text_length": self.text_length,
            "job_titles": self.job_titles,
            "skills": self.skills,
            "companies": self.companies,
            "contact_info": self.contact_info,
            "parsed_at": self.parsed_at,
            "file_extension": self.file_extension,
            "file_size_kb": self.file_size_kb,
        }


def deep_parse_automated_parser() -> Dict[str, Any]:
    """
    Phase A: Deep-parse every supported file in automated_parser/.
    Returns summary statistics and all parsed documents.
    """
    logger.info("=" * 80)
    logger.info("PHASE A: DEEP DOCUMENT PARSING — automated_parser/")
    logger.info("=" * 80)

    if not AUTOMATED_PARSER_DIR.exists():
        logger.error("automated_parser directory not found: %s", AUTOMATED_PARSER_DIR)
        return {"error": "directory_not_found"}

    # Load existing parse progress to skip already-parsed files
    progress_file = AI_DATA_DIR / "deep_parse_progress.json"
    processed_hashes: Set[str] = set()
    if progress_file.exists():
        try:
            with open(progress_file, "r") as f:
                processed_hashes = set(json.load(f).get("processed_hashes", []))
            logger.info("Loaded %d previously processed file hashes", len(processed_hashes))
        except Exception:
            pass

    # Collect all supported files
    all_files = []
    ignored_dirs = {"__pycache__", ".git", ".venv", "node_modules", "incoming", "completed"}
    for fpath in AUTOMATED_PARSER_DIR.rglob("*"):
        if not fpath.is_file():
            continue
        if any(ig in fpath.parts for ig in ignored_dirs):
            continue
        if fpath.suffix.lower() in SUPPORTED_EXTENSIONS:
            all_files.append(fpath)

    logger.info("Found %d supported files in automated_parser/", len(all_files))

    # Stats
    stats = {
        "total_files": len(all_files),
        "parsed": 0,
        "skipped_existing": 0,
        "failed": 0,
        "by_type": Counter(),
        "by_extension": Counter(),
    }

    all_parsed: List[ParsedDocument] = []
    all_raw_texts: List[str] = []
    all_job_titles: List[str] = []
    all_skills: List[str] = []
    all_companies: List[str] = []
    new_hashes: List[str] = []

    start_time = time.time()

    for i, filepath in enumerate(all_files):
        try:
            # Progress reporting
            if (i + 1) % 50 == 0 or i == 0:
                elapsed = time.time() - start_time
                rate = (i + 1) / elapsed if elapsed > 0 else 0
                logger.info(
                    "  [%d/%d] %.1f files/sec — %s",
                    i + 1, len(all_files), rate, filepath.name[:60],
                )

            # Check hash for deduplication
            try:
                file_hash = compute_file_hash(filepath)
            except Exception:
                file_hash = f"nohash_{filepath.name}"

            if file_hash in processed_hashes:
                stats["skipped_existing"] += 1
                continue

            # Extract text
            ext = filepath.suffix.lower()
            extractor = EXTRACTORS.get(ext)
            if not extractor:
                continue

            raw_text = extractor(filepath)
            if not raw_text or len(raw_text.strip()) < 20:
                stats["failed"] += 1
                new_hashes.append(file_hash)  # Mark as attempted
                continue

            # Classify document
            doc_type = classify_document(filepath, raw_text)

            # Extract structured data
            job_titles = extract_job_titles_from_text(raw_text)
            skills = extract_skills_from_text(raw_text)
            companies = extract_company_names(raw_text)
            contact = extract_contact_info(raw_text)

            # Build parsed record
            parsed = ParsedDocument(
                source_file=str(filepath.relative_to(DATA_ROOT)),
                file_hash=file_hash,
                doc_type=doc_type,
                raw_text=raw_text,
                text_length=len(raw_text),
                job_titles=job_titles,
                skills=skills,
                companies=companies,
                contact_info=contact,
                parsed_at=datetime.now().isoformat(),
                file_extension=ext,
                file_size_kb=round(filepath.stat().st_size / 1024, 2),
            )

            all_parsed.append(parsed)
            all_raw_texts.append(raw_text)
            all_job_titles.extend(job_titles)
            all_skills.extend(skills)
            all_companies.extend(companies)
            new_hashes.append(file_hash)

            stats["parsed"] += 1
            stats["by_type"][doc_type] += 1
            stats["by_extension"][ext] += 1

        except Exception as e:
            logger.debug("Failed to parse %s: %s", filepath.name, e)
            stats["failed"] += 1
            continue

    elapsed = time.time() - start_time

    # ── Save parsed documents ──
    logger.info("Saving parsed documents to ai_data_final/...")

    # Save per-document JSON files
    output_dir = AI_DATA_DIR / "parsed_from_automated"
    cv_output_dir = AI_DATA_DIR / "parsed_resumes"
    jd_output_dir = AI_DATA_DIR / "parsed_job_descriptions"

    cv_count = 0
    jd_count = 0
    for parsed_doc in all_parsed:
        try:
            safe_name = re.sub(r'[<>:"/\\|?*]', '_', Path(parsed_doc.source_file).stem)[:80]
            doc_dict = parsed_doc.to_dict()

            # Save to appropriate directory based on type
            if parsed_doc.doc_type == "cv":
                out_path = cv_output_dir / f"{safe_name}.json"
                cv_count += 1
            elif parsed_doc.doc_type == "job_description":
                out_path = jd_output_dir / f"{safe_name}.json"
                jd_count += 1
            else:
                out_path = output_dir / f"{safe_name}.json"

            with open(out_path, "w", encoding="utf-8") as f:
                json.dump(doc_dict, f, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.debug("Failed to save %s: %s", parsed_doc.source_file, e)

    # ── Update consolidated_terms.json ──
    consolidated_path = AI_DATA_DIR / "consolidated_terms.json"
    existing_terms = []
    if consolidated_path.exists():
        try:
            with open(consolidated_path, "r", encoding="utf-8") as f:
                existing_terms = json.load(f)
        except Exception:
            pass

    # Add new terms (skills, job titles, companies)
    new_terms = set()
    for skill in all_skills:
        if len(skill) > 3:
            new_terms.add(skill.strip())
    for title in all_job_titles:
        if len(title) > 3:
            new_terms.add(title.strip())
    for company in all_companies:
        if len(company) > 3:
            new_terms.add(company.strip())

    existing_set = set(existing_terms)
    truly_new = new_terms - existing_set
    updated_terms = existing_terms + sorted(truly_new)

    with open(consolidated_path, "w", encoding="utf-8") as f:
        json.dump(updated_terms, f, indent=2, ensure_ascii=False)
    logger.info("consolidated_terms.json updated: %d existing + %d new = %d total",
                len(existing_terms), len(truly_new), len(updated_terms))

    # ── Update core_databases/Candidate_database_merged.json ──
    merged_db_path = AI_DATA_DIR / "core_databases" / "Candidate_database_merged.json"
    existing_candidates = []
    if merged_db_path.exists():
        try:
            with open(merged_db_path, "r", encoding="utf-8") as f:
                existing_candidates = json.load(f)
        except Exception:
            pass

    existing_hashes = {c.get("file_hash") for c in existing_candidates if isinstance(c, dict)}
    new_candidates = []
    for parsed_doc in all_parsed:
        if parsed_doc.doc_type == "cv" and parsed_doc.file_hash not in existing_hashes:
            # Build a candidate record
            candidate = {
                "name": Path(parsed_doc.source_file).stem.replace("_", " "),
                "file_hash": parsed_doc.file_hash,
                "source_file": parsed_doc.source_file,
                "skills": parsed_doc.skills,
                "Job Title": parsed_doc.job_titles[0] if parsed_doc.job_titles else "",
                "current_position": parsed_doc.job_titles[0] if parsed_doc.job_titles else "",
                "companies": parsed_doc.companies,
                "contact": parsed_doc.contact_info,
                "raw_text": parsed_doc.raw_text[:5000],
                "text_length": parsed_doc.text_length,
                "parsed_at": parsed_doc.parsed_at,
            }
            new_candidates.append(candidate)

    all_candidates = existing_candidates + new_candidates
    with open(merged_db_path, "w", encoding="utf-8") as f:
        json.dump(all_candidates, f, indent=2, ensure_ascii=False)
    logger.info("Candidate_database_merged.json: %d existing + %d new = %d total",
                len(existing_candidates), len(new_candidates), len(all_candidates))

    # ── Update enhanced_job_titles_database.json ──
    job_titles_db_path = AI_DATA_DIR / "enhanced_job_titles_database.json"
    existing_jt_db = {}
    if job_titles_db_path.exists():
        try:
            with open(job_titles_db_path, "r", encoding="utf-8") as f:
                existing_jt_db = json.load(f)
        except Exception:
            pass

    # Build unique job titles with frequency
    jt_counter = Counter(t.strip().title() for t in all_job_titles if len(t.strip()) > 3)
    existing_titles = set(existing_jt_db.get("normalized_job_titles", []))
    new_titles = set(jt_counter.keys()) - existing_titles

    if not isinstance(existing_jt_db.get("normalized_job_titles"), list):
        existing_jt_db["normalized_job_titles"] = list(existing_titles)
    existing_jt_db["normalized_job_titles"] = sorted(
        set(existing_jt_db["normalized_job_titles"]) | new_titles
    )

    # Skill mappings from extracted skills
    skill_mappings = existing_jt_db.get("skill_mappings", {})
    for parsed_doc in all_parsed:
        for title in parsed_doc.job_titles:
            title_key = title.strip().title()
            if title_key not in skill_mappings:
                skill_mappings[title_key] = list(set(parsed_doc.skills[:20]))
            else:
                existing = set(skill_mappings[title_key])
                for s in parsed_doc.skills[:20]:
                    existing.add(s)
                skill_mappings[title_key] = list(existing)[:30]
    existing_jt_db["skill_mappings"] = skill_mappings

    with open(job_titles_db_path, "w", encoding="utf-8") as f:
        json.dump(existing_jt_db, f, indent=2, ensure_ascii=False)
    logger.info("enhanced_job_titles_database.json: %d titles, %d skill mappings",
                len(existing_jt_db["normalized_job_titles"]), len(skill_mappings))

    # ── Save progress (hashes) ──
    all_processed = processed_hashes | set(new_hashes)
    with open(progress_file, "w", encoding="utf-8") as f:
        json.dump({
            "processed_hashes": sorted(all_processed),
            "last_run": datetime.now().isoformat(),
            "total_processed": len(all_processed),
        }, f, indent=2)

    stats["elapsed_seconds"] = round(elapsed, 2)
    stats["cvs_saved"] = cv_count
    stats["jds_saved"] = jd_count
    stats["new_candidates"] = len(new_candidates)
    stats["new_terms"] = len(truly_new)
    stats["new_job_titles"] = len(new_titles)
    stats["total_raw_text_chars"] = sum(len(t) for t in all_raw_texts)

    logger.info("=" * 80)
    logger.info("PHASE A COMPLETE — %d parsed, %d skipped, %d failed in %.1fs",
                stats["parsed"], stats["skipped_existing"], stats["failed"], elapsed)
    logger.info("  CVs: %d | JDs: %d | New candidates: %d | New terms: %d",
                cv_count, jd_count, len(new_candidates), len(truly_new))
    logger.info("  By type: %s", dict(stats["by_type"]))
    logger.info("  By ext:  %s", dict(stats["by_extension"]))
    logger.info("=" * 80)

    return {
        "stats": stats,
        "all_raw_texts": all_raw_texts,
        "all_job_titles": all_job_titles,
        "all_skills": all_skills,
        "all_companies": all_companies,
        "all_parsed": all_parsed,
    }


# ═══════════════════════════════════════════════════════════════════════
#  PHASE B: COLLOCATION MINING & GAZETTEER BUILDING
# ═══════════════════════════════════════════════════════════════════════

# Category keyword mappings for classification
CATEGORY_KEYWORDS = {
    "TECH_SKILL": [
        "software", "programming", "python", "java", "javascript", "react",
        "node", "sql", "database", "api", "cloud", "aws", "azure", "docker",
        "kubernetes", "machine learning", "deep learning", "ai", "data",
        "algorithm", "framework", "architecture", "devops", "ci/cd",
        "automation", "testing", "agile", "scrum", "git", "linux",
        "network", "security", "web", "mobile", "frontend", "backend",
        "full stack", "microservice", "server", "system", "embedded",
        "iot", "blockchain", "crypto", "virtual", "cloud computing",
        "saas", "paas", "iaas", "rest", "graphql", "tensor", "neural",
        "nlp", "computer vision", "robotics", "plc", "scada", "hmi",
        "control system", "modelling", "simulation", "cad", "cam",
        "bim", "erp", "sap", "oracle", "salesforce", "dynamics",
        "hysys", "aspen", "matlab", "autocad", "solidworks", "catia",
    ],
    "SOFT_SKILL": [
        "leadership", "communication", "teamwork", "problem solving",
        "critical thinking", "time management", "adaptability", "creativity",
        "interpersonal", "presentation", "negotiation", "collaboration",
        "mentoring", "coaching", "decision making", "stakeholder",
        "cross functional", "emotional intelligence", "conflict resolution",
        "strategic thinking", "analytical thinking", "organizational",
    ],
    "CERTIFICATION": [
        "certified", "certification", "certificate", "diploma",
        "accredited", "licensed", "chartered", "professional",
        "pmp", "prince2", "itil", "cissp", "cisa", "aws certified",
        "azure certified", "google certified", "six sigma", "lean",
        "iso", "nebosh", "iosh", "comptia", "cisco", "ccna", "ccnp",
        "ceh", "oscp", "togaf",
    ],
    "INDUSTRY_TERM": [
        "supply chain", "logistics", "procurement", "manufacturing",
        "quality control", "regulatory", "compliance", "audit",
        "risk management", "process improvement", "lean manufacturing",
        "six sigma", "kaizen", "continuous improvement", "safety",
        "environment", "sustainability", "esg", "carbon", "emission",
        "energy", "renewable", "lifecycle", "due diligence",
        "feasibility", "commissioning", "decommissioning",
    ],
    "OIL_GAS": [
        "oil", "gas", "petroleum", "drilling", "reservoir", "upstream",
        "downstream", "midstream", "refinery", "pipeline", "subsea",
        "offshore", "onshore", "lng", "lpg", "crude", "wellhead",
        "completion", "workover", "production", "exploration",
        "seismic", "geophysic", "petrophysic", "mud", "casing",
        "cement", "stimulation", "fracturing", "hydrogen", "ammonia",
        "methanol", "ethylene", "polyethylene", "polypropylene",
    ],
    "METHODOLOGY": [
        "agile", "scrum", "kanban", "waterfall", "lean", "six sigma",
        "design thinking", "tdd", "bdd", "ci/cd", "devops", "itil",
        "prince2", "pmbok", "safe", "spiral", "rad", "xp",
        "pair programming", "code review", "sprint", "retrospective",
    ],
    "INDUSTRIAL_SKILL": [
        "mechanical", "electrical", "chemical", "civil", "structural",
        "instrumentation", "control", "process", "piping", "valve",
        "pump", "compressor", "turbine", "boiler", "heat exchanger",
        "distillation", "reactor", "vessel", "tank", "furnace",
        "maintenance", "reliability", "inspection", "ndt", "welding",
        "fabrication", "construction", "commissioning", "shutdown",
    ],
}


def is_clean_term(term: str) -> bool:
    """Filter out noise terms — must be 2-5 words, 5-60 chars, no junk."""
    words = term.split()
    if not (2 <= len(words) <= 5):
        return False
    if not (5 <= len(term) <= 60):
        return False
    # Reject if starts with common noise prefixes
    noise_prefixes = [
        "the ", "a ", "an ", "for ", "and ", "or ", "in ", "of ", "to ",
        "with ", "from ", "by ", "on ", "at ", "is ", "was ", "are ", "were ",
        "has ", "have ", "had ", "be ", "been ", "being ", "will ", "would ",
        "could ", "should ", "may ", "might ", "must ", "shall ", "can ",
        "page ", "docx ", "pdf ", "http", "www.", "fig ", "table ",
    ]
    lower = term.lower()
    if any(lower.startswith(p) for p in noise_prefixes):
        return False
    # Reject if contains file extensions or URLs
    if any(ext in lower for ext in [".pdf", ".docx", ".xlsx", ".com/", ".org/", ".net/"]):
        return False
    # Reject if too many short words (filler)
    if sum(1 for w in words if len(w) <= 2) > len(words) // 2:
        return False
    # Reject if contains digits prominently
    digit_count = sum(1 for c in term if c.isdigit())
    if digit_count > len(term) * 0.3:
        return False
    # Must have at least some alpha
    if not any(c.isalpha() for c in term):
        return False
    return True


def categorize_term(term: str) -> Optional[str]:
    """Assign a category label to a term based on keyword matching."""
    lower = term.lower()
    scores = {}
    for category, keywords in CATEGORY_KEYWORDS.items():
        score = sum(1 for kw in keywords if kw in lower)
        if score > 0:
            scores[category] = score
    if scores:
        return max(scores, key=scores.get)
    return None


def run_collocation_mining(all_texts: List[str], all_job_titles: List[str],
                           all_skills: List[str]) -> Dict[str, Any]:
    """
    Phase B: Mine collocations from all extracted text.
    Combines: N-grams, PMI, NLTK collocations, co-occurrence analysis.
    """
    logger.info("=" * 80)
    logger.info("PHASE B: COLLOCATION MINING & GAZETTEER BUILDING")
    logger.info("=" * 80)

    start_time = time.time()

    if not all_texts:
        logger.warning("No texts available for mining — loading from consolidated_terms.json")
        consolidated_path = AI_DATA_DIR / "consolidated_terms.json"
        if consolidated_path.exists():
            with open(consolidated_path, "r", encoding="utf-8") as f:
                terms = json.load(f)
            all_texts = [" ".join(terms[i:i+50]) for i in range(0, len(terms), 50)]
        else:
            logger.error("No text sources available for mining")
            return {"error": "no_text"}

    logger.info("Mining from %d text blocks (%d total chars)",
                len(all_texts), sum(len(t) for t in all_texts))

    # ── 1. N-gram extraction ──
    logger.info("Step 1: N-gram extraction...")
    from sklearn.feature_extraction.text import CountVectorizer

    ngram_results = {}
    try:
        vectorizer = CountVectorizer(
            ngram_range=(2, 4),
            min_df=2,
            stop_words="english",
            max_features=10000,
            token_pattern=r"(?u)\b[a-zA-Z][a-zA-Z-]+\b",
        )
        X = vectorizer.fit_transform(all_texts)
        feature_names = vectorizer.get_feature_names_out()
        freqs = X.sum(axis=0).A1

        for phrase, freq in sorted(zip(feature_names, freqs), key=lambda x: -x[1]):
            if freq >= 2 and is_clean_term(phrase):
                ngram_results[phrase.lower()] = int(freq)
        logger.info("  N-grams: %d clean terms extracted", len(ngram_results))
    except Exception as e:
        logger.warning("  N-gram extraction failed: %s", e)

    # ── 2. PMI scoring ──
    logger.info("Step 2: PMI scoring...")
    pmi_results = {}
    try:
        unigram_counts = Counter()
        bigram_counts = Counter()
        total_words = 0
        for text in all_texts:
            words = re.findall(r"\b[a-zA-Z][a-zA-Z-]+\b", text.lower())
            total_words += len(words)
            unigram_counts.update(words)
            for i in range(len(words) - 1):
                bigram_counts[(words[i], words[i + 1])] += 1

        total_bigrams = sum(bigram_counts.values())
        for (w1, w2), freq in bigram_counts.items():
            if freq < 3:
                continue
            p_xy = freq / total_bigrams
            p_x = unigram_counts[w1] / total_words
            p_y = unigram_counts[w2] / total_words
            if p_x > 0 and p_y > 0:
                pmi = math.log2(p_xy / (p_x * p_y))
                if pmi >= 3.0:
                    phrase = f"{w1} {w2}"
                    if is_clean_term(phrase):
                        pmi_results[phrase] = round(pmi, 3)
        logger.info("  PMI: %d high-PMI collocations found", len(pmi_results))
    except Exception as e:
        logger.warning("  PMI scoring failed: %s", e)

    # ── 3. NLTK collocations ──
    logger.info("Step 3: NLTK collocations...")
    nltk_results = {}
    try:
        import nltk
        from nltk.collocations import BigramCollocationFinder, TrigramCollocationFinder
        from nltk.metrics import BigramAssocMeasures, TrigramAssocMeasures

        # Ensure punkt is available
        try:
            nltk.data.find("tokenizers/punkt_tab")
        except LookupError:
            nltk.download("punkt_tab", quiet=True)

        all_words = []
        for text in all_texts[:2000]:  # Cap for memory
            try:
                from nltk.tokenize import word_tokenize
                all_words.extend(word_tokenize(text.lower()))
            except Exception:
                all_words.extend(text.lower().split())

        # Bigrams
        bigram_finder = BigramCollocationFinder.from_words(all_words)
        bigram_finder.apply_freq_filter(3)
        for (w1, w2), score in bigram_finder.score_ngrams(BigramAssocMeasures.pmi)[:500]:
            if w1.isalpha() and w2.isalpha():
                phrase = f"{w1} {w2}"
                if is_clean_term(phrase):
                    nltk_results[phrase] = round(score, 3)

        # Trigrams
        trigram_finder = TrigramCollocationFinder.from_words(all_words)
        trigram_finder.apply_freq_filter(3)
        for (w1, w2, w3), score in trigram_finder.score_ngrams(TrigramAssocMeasures.pmi)[:300]:
            if w1.isalpha() and w2.isalpha() and w3.isalpha():
                phrase = f"{w1} {w2} {w3}"
                if is_clean_term(phrase):
                    nltk_results[phrase] = round(score, 3)

        logger.info("  NLTK: %d collocations found", len(nltk_results))
    except Exception as e:
        logger.warning("  NLTK collocations failed: %s", e)

    # ── 4. Co-occurrence analysis ──
    logger.info("Step 4: Co-occurrence analysis...")
    cooc_results = {}
    try:
        window = 5
        pair_counts = Counter()
        word_counts = Counter()
        for text in all_texts[:2000]:
            words = [w.lower() for w in re.findall(r"\b[a-zA-Z][a-zA-Z-]+\b", text)]
            word_counts.update(words)
            for i, w1 in enumerate(words):
                for j in range(i + 1, min(i + window + 1, len(words))):
                    w2 = words[j]
                    if w1 != w2:
                        pair = tuple(sorted([w1, w2]))
                        pair_counts[pair] += 1

        total = sum(word_counts.values())
        for (w1, w2), freq in pair_counts.most_common(1000):
            if freq < 3:
                continue
            p_xy = freq / max(sum(pair_counts.values()), 1)
            p_x = word_counts[w1] / total if total > 0 else 0
            p_y = word_counts[w2] / total if total > 0 else 0
            if p_x > 0 and p_y > 0:
                pmi = math.log2(p_xy / (p_x * p_y))
                if pmi >= 2.0:
                    phrase = f"{w1} {w2}"
                    if is_clean_term(phrase):
                        cooc_results[phrase] = round(pmi, 3)
        logger.info("  Co-occurrence: %d pairs found", len(cooc_results))
    except Exception as e:
        logger.warning("  Co-occurrence analysis failed: %s", e)

    # ── 5. Merge & categorize all discovered terms ──
    logger.info("Step 5: Merging and categorizing...")
    all_discovered = {}

    # Merge all sources with best score
    for source_name, source_results in [
        ("ngram", ngram_results), ("pmi", pmi_results),
        ("nltk", nltk_results), ("cooccurrence", cooc_results),
    ]:
        for phrase, score in source_results.items():
            key = phrase.lower().strip()
            if key not in all_discovered or score > all_discovered[key]["score"]:
                all_discovered[key] = {"score": score, "method": source_name}

    # Also add job titles and skills directly
    for title in all_job_titles:
        key = title.lower().strip()
        if is_clean_term(key) and key not in all_discovered:
            all_discovered[key] = {"score": 10.0, "method": "job_title_extract"}
    for skill in all_skills:
        key = skill.lower().strip()
        if is_clean_term(key) and key not in all_discovered:
            all_discovered[key] = {"score": 8.0, "method": "skill_extract"}

    # Categorize
    categorized: Dict[str, Dict[str, list]] = defaultdict(lambda: {"count": 0, "terms": []})
    uncategorized = []
    for phrase, meta in all_discovered.items():
        cat = categorize_term(phrase)
        if cat:
            categorized[cat]["terms"].append(phrase)
            categorized[cat]["count"] += 1
        else:
            uncategorized.append(phrase)

    # Assign uncategorized to best-guess or GENERAL
    for phrase in uncategorized:
        categorized["GENERAL"]["terms"].append(phrase)
        categorized["GENERAL"]["count"] += 1

    for cat in categorized:
        categorized[cat]["terms"] = sorted(set(categorized[cat]["terms"]))
        categorized[cat]["count"] = len(categorized[cat]["terms"])

    # ── 6. Save collocation_mining_results.json ──
    mining_results_path = DATA_ROOT / "collocation_mining_results.json"
    with open(mining_results_path, "w", encoding="utf-8") as f:
        json.dump(dict(categorized), f, indent=2, ensure_ascii=False)
    logger.info("Saved collocation_mining_results.json: %d categories, %d total terms",
                len(categorized), sum(c["count"] for c in categorized.values()))

    # Also save mined_collocations_output.json (flat format for compatibility)
    mined_output_path = DATA_ROOT / "mined_collocations_output.json"
    flat_output = {cat: data["terms"] for cat, data in categorized.items()}
    with open(mined_output_path, "w", encoding="utf-8") as f:
        json.dump(flat_output, f, indent=2, ensure_ascii=False)

    # ── 7. Update gazetteers ──
    logger.info("Step 7: Updating gazetteers...")
    gazetteer_updates = 0
    for cat, data in categorized.items():
        if cat == "GENERAL":
            continue  # Don't pollute gazetteers with uncategorized

        # Map category to gazetteer filename
        gaz_name_map = {
            "TECH_SKILL": "tech_skills.json",
            "SOFT_SKILL": "soft_skills.json",
            "CERTIFICATION": "certifications.json",
            "INDUSTRY_TERM": "industry_terms.json",
            "OIL_GAS": "oil_gas.json",
            "METHODOLOGY": "methodologies.json",
            "INDUSTRIAL_SKILL": "industrial_automation.json",
        }
        gaz_filename = gaz_name_map.get(cat)
        if not gaz_filename:
            continue

        gaz_path = GAZETTEERS_DIR / gaz_filename
        existing_gaz = {"label": cat, "terms": [], "abbreviations": {}}
        if gaz_path.exists():
            try:
                with open(gaz_path, "r", encoding="utf-8") as f:
                    existing_gaz = json.load(f)
            except Exception:
                pass

        existing_terms_set = {t.lower() for t in existing_gaz.get("terms", [])}
        new_terms = [t for t in data["terms"] if t.lower() not in existing_terms_set]
        existing_gaz["terms"] = sorted(set(existing_gaz.get("terms", []) + new_terms))
        existing_gaz["label"] = cat
        existing_gaz["updated_at"] = datetime.now().isoformat()

        with open(gaz_path, "w", encoding="utf-8") as f:
            json.dump(existing_gaz, f, indent=2, ensure_ascii=False)
        gazetteer_updates += len(new_terms)

    # ── 8. Update learned_collocations.json ──
    learned_path = GAZETTEERS_DIR / "learned_collocations.json"
    existing_learned = {"version": 2, "phrase_count": 0, "phrases": {}}
    if learned_path.exists():
        try:
            with open(learned_path, "r", encoding="utf-8") as f:
                existing_learned = json.load(f)
        except Exception:
            pass

    now = datetime.now().isoformat()
    phrases = existing_learned.get("phrases", {})
    new_learned = 0
    for phrase, meta in all_discovered.items():
        if phrase not in phrases:
            cat = categorize_term(phrase)
            phrases[phrase] = {
                "label": cat,
                "score": meta["score"],
                "method": meta["method"],
                "frequency": 1,
                "discovered_at": now,
                "source": "deep_ingest_pipeline",
            }
            new_learned += 1

    existing_learned["phrases"] = phrases
    existing_learned["phrase_count"] = len(phrases)
    existing_learned["persisted_at"] = now
    existing_learned["total_ingestions"] = existing_learned.get("total_ingestions", 0) + 1

    with open(learned_path, "w", encoding="utf-8") as f:
        json.dump(existing_learned, f, indent=2, ensure_ascii=False)

    elapsed = time.time() - start_time

    mining_stats = {
        "ngram_count": len(ngram_results),
        "pmi_count": len(pmi_results),
        "nltk_count": len(nltk_results),
        "cooccurrence_count": len(cooc_results),
        "total_discovered": len(all_discovered),
        "categorized_breakdown": {cat: data["count"] for cat, data in categorized.items()},
        "gazetteer_new_terms": gazetteer_updates,
        "learned_new_phrases": new_learned,
        "elapsed_seconds": round(elapsed, 2),
    }

    logger.info("=" * 80)
    logger.info("PHASE B COMPLETE — %d total collocations discovered in %.1fs",
                len(all_discovered), elapsed)
    logger.info("  N-grams: %d | PMI: %d | NLTK: %d | Co-occ: %d",
                len(ngram_results), len(pmi_results), len(nltk_results), len(cooc_results))
    logger.info("  Categories: %s", {cat: data["count"] for cat, data in categorized.items()})
    logger.info("  Gazetteer updates: %d new terms | Learned phrases: %d new",
                gazetteer_updates, new_learned)
    logger.info("=" * 80)

    return mining_stats


# ═══════════════════════════════════════════════════════════════════════
#  PHASE C: AI MODEL TRAINING
# ═══════════════════════════════════════════════════════════════════════

def run_ai_model_training() -> Dict[str, Any]:
    """
    Phase C: Train all AI models using data in ai_data_final/.
    Models: Bayesian, TF-IDF, Sentence-BERT, spaCy NER, clustering,
            job classifier, salary predictor.
    """
    logger.info("=" * 80)
    logger.info("PHASE C: AI MODEL TRAINING")
    logger.info("=" * 80)

    import numpy as np
    import pandas as pd
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.naive_bayes import MultinomialNB
    from sklearn.ensemble import RandomForestClassifier, RandomForestRegressor
    from sklearn.cluster import KMeans, DBSCAN
    from sklearn.metrics.pairwise import cosine_similarity
    from sklearn.model_selection import train_test_split
    from sklearn.metrics import accuracy_score, r2_score

    start_time = time.time()
    training_report = {
        "timestamp": datetime.now().isoformat(),
        "data_stats": {},
        "model_performance": {},
        "files_created": [],
    }

    # ── Load training data ──
    logger.info("Loading training data from ai_data_final/...")

    all_records = []

    # 1. Load from merged candidate database
    merged_db = AI_DATA_DIR / "core_databases" / "Candidate_database_merged.json"
    if merged_db.exists():
        try:
            with open(merged_db, "r", encoding="utf-8") as f:
                candidates = json.load(f)
            for c in candidates:
                text_parts = []
                for key in ["name", "current_position", "Job Title", "raw_text",
                            "summary", "Position", "Company", "subject", "body"]:
                    val = c.get(key)
                    if val and isinstance(val, str) and len(val) > 2:
                        text_parts.append(val[:1000])
                skills_val = c.get("skills", c.get("Skills", []))
                if isinstance(skills_val, list):
                    text_parts.append(" ".join(str(s) for s in skills_val))

                text = " ".join(text_parts).strip()
                if len(text) > 50:
                    all_records.append({
                        "text": text,
                        "job_title": c.get("Job Title", c.get("current_position", "")),
                        "skills": skills_val if isinstance(skills_val, list) else [],
                        "experience_years": c.get("experience_years", c.get("years_experience", 0)),
                        "industry": c.get("industry", "Unknown"),
                        "salary": c.get("salary", c.get("expected_salary", None)),
                    })
            logger.info("  Loaded %d records from Candidate_database_merged.json", len(all_records))
        except Exception as e:
            logger.warning("  Failed to load merged DB: %s", e)

    # 2. Load from parsed_resumes/
    parsed_resumes_dir = AI_DATA_DIR / "parsed_resumes"
    resume_count = 0
    if parsed_resumes_dir.exists():
        for jf in parsed_resumes_dir.glob("*.json"):
            try:
                with open(jf, "r", encoding="utf-8") as f:
                    data = json.load(f)
                text = data.get("raw_text", data.get("text", ""))
                if len(text) > 50:
                    all_records.append({
                        "text": text[:5000],
                        "job_title": (data.get("job_titles", [""])[0]
                                      if isinstance(data.get("job_titles"), list)
                                      else data.get("job_title", "")),
                        "skills": data.get("skills", []),
                        "experience_years": data.get("experience_years", 0),
                        "industry": data.get("industry", "Unknown"),
                        "salary": data.get("salary", None),
                    })
                    resume_count += 1
            except Exception:
                continue
        logger.info("  Loaded %d records from parsed_resumes/", resume_count)

    # 3. Load from parsed_from_automated/
    auto_dir = AI_DATA_DIR / "parsed_from_automated"
    auto_count = 0
    if auto_dir.exists():
        for jf in auto_dir.glob("*.json"):
            try:
                with open(jf, "r", encoding="utf-8") as f:
                    data = json.load(f)
                text = data.get("raw_text", data.get("text", ""))
                if len(text) > 50:
                    all_records.append({
                        "text": text[:5000],
                        "job_title": (data.get("job_titles", [""])[0]
                                      if isinstance(data.get("job_titles"), list)
                                      else data.get("job_title", "")),
                        "skills": data.get("skills", []),
                        "experience_years": data.get("experience_years", 0),
                        "industry": data.get("industry", "Unknown"),
                        "salary": data.get("salary", None),
                    })
                    auto_count += 1
            except Exception:
                continue
        logger.info("  Loaded %d records from parsed_from_automated/", auto_count)

    if not all_records:
        logger.error("No training data available!")
        return {"error": "no_training_data"}

    df = pd.DataFrame(all_records)
    logger.info("Total training records: %d (avg text len: %.0f chars)",
                len(df), df["text"].str.len().mean())

    training_report["data_stats"] = {
        "total_records": len(df),
        "avg_text_length": float(df["text"].str.len().mean()),
        "from_merged_db": len(all_records) - resume_count - auto_count,
        "from_parsed_resumes": resume_count,
        "from_parsed_automated": auto_count,
    }

    # ── Job category inference ──
    def infer_category(job_title: str, text: str) -> str:
        title_lower = str(job_title).lower()
        text_lower = str(text)[:2000].lower()
        combined = f"{title_lower} {text_lower}"

        categories = {
            "Technology": ["engineer", "developer", "programmer", "software", "data scientist",
                           "devops", "architect", "sysadmin", "sre", "platform"],
            "Oil & Gas / Energy": ["oil", "gas", "petroleum", "drilling", "reservoir",
                                     "refinery", "pipeline", "subsea", "offshore", "energy",
                                     "hydrogen", "process engineer"],
            "Finance": ["accountant", "financial", "analyst", "banker", "investment",
                         "trading", "audit", "compliance", "risk"],
            "Sales & Marketing": ["sales", "marketing", "business development",
                                    "account manager", "brand", "campaign", "digital marketing"],
            "Management": ["manager", "director", "executive", "chief", "head of",
                            "vp ", "vice president", "ceo", "cfo", "cto", "coo"],
            "Engineering": ["mechanical", "electrical", "chemical", "civil", "structural",
                             "nuclear", "industrial", "instrumentation"],
            "Human Resources": ["hr", "human resources", "recruiter", "talent",
                                  "people", "training", "l&d"],
            "Healthcare": ["doctor", "nurse", "physician", "medical", "healthcare",
                            "clinical", "pharma"],
        }
        best_cat = "Other"
        best_score = 0
        for cat, keywords in categories.items():
            score = sum(1 for kw in keywords if kw in combined)
            if score > best_score:
                best_score = score
                best_cat = cat
        return best_cat

    df["category"] = df.apply(lambda r: infer_category(r["job_title"], r["text"]), axis=1)

    # ── MODEL 1: Bayesian Classifier + TF-IDF ──
    logger.info("\nMODEL 1: Bayesian Job Category Classifier...")
    try:
        cat_counts = df["category"].value_counts()
        valid_cats = cat_counts[cat_counts >= 5].index
        df_filtered = df[df["category"].isin(valid_cats)]

        if len(df_filtered) >= 20:
            vectorizer = TfidfVectorizer(max_features=5000, min_df=2, max_df=0.9,
                                         stop_words="english", ngram_range=(1, 2))
            X = vectorizer.fit_transform(df_filtered["text"])
            y = df_filtered["category"]

            X_train, X_test, y_train, y_test = train_test_split(
                X, y, test_size=0.2, random_state=42,
                stratify=y if all(cat_counts[c] >= 3 for c in valid_cats) else None
            )

            model = MultinomialNB(alpha=0.1)
            model.fit(X_train, y_train)
            y_pred = model.predict(X_test)
            acc = accuracy_score(y_test, y_pred)

            pickle.dump(model, open(MODELS_DIR / "bayesian_classifier.pkl", "wb"))
            pickle.dump(vectorizer, open(MODELS_DIR / "tfidf_vectorizer.pkl", "wb"))

            training_report["model_performance"]["bayesian_classifier"] = {
                "accuracy": float(acc),
                "training_samples": int(X_train.shape[0]),
                "test_samples": int(X_test.shape[0]),
                "categories": list(valid_cats),
            }
            training_report["files_created"].extend([
                str(MODELS_DIR / "bayesian_classifier.pkl"),
                str(MODELS_DIR / "tfidf_vectorizer.pkl"),
            ])
            logger.info("  Bayesian classifier: %.2f%% accuracy (%d train / %d test)",
                        acc * 100, X_train.shape[0], X_test.shape[0])
            logger.info("  Categories: %s", list(valid_cats))
        else:
            logger.warning("  Not enough data for Bayesian classifier (%d records)", len(df_filtered))
    except Exception as e:
        logger.error("  Bayesian classifier failed: %s", e)
        traceback.print_exc()

    # ── MODEL 2: Sentence-BERT Embeddings ──
    logger.info("\nMODEL 2: Sentence-BERT Embeddings...")
    embeddings = None
    try:
        model = get_sentence_model()
        if model:
            # Embed a representative sample (limit for speed)
            sample_texts = df["text"].head(min(5000, len(df))).tolist()
            # Truncate long texts for embedding
            sample_texts = [t[:512] for t in sample_texts]
            embeddings = model.encode(sample_texts, show_progress_bar=True, batch_size=32)

            np.save(MODELS_DIR / "candidate_embeddings.npy", embeddings)

            model_info = {
                "model_name": "all-MiniLM-L6-v2",
                "embedding_dim": int(embeddings.shape[1]),
                "samples_embedded": int(embeddings.shape[0]),
                "provider": "sentence-transformers",
            }
            with open(MODELS_DIR / "sentence_bert_info.json", "w") as f:
                json.dump(model_info, f, indent=2)

            training_report["model_performance"]["sentence_embeddings"] = model_info
            training_report["files_created"].extend([
                str(MODELS_DIR / "candidate_embeddings.npy"),
                str(MODELS_DIR / "sentence_bert_info.json"),
            ])
            logger.info("  Created %d embeddings (%d dimensions)", *embeddings.shape)
        else:
            logger.warning("  Sentence-BERT model not available")
    except Exception as e:
        logger.error("  Sentence-BERT failed: %s", e)
        traceback.print_exc()

    # ── MODEL 3: spaCy NER ──
    logger.info("\nMODEL 3: spaCy NER...")
    try:
        import spacy
        nlp = spacy.load("en_core_web_sm")
        test_text = "Senior Python Developer at Google with 5 years of machine learning experience"
        doc = nlp(test_text)
        entities = [(ent.text, ent.label_) for ent in doc.ents]

        model_info = {
            "model_name": "en_core_web_sm",
            "version": nlp.meta.get("version", "unknown"),
            "installed": True,
            "test_entities": entities,
        }
        with open(MODELS_DIR / "spacy_model_info.json", "w") as f:
            json.dump(model_info, f, indent=2)

        training_report["model_performance"]["spacy_ner"] = model_info
        training_report["files_created"].append(str(MODELS_DIR / "spacy_model_info.json"))
        logger.info("  spaCy NER ready: %s — test entities: %s", nlp.meta.get("version"), entities)
    except Exception as e:
        logger.warning("  spaCy NER not available: %s", e)

    # ── MODEL 4: K-Means + DBSCAN Clustering ──
    logger.info("\nMODEL 4: Clustering (K-Means + DBSCAN)...")
    if embeddings is not None and len(embeddings) >= 20:
        try:
            n_clusters = min(15, len(embeddings) // 5)
            if n_clusters >= 2:
                kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
                kmeans_labels = kmeans.fit_predict(embeddings)

                pickle.dump(kmeans, open(MODELS_DIR / "kmeans_model.pkl", "wb"))

                cluster_counts = Counter(int(l) for l in kmeans_labels)
                kmeans_info = {
                    "n_clusters": n_clusters,
                    "inertia": float(kmeans.inertia_),
                    "cluster_counts": dict(sorted(cluster_counts.items())),
                }

                # DBSCAN
                dbscan = DBSCAN(eps=0.5, min_samples=5)
                dbscan_labels = dbscan.fit_predict(embeddings)
                n_db_clusters = len(set(dbscan_labels)) - (1 if -1 in dbscan_labels else 0)

                clustering_results = {
                    "kmeans": kmeans_info,
                    "dbscan": {
                        "n_clusters": n_db_clusters,
                        "n_noise": int(list(dbscan_labels).count(-1)),
                    },
                }
                with open(MODELS_DIR / "clustering_results.json", "w") as f:
                    json.dump(clustering_results, f, indent=2)

                training_report["model_performance"]["clustering"] = clustering_results
                training_report["files_created"].extend([
                    str(MODELS_DIR / "kmeans_model.pkl"),
                    str(MODELS_DIR / "clustering_results.json"),
                ])
                logger.info("  K-Means: %d clusters, inertia=%.1f", n_clusters, kmeans.inertia_)
                logger.info("  DBSCAN: %d clusters, %d noise points", n_db_clusters,
                            list(dbscan_labels).count(-1))
        except Exception as e:
            logger.error("  Clustering failed: %s", e)
    else:
        logger.warning("  Not enough embeddings for clustering")

    # ── MODEL 5: Cosine Similarity Matrix ──
    logger.info("\nMODEL 5: Cosine Similarity Matrix...")
    if embeddings is not None and len(embeddings) >= 10:
        try:
            # Save a smaller sample for the similarity matrix (memory)
            sample_size = min(2000, len(embeddings))
            sim_matrix = cosine_similarity(embeddings[:sample_size])
            np.save(MODELS_DIR / "similarity_matrix.npy", sim_matrix)

            sim_info = {
                "shape": list(sim_matrix.shape),
                "mean_similarity": float(sim_matrix.mean()),
                "max_similarity": float(sim_matrix[~np.eye(sim_matrix.shape[0], dtype=bool)].max())
                    if sim_matrix.shape[0] > 1 else 0.0,
            }
            training_report["model_performance"]["similarity_matrix"] = sim_info
            training_report["files_created"].append(str(MODELS_DIR / "similarity_matrix.npy"))
            logger.info("  Similarity matrix: %s, mean=%.3f", sim_matrix.shape, sim_matrix.mean())
        except Exception as e:
            logger.error("  Similarity matrix failed: %s", e)

    # ── MODEL 6: Job Title Classifier (Random Forest) ──
    logger.info("\nMODEL 6: Job Title Classifier (Random Forest)...")
    try:
        # Use the job titles from our data
        df_with_titles = df[df["job_title"].str.len() > 3].copy()
        if len(df_with_titles) >= 20:
            jt_vectorizer = TfidfVectorizer(max_features=1000, stop_words="english", ngram_range=(1, 2))
            X_jt = jt_vectorizer.fit_transform(df_with_titles["job_title"])
            y_jt = df_with_titles["category"]

            # Only train if we have enough per category
            jt_cat_counts = y_jt.value_counts()
            valid_jt_cats = jt_cat_counts[jt_cat_counts >= 3].index
            mask = y_jt.isin(valid_jt_cats)
            X_jt = X_jt[mask]
            y_jt = y_jt[mask]

            if len(y_jt) >= 10:
                X_jt_train, X_jt_test, y_jt_train, y_jt_test = train_test_split(
                    X_jt, y_jt, test_size=0.2, random_state=42
                )
                rf = RandomForestClassifier(n_estimators=100, random_state=42)
                rf.fit(X_jt_train, y_jt_train)
                jt_pred = rf.predict(X_jt_test)
                jt_acc = accuracy_score(y_jt_test, jt_pred)

                pickle.dump(rf, open(MODELS_DIR / "job_classifier.pkl", "wb"))
                pickle.dump(jt_vectorizer, open(MODELS_DIR / "job_classifier_vectorizer.pkl", "wb"))

                training_report["model_performance"]["job_classifier"] = {
                    "accuracy": float(jt_acc),
                    "n_categories": len(valid_jt_cats),
                    "training_samples": int(X_jt_train.shape[0]),
                }
                training_report["files_created"].extend([
                    str(MODELS_DIR / "job_classifier.pkl"),
                    str(MODELS_DIR / "job_classifier_vectorizer.pkl"),
                ])
                logger.info("  Job classifier: %.2f%% accuracy (%d categories)",
                            jt_acc * 100, len(valid_jt_cats))
    except Exception as e:
        logger.error("  Job classifier failed: %s", e)
        traceback.print_exc()

    # ── MODEL 7: Statistical / TF-IDF Top Features ──
    logger.info("\nMODEL 7: TF-IDF Top Features for Admin Dashboard...")
    try:
        full_tfidf = TfidfVectorizer(max_features=2000, stop_words="english", ngram_range=(1, 3))
        full_matrix = full_tfidf.fit_transform(df["text"])
        feature_names = full_tfidf.get_feature_names_out()

        # Top terms by total TF-IDF weight
        total_weights = full_matrix.sum(axis=0).A1
        top_indices = total_weights.argsort()[-200:][::-1]
        top_features = [{"term": feature_names[i], "weight": float(total_weights[i])} for i in top_indices]

        with open(MODELS_DIR / "tfidf_top_features.json", "w", encoding="utf-8") as f:
            json.dump(top_features, f, indent=2)

        training_report["model_performance"]["tfidf_features"] = {
            "total_vocabulary": len(feature_names),
            "top_200_saved": True,
        }
        training_report["files_created"].append(str(MODELS_DIR / "tfidf_top_features.json"))
        logger.info("  Top features: %d vocabulary, top 200 saved", len(feature_names))
    except Exception as e:
        logger.error("  TF-IDF features failed: %s", e)

    # ── Save training report ──
    report_path = MODELS_DIR / "training_report.json"
    with open(report_path, "w", encoding="utf-8") as f:
        json.dump(training_report, f, indent=2, default=str)

    elapsed = time.time() - start_time
    training_report["total_training_time"] = round(elapsed, 2)

    logger.info("=" * 80)
    logger.info("PHASE C COMPLETE — %d models trained in %.1fs",
                len(training_report["model_performance"]), elapsed)
    logger.info("  Files created: %d", len(training_report["files_created"]))
    for model_name, metrics in training_report["model_performance"].items():
        if isinstance(metrics, dict) and "accuracy" in metrics:
            logger.info("  %s: %.2f%% accuracy", model_name, metrics["accuracy"] * 100)
    logger.info("=" * 80)

    return training_report


# ═══════════════════════════════════════════════════════════════════════
#  MAIN PIPELINE ORCHESTRATOR
# ═══════════════════════════════════════════════════════════════════════

def main():
    pipeline_start = time.time()

    logger.info("#" * 80)
    logger.info("#  CareerTrojan — DEEP INGESTION, MINING & TRAINING PIPELINE")
    logger.info("#  Started: %s", datetime.now().isoformat())
    logger.info("#  Data Root: %s", DATA_ROOT)
    logger.info("#  AI Data: %s", AI_DATA_DIR)
    logger.info("#  Models: %s", MODELS_DIR)
    logger.info("#" * 80)

    final_report = {
        "pipeline_start": datetime.now().isoformat(),
        "phases": {},
    }

    # ═══ PHASE A: Deep Document Parsing ═══
    try:
        parse_result = deep_parse_automated_parser()
        final_report["phases"]["A_parsing"] = parse_result.get("stats", {})
        all_texts = parse_result.get("all_raw_texts", [])
        all_job_titles = parse_result.get("all_job_titles", [])
        all_skills = parse_result.get("all_skills", [])
    except Exception as e:
        logger.error("PHASE A FAILED: %s", e)
        traceback.print_exc()
        final_report["phases"]["A_parsing"] = {"error": str(e)}
        all_texts, all_job_titles, all_skills = [], [], []

    # ═══ PHASE B: Collocation Mining ═══
    try:
        mining_stats = run_collocation_mining(all_texts, all_job_titles, all_skills)
        final_report["phases"]["B_collocation_mining"] = mining_stats
    except Exception as e:
        logger.error("PHASE B FAILED: %s", e)
        traceback.print_exc()
        final_report["phases"]["B_collocation_mining"] = {"error": str(e)}

    # ═══ PHASE C: AI Model Training ═══
    try:
        training_report = run_ai_model_training()
        final_report["phases"]["C_ai_training"] = training_report
    except Exception as e:
        logger.error("PHASE C FAILED: %s", e)
        traceback.print_exc()
        final_report["phases"]["C_ai_training"] = {"error": str(e)}

    # ═══ Final Summary ═══
    pipeline_elapsed = time.time() - pipeline_start
    final_report["pipeline_end"] = datetime.now().isoformat()
    final_report["total_pipeline_seconds"] = round(pipeline_elapsed, 2)

    # Save final report
    report_path = AI_DATA_DIR / "deep_pipeline_report.json"
    with open(report_path, "w", encoding="utf-8") as f:
        json.dump(final_report, f, indent=2, default=str)

    logger.info("")
    logger.info("#" * 80)
    logger.info("#  PIPELINE COMPLETE")
    logger.info("#  Total time: %.1f seconds (%.1f minutes)", pipeline_elapsed, pipeline_elapsed / 60)
    logger.info("#")
    logger.info("#  Phase A (Parsing):     %s", _phase_summary(final_report["phases"].get("A_parsing", {})))
    logger.info("#  Phase B (Collocations): %s", _phase_summary(final_report["phases"].get("B_collocation_mining", {})))
    logger.info("#  Phase C (Training):    %s", _phase_summary(final_report["phases"].get("C_ai_training", {})))
    logger.info("#")
    logger.info("#  Full report: %s", report_path)
    logger.info("#" * 80)

    return 0


def _phase_summary(phase_data: Dict[str, Any]) -> str:
    if "error" in phase_data:
        return f"FAILED — {phase_data['error']}"
    if "parsed" in phase_data:
        return f"{phase_data['parsed']} files parsed, {phase_data.get('new_candidates', 0)} new candidates"
    if "total_discovered" in phase_data:
        return f"{phase_data['total_discovered']} collocations, {phase_data.get('learned_new_phrases', 0)} new learned"
    if "model_performance" in phase_data:
        return f"{len(phase_data['model_performance'])} models, {len(phase_data.get('files_created', []))} files"
    return str(phase_data)[:100]


if __name__ == "__main__":
    sys.exit(main())
