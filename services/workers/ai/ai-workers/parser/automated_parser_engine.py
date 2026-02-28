"""
Automated Parser Engine - Production Ready Data Extraction
===========================================================

Converts raw files (PDF, DOCX, CSV, etc.) to structured JSON
Handles 2000+ files in automated_parser/ directory
Output: JSON files in automated_parser/completed/

Usage:
    python automated_parser_engine.py

    # or

    from automated_parser_engine import AutomatedParserEngine
    parser = AutomatedParserEngine()
    results = parser.run()
"""

import json
import hashlib
import re
import zipfile
import csv
import io
from pathlib import Path
from datetime import datetime
import logging
from typing import Dict, List, Any, Optional
import sys

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class AutomatedParserEngine:
    """
    Production-ready parser for multiple file formats
    Supports: PDF, DOCX, DOC, CSV, XLSX, MSG, TXT, JSON
    """

    def __init__(self, parser_root='automated_parser', output_root='ai_data_final'):
        """Initialize parser with folder paths"""
        self.parser_root = Path(parser_root)
        self.incoming = self.parser_root / 'incoming'
        self.completed = self.parser_root / 'completed'
        self.output_root = Path(output_root)

        # Create output folders if needed
        self.completed.mkdir(parents=True, exist_ok=True)
        self.output_root.mkdir(parents=True, exist_ok=True)
        self.processed_index_file = self.output_root / '_processed_sources.txt'
        self.failed_index_file = self.output_root / '_failed_sources.json'
        self.status_manifest_file = self.output_root / 'parser_ingestion_status.json'
        self.status_summary_file = self.output_root / 'parser_ingestion_status_summary.json'
        self.processed_sources = self._load_processed_sources()
        self.failed_sources = self._load_failed_sources()
        self.optional_dependency_status = self._detect_optional_dependencies()

        self.results = {
            'total_files': 0,
            'processed': 0,
            'skipped': 0,
            'skipped_existing': 0,
            'errors': [],
            'file_types': {},
            'created_at': datetime.now().isoformat()
        }

        logger.info(f"Parser initialized: {self.parser_root}")
        logger.info(f"Output folder (legacy): {self.completed}")
        logger.info(f"Output folder (primary): {self.output_root}")
        logger.info(f"Loaded processed index entries: {len(self.processed_sources)}")
        logger.info(f"Loaded failed index entries: {len(self.failed_sources)}")

    def _detect_optional_dependencies(self) -> Dict[str, bool]:
        status = {
            'xlrd': False,
            'extract_msg': False,
            'openpyxl': False,
            'python_docx': False,
        }
        try:
            import xlrd  # type: ignore
            status['xlrd'] = True
        except Exception:
            pass
        try:
            import extract_msg  # type: ignore
            status['extract_msg'] = True
        except Exception:
            pass
        try:
            import openpyxl  # type: ignore
            status['openpyxl'] = True
        except Exception:
            pass
        try:
            import docx  # type: ignore
            status['python_docx'] = True
        except Exception:
            pass
        return status

    def _discover_all_parseable_sources(self) -> List[Path]:
        """Discover all parseable source files regardless of processed index state."""
        extensions = {
            '.pdf', '.docx', '.doc', '.csv', '.xlsx', '.xls', '.zip', '.msg', '.eml', '.mbox',
            '.txt', '.json', '.rtf', '.odt', '.ods', '.xml', '.yaml', '.yml', '.html', '.htm', '.md'
        }
        ignored_folders = {'incoming', 'completed', '__pycache__', '.git', '.venv', 'node_modules'}
        skip_files = {'data_ingestion_tracker.py', 'README.md', '.gitignore', 'desktop.ini', 'Thumbs.db'}

        discovered: List[Path] = []
        for file_path in self.parser_root.rglob('*'):
            if not file_path.is_file():
                continue
            if any(ignored in file_path.parts for ignored in ignored_folders):
                continue
            if file_path.name in skip_files:
                continue
            if file_path.name.startswith('~$') or file_path.name.startswith('.~lock.'):
                continue
            if '~$' in file_path.name:
                continue
            if file_path.suffix.lower() not in extensions:
                continue
            discovered.append(file_path)

        return discovered

    def _status_for_source(self, file_path: Path) -> str:
        rel = str(file_path.relative_to(self.parser_root))
        if rel in self.processed_sources:
            return 'ingested'

        if rel in self.failed_sources:
            return 'pending_parse_error'

        ext = file_path.suffix.lower()
        if ext == '.xls' and not self.optional_dependency_status.get('xlrd', False):
            return 'blocked_missing_dependency'
        if ext == '.msg' and not self.optional_dependency_status.get('extract_msg', False):
            return 'blocked_missing_dependency'
        if ext == '.xlsx' and not self.optional_dependency_status.get('openpyxl', False):
            return 'blocked_missing_dependency'
        if ext == '.docx' and not self.optional_dependency_status.get('python_docx', False):
            return 'blocked_missing_dependency'

        return 'pending_not_attempted'

    def generate_ingestion_status_manifest(self) -> Path:
        """Generate full file-level ingestion status manifest for monitoring and trap hooks."""
        all_sources = self._discover_all_parseable_sources()
        rows = []
        status_counts = {
            'ingested': 0,
            'pending_not_attempted': 0,
            'pending_parse_error': 0,
            'blocked_missing_dependency': 0,
        }

        for source in all_sources:
            rel = str(source.relative_to(self.parser_root))
            status = self._status_for_source(source)
            status_counts[status] = status_counts.get(status, 0) + 1
            rows.append(
                {
                    'source_path': rel,
                    'file_name': source.name,
                    'file_type': source.suffix.lower(),
                    'size_bytes': source.stat().st_size if source.exists() else 0,
                    'status': status,
                    'indexed': rel in self.processed_sources,
                    'failure_reason': (self.failed_sources.get(rel) or {}).get('reason', ''),
                    'last_attempt': (self.failed_sources.get(rel) or {}).get('last_attempt', ''),
                }
            )

        trap_flags = []
        if status_counts.get('pending_not_attempted', 0) > 0:
            trap_flags.append('pending_parser_data_detected')
        if status_counts.get('pending_parse_error', 0) > 0:
            trap_flags.append('parser_parse_error_backlog_detected')
        if status_counts.get('blocked_missing_dependency', 0) > 0:
            trap_flags.append('parser_blocked_by_missing_dependency')

        payload = {
            'created_at': datetime.now().isoformat(),
            'parser_root': str(self.parser_root),
            'output_root': str(self.output_root),
            'processed_index_file': str(self.processed_index_file),
            'processed_index_entries': len(self.processed_sources),
            'optional_dependencies': self.optional_dependency_status,
            'status_counts': status_counts,
            'trap_flags': trap_flags,
            'records': rows,
        }

        summary_payload = {
            'created_at': payload['created_at'],
            'parser_root': payload['parser_root'],
            'output_root': payload['output_root'],
            'processed_index_entries': payload['processed_index_entries'],
            'optional_dependencies': payload['optional_dependencies'],
            'status_counts': status_counts,
            'trap_flags': trap_flags,
        }

        with open(self.status_manifest_file, 'w', encoding='utf-8') as f:
            json.dump(payload, f, indent=2, ensure_ascii=False)
        with open(self.status_summary_file, 'w', encoding='utf-8') as f:
            json.dump(summary_payload, f, indent=2, ensure_ascii=False)

        logger.info(f"Ingestion status manifest saved: {self.status_manifest_file}")
        logger.info(f"Ingestion status summary saved: {self.status_summary_file}")
        return self.status_manifest_file

    def _load_processed_sources(self) -> set:
        """Load previously processed source paths from index file."""
        if not self.processed_index_file.exists():
            return set()

        processed = set()
        try:
            with open(self.processed_index_file, 'r', encoding='utf-8') as f:
                for line in f:
                    source = line.strip()
                    if source:
                        processed.add(source)
        except Exception as e:
            logger.warning(f"Could not load processed index: {e}")
        return processed

    def _load_failed_sources(self) -> Dict[str, Dict[str, Any]]:
        """Load previously failed source paths and reasons from index file."""
        if not self.failed_index_file.exists():
            return {}

        try:
            raw = json.loads(self.failed_index_file.read_text(encoding='utf-8'))
            if isinstance(raw, dict):
                normalized: Dict[str, Dict[str, Any]] = {}
                for source_path, details in raw.items():
                    if not isinstance(source_path, str):
                        continue
                    if not isinstance(details, dict):
                        details = {}
                    normalized[source_path] = {
                        'status': str(details.get('status') or 'error'),
                        'reason': str(details.get('reason') or ''),
                        'last_attempt': str(details.get('last_attempt') or ''),
                    }
                return normalized
        except Exception as e:
            logger.warning(f"Could not load failed index: {e}")
        return {}

    def _save_failed_sources(self) -> None:
        try:
            with open(self.failed_index_file, 'w', encoding='utf-8') as f:
                json.dump(self.failed_sources, f, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.warning(f"Could not save failed index: {e}")

    def _mark_failed_source(self, source_path: str, status: str, reason: str) -> None:
        self.failed_sources[source_path] = {
            'status': status,
            'reason': reason,
            'last_attempt': datetime.now().isoformat(),
        }

    def _clear_failed_source(self, source_path: str) -> None:
        if source_path in self.failed_sources:
            self.failed_sources.pop(source_path, None)

    def _append_processed_source(self, source_path: str) -> None:
        """Append a processed source path to index file."""
        try:
            with open(self.processed_index_file, 'a', encoding='utf-8') as f:
                f.write(source_path + '\n')
        except Exception as e:
            logger.warning(f"Could not update processed index: {e}")

    def discover_source_files(self) -> Dict[str, List[Path]]:
        """
        Discover all parseable files in parser root RECURSIVELY
        Returns: dict mapping file type to list of paths
        """
        extensions = {
            '.pdf': 'PDF',
            '.docx': 'DOCX',
            '.doc': 'DOC',
            '.csv': 'CSV',
            '.xlsx': 'XLSX',
            '.xls': 'XLS',      # NEW: Legacy Excel
            '.zip': 'ZIP',      # NEW: Archive bundles (e.g., LinkedIn exports)
            '.msg': 'MSG',
            '.eml': 'EML',      # NEW: Email files
            '.mbox': 'MBOX',    # NEW: Mailbox archives
            '.txt': 'TEXT',
            '.json': 'JSON',
            '.rtf': 'TEXT',
            '.odt': 'ODT',      # NEW: OpenDocument Text
            '.ods': 'ODS',      # NEW: OpenDocument Spreadsheet
            '.xml': 'XML',      # NEW: XML files
            '.yaml': 'YAML',    # NEW: YAML files
            '.yml': 'YAML',     # NEW: YAML files
            '.html': 'HTML',    # NEW: HTML files
            '.htm': 'HTML',     # NEW: HTML files
            '.md': 'MARKDOWN'   # NEW: Markdown files
        }

        source_files = {}
        ignored_folders = {'incoming', 'completed', '__pycache__', '.git', '.venv', 'node_modules'}
        skip_files = {'data_ingestion_tracker.py', 'README.md', '.gitignore', 'desktop.ini', 'Thumbs.db'}

        logger.info("Scanning ALL subdirectories recursively...")

        # Recursively scan all subdirectories using rglob
        for file_path in self.parser_root.rglob('*'):
            # Skip if it's a directory
            if file_path.is_dir():
                continue

            # Skip if in ignored folders
            if any(ignored in file_path.parts for ignored in ignored_folders):
                continue

            # Skip certain files
            if file_path.name in skip_files:
                continue

            # Skip Office/LibreOffice temp lock files and similar volatile artifacts
            if file_path.name.startswith('~$') or file_path.name.startswith('.~lock.'):
                continue

            # Skip if name contains embedded temp markers (e.g., 123_~$foo.docx)
            if '~$' in file_path.name:
                continue

            # Check extension
            ext = file_path.suffix.lower()
            if ext in extensions:
                source_rel_path = str(file_path.relative_to(self.parser_root))
                if source_rel_path in self.processed_sources:
                    self.results['skipped_existing'] += 1
                    continue

                file_type = extensions[ext]
                if file_type not in source_files:
                    source_files[file_type] = []
                source_files[file_type].append(file_path)
                self.results['total_files'] += 1

                # Track file type stats
                self.results['file_types'][file_type] = self.results['file_types'].get(file_type, 0) + 1

        logger.info(f"Found {self.results['total_files']} files to process")
        for ft, count in self.results['file_types'].items():
            logger.info(f"  {ft}: {count} files")

        return source_files

    # FILE PARSING METHODS

    def parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF"""
        try:
            # Try pdfplumber first (better accuracy)
            try:
                import pdfplumber
                text = ""
                metadata = {}

                with pdfplumber.open(file_path) as pdf:
                    metadata = pdf.metadata or {}
                    for page in pdf.pages:
                        text += page.extract_text() or ""

                return {
                    'file_name': file_path.name,
                    'file_type': 'PDF',
                    'text': text.strip(),
                    'metadata': metadata,
                    'page_count': len(pdf.pages) if hasattr(pdf, 'pages') else None,
                    'extraction_method': 'pdfplumber'
                }
            except ImportError:
                # Fallback to PyPDF2
                try:
                    from PyPDF2 import PdfReader
                    text = ""

                    with open(file_path, 'rb') as f:
                        reader = PdfReader(f)
                        for page in reader.pages:
                            text += page.extract_text() or ""

                    return {
                        'file_name': file_path.name,
                        'file_type': 'PDF',
                        'text': text.strip(),
                        'page_count': len(reader.pages),
                        'extraction_method': 'PyPDF2'
                    }
                except ImportError:
                    logger.warning("PDF parsing libraries not installed. Install: pdfplumber or PyPDF2")
                    return {
                        'file_name': file_path.name,
                        'file_type': 'PDF',
                        'error': 'PDF libraries not installed',
                        'status': 'skipped'
                    }
        except Exception as e:
            err_msg = str(e)
            logger.error(f"PDF parsing error {file_path.name}: {err_msg}")
            status = 'skipped' if ('no /root object' in err_msg.lower() or 'seek of closed file' in err_msg.lower()) else 'error'
            return {
                'file_name': file_path.name,
                'file_type': 'PDF',
                'error': err_msg,
                'status': status
            }

    def parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX"""
        try:
            from docx import Document

            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])

            # Extract tables
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(table_data)

            return {
                'file_name': file_path.name,
                'file_type': 'DOCX',
                'text': text.strip(),
                'tables': tables_data if tables_data else None,
                'paragraph_count': len(doc.paragraphs),
                'extraction_method': 'python-docx'
            }
        except ImportError:
            logger.warning("python-docx not installed. Install: pip install python-docx")
            return {
                'file_name': file_path.name,
                'file_type': 'DOCX',
                'error': 'python-docx not installed',
                'status': 'skipped'
            }
        except Exception as e:
            err_msg = str(e)
            logger.error(f"DOCX parsing error {file_path.name}: {err_msg}")
            # Treat corrupt/non-zip docx as skipped instead of fatal
            status = 'skipped' if 'zip file' in err_msg.lower() else 'error'
            return {
                'file_name': file_path.name,
                'file_type': 'DOCX',
                'error': err_msg,
                'status': status
            }

    def parse_doc(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOC (try to convert or use fallback)"""
        # Try python-docx first in case a .doc is actually OOXML content.
        docx_result = self.parse_docx(file_path)
        if docx_result.get('status') not in {'error', 'skipped'}:
            docx_result['file_type'] = 'DOC'
            return docx_result

        # Fallback: extract printable text from binary payload.
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
                text = ''.join(chr(b) for b in content if 32 <= b < 127)

            return {
                'file_name': file_path.name,
                'file_type': 'DOC',
                'text': text.strip(),
                'extraction_method': 'binary_fallback',
                'note': 'Limited text extraction from legacy binary DOC format'
            }
        except Exception as e:
            logger.error(f"DOC parsing error {file_path.name}: {str(e)}")
            return {
                'file_name': file_path.name,
                'file_type': 'DOC',
                'error': str(e),
                'status': 'error'
            }

    def parse_csv(self, file_path: Path) -> Dict[str, Any]:
        """Parse CSV data"""
        try:
            import csv
            csv.field_size_limit(1024 * 1024)  # raise limit for large fields

            rows = []
            headers = []

            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                reader = csv.DictReader(f)
                headers = reader.fieldnames or []

                for row in reader:
                    rows.append(row)

            return {
                'file_name': file_path.name,
                'file_type': 'CSV',
                'headers': headers,
                'rows': rows,
                'row_count': len(rows),
                'column_count': len(headers),
                'extraction_method': 'csv'
            }
        except Exception as e:
            logger.error(f"CSV parsing error {file_path.name}: {str(e)}")
            return {
                'file_name': file_path.name,
                'file_type': 'CSV',
                'error': str(e),
                'status': 'error'
            }

    def parse_xlsx(self, file_path: Path) -> Dict[str, Any]:
        """Extract data from Excel"""
        try:
            from openpyxl import load_workbook

            wb = load_workbook(file_path)
            sheets_data = {}

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = []

                for row in ws.iter_rows(values_only=True):
                    rows.append(list(row))

                sheets_data[sheet_name] = {
                    'rows': rows,
                    'row_count': ws.max_row,
                    'column_count': ws.max_column
                }

            return {
                'file_name': file_path.name,
                'file_type': 'XLSX',
                'sheets': sheets_data,
                'sheet_count': len(wb.sheetnames),
                'extraction_method': 'openpyxl'
            }
        except ImportError:
            logger.warning("openpyxl not installed. Install: pip install openpyxl")
            return {
                'file_name': file_path.name,
                'file_type': 'XLSX',
                'error': 'openpyxl not installed',
                'status': 'skipped'
            }
        except Exception as e:
            logger.error(f"XLSX parsing error {file_path.name}: {str(e)}")
            return {
                'file_name': file_path.name,
                'file_type': 'XLSX',
                'error': str(e),
                'status': 'error'
            }

    def parse_msg(self, file_path: Path) -> Dict[str, Any]:
        """Extract email data from MSG"""
        try:
            from extract_msg import Message

            msg = Message(file_path)
            attachments = []
            if hasattr(msg, 'attachments'):
                for att in msg.attachments:
                    name = getattr(att, 'filename', None) or getattr(att, 'longFilename', None) or ''
                    attachments.append(name)

            return {
                'file_name': file_path.name,
                'file_type': 'MSG',
                'subject': msg.subject,
                'sender': msg.sender,
                'to': msg.to,
                'cc': msg.cc or [],
                'bcc': msg.bcc or [],
                'body': msg.body,
                'date': str(msg.date) if hasattr(msg, 'date') else None,
                'attachments': attachments,
                'extraction_method': 'extract_msg'
            }
        except ImportError:
            logger.warning("extract_msg not installed. Install: pip install extract_msg")
            return {
                'file_name': file_path.name,
                'file_type': 'MSG',
                'error': 'extract_msg not installed',
                'status': 'skipped'
            }
        except Exception as e:
            logger.error(f"MSG parsing error {file_path.name}: {str(e)}")
            return {
                'file_name': file_path.name,
                'file_type': 'MSG',
                'error': str(e),
                'status': 'error'
            }

    def parse_text(self, file_path: Path) -> Dict[str, Any]:
        """Read plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()

            return {
                'file_name': file_path.name,
                'file_type': 'TEXT',
                'text': text.strip(),
                'character_count': len(text),
                'line_count': len(text.split('\n')),
                'extraction_method': 'text'
            }
        except Exception as e:
            logger.error(f"Text parsing error {file_path.name}: {str(e)}")
            return {
                'file_name': file_path.name,
                'file_type': 'TEXT',
                'error': str(e),
                'status': 'error'
            }

    def parse_json(self, file_path: Path) -> Dict[str, Any]:
        """Read JSON file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            return {
                'file_name': file_path.name,
                'file_type': 'JSON',
                'data': data,
                'extraction_method': 'json'
            }
        except Exception as e:
            logger.error(f"JSON parsing error {file_path.name}: {str(e)}")
            return {
                'file_name': file_path.name,
                'file_type': 'JSON',
                'error': str(e),
                'status': 'error'
            }
    def parse_xls(self, file_path: Path) -> Dict[str, Any]:
        """Extract data from Legacy Excel (.xls)"""
        try:
            import xlrd

            workbook = xlrd.open_workbook(file_path)
            sheets_data = {}

            for sheet_name in workbook.sheet_names():
                sheet = workbook.sheet_by_name(sheet_name)
                rows = []

                for row_idx in range(sheet.nrows):
                    row_data = [sheet.cell_value(row_idx, col_idx) for col_idx in range(sheet.ncols)]
                    rows.append(row_data)

                sheets_data[sheet_name] = {
                    'rows': rows,
                    'row_count': sheet.nrows,
                    'column_count': sheet.ncols
                }

            return {
                'file_name': file_path.name,
                'file_type': 'XLS',
                'sheets': sheets_data,
                'sheet_count': len(workbook.sheet_names()),
                'extraction_method': 'xlrd'
            }
        except ImportError:
            logger.warning("xlrd not installed. Install: pip install xlrd")
            return {
                'file_name': file_path.name,
                'file_type': 'XLS',
                'error': 'xlrd not installed',
                'status': 'skipped'
            }
        except Exception as e:
            logger.error(f"XLS parsing error {file_path.name}: {str(e)}")
            return {
                'file_name': file_path.name,
                'file_type': 'XLS',
                'error': str(e),
                'status': 'error'
            }

    def parse_odt(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from OpenDocument Text (.odt)"""
        try:
            import zipfile
            import xml.etree.ElementTree as ET

            # ODT files are zip archives containing content.xml
            with zipfile.ZipFile(file_path, 'r') as odt_zip:
                # Read content.xml
                content_xml = odt_zip.read('content.xml')

                # Parse XML and extract text
                root = ET.fromstring(content_xml)

                # ODT uses namespaces
                namespaces = {
                    'text': 'urn:oasis:names:tc:opendocument:xmlns:text:1.0',
                    'office': 'urn:oasis:names:tc:opendocument:xmlns:office:1.0'
                }

                # Extract all text elements
                text_elements = []
                for elem in root.iter():
                    if elem.text:
                        text_elements.append(elem.text.strip())
                    if elem.tail:
                        text_elements.append(elem.tail.strip())

                text = '\n'.join([t for t in text_elements if t])

                return {
                    'file_name': file_path.name,
                    'file_type': 'ODT',
                    'text': text.strip(),
                    'extraction_method': 'zipfile_xml',
                    'note': 'OpenDocument format parsed as XML from ZIP'
                }
        except Exception as e:
            logger.error(f"ODT parsing error {file_path.name}: {str(e)}")
            return {
                'file_name': file_path.name,
                'file_type': 'ODT',
                'error': str(e),
                'status': 'error'
            }

    def parse_ods(self, file_path: Path) -> Dict[str, Any]:
        """Extract data from OpenDocument Spreadsheet (.ods)"""
        try:
            import zipfile
            import xml.etree.ElementTree as ET

            with zipfile.ZipFile(file_path, 'r') as ods_zip:
                content_xml = ods_zip.read('content.xml')
                root = ET.fromstring(content_xml)

                # Extract table data
                tables_data = {}
                for table in root.iter('{urn:oasis:names:tc:opendocument:xmlns:table:1.0}table'):
                    table_name = table.get('{urn:oasis:names:tc:opendocument:xmlns:table:1.0}name', 'Sheet')
                    rows = []

                    for row in table.iter('{urn:oasis:names:tc:opendocument:xmlns:table:1.0}table-row'):
                        row_data = []
                        for cell in row.iter('{urn:oasis:names:tc:opendocument:xmlns:table:1.0}table-cell'):
                            cell_text = ''.join(cell.itertext()).strip()
                            row_data.append(cell_text)
                        if row_data:
                            rows.append(row_data)

                    tables_data[table_name] = {
                        'rows': rows,
                        'row_count': len(rows)
                    }

                return {
                    'file_name': file_path.name,
                    'file_type': 'ODS',
                    'sheets': tables_data,
                    'sheet_count': len(tables_data),
                    'extraction_method': 'zipfile_xml'
                }
        except Exception as e:
            logger.error(f"ODS parsing error {file_path.name}: {str(e)}")
            return {
                'file_name': file_path.name,
                'file_type': 'ODS',
                'error': str(e),
                'status': 'error'
            }

    def parse_eml(self, file_path: Path) -> Dict[str, Any]:
        """Extract email data from EML files"""
        try:
            import email
            from email import policy

            with open(file_path, 'rb') as f:
                msg = email.message_from_binary_file(f, policy=policy.default)

            # Extract body
            body = ""
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    if content_type == 'text/plain':
                        body += part.get_content()
            else:
                body = msg.get_content()

            # Extract attachments
            attachments = []
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_disposition() == 'attachment':
                        attachments.append(part.get_filename())

            return {
                'file_name': file_path.name,
                'file_type': 'EML',
                'subject': msg.get('Subject', ''),
                'sender': msg.get('From', ''),
                'to': msg.get('To', ''),
                'cc': msg.get('Cc', ''),
                'date': msg.get('Date', ''),
                'body': body.strip() if body else '',
                'attachments': attachments,
                'extraction_method': 'email.parser'
            }
        except Exception as e:
            logger.error(f"EML parsing error {file_path.name}: {str(e)}")
            return {
                'file_name': file_path.name,
                'file_type': 'EML',
                'error': str(e),
                'status': 'error'
            }

    def parse_mbox(self, file_path: Path) -> Dict[str, Any]:
        """Extract emails from MBOX mailbox archive"""
        try:
            import mailbox

            mbox = mailbox.mbox(file_path)
            messages = []

            for message in mbox:
                messages.append({
                    'subject': message.get('Subject', ''),
                    'sender': message.get('From', ''),
                    'to': message.get('To', ''),
                    'date': message.get('Date', ''),
                    'body': message.get_payload()[:500] if message.get_payload() else ''  # Limit body length
                })

            return {
                'file_name': file_path.name,
                'file_type': 'MBOX',
                'message_count': len(messages),
                'messages': messages[:100],  # Limit to first 100 messages
                'extraction_method': 'mailbox'
            }
        except Exception as e:
            logger.error(f"MBOX parsing error {file_path.name}: {str(e)}")
            return {
                'file_name': file_path.name,
                'file_type': 'MBOX',
                'error': str(e),
                'status': 'error'
            }

    def parse_xml(self, file_path: Path) -> Dict[str, Any]:
        """Parse XML files"""
        try:
            import xml.etree.ElementTree as ET

            tree = ET.parse(file_path)
            root = tree.getroot()

            # Extract text from all elements
            text_elements = []
            for elem in root.iter():
                if elem.text and elem.text.strip():
                    text_elements.append(elem.text.strip())

            # Also capture XML as string for structure
            xml_string = ET.tostring(root, encoding='unicode')[:5000]  # Limit length

            return {
                'file_name': file_path.name,
                'file_type': 'XML',
                'text': '\n'.join(text_elements),
                'root_tag': root.tag,
                'xml_preview': xml_string,
                'extraction_method': 'xml.etree'
            }
        except Exception as e:
            logger.error(f"XML parsing error {file_path.name}: {str(e)}")
            return {
                'file_name': file_path.name,
                'file_type': 'XML',
                'error': str(e),
                'status': 'error'
            }

    def parse_yaml(self, file_path: Path) -> Dict[str, Any]:
        """Parse YAML files"""
        try:
            import yaml

            with open(file_path, 'r', encoding='utf-8') as f:
                data = yaml.safe_load(f)

            return {
                'file_name': file_path.name,
                'file_type': 'YAML',
                'data': data,
                'extraction_method': 'pyyaml'
            }
        except ImportError:
            logger.warning("PyYAML not installed. Install: pip install PyYAML")
            return {
                'file_name': file_path.name,
                'file_type': 'YAML',
                'error': 'PyYAML not installed',
                'status': 'skipped'
            }
        except Exception as e:
            logger.error(f"YAML parsing error {file_path.name}: {str(e)}")
            return {
                'file_name': file_path.name,
                'file_type': 'YAML',
                'error': str(e),
                'status': 'error'
            }

    def parse_html(self, file_path: Path) -> Dict[str, Any]:
        """Parse HTML files and extract text"""
        try:
            from bs4 import BeautifulSoup

            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()

            soup = BeautifulSoup(html_content, 'lxml')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Extract text
            text = soup.get_text(separator='\n', strip=True)

            # Extract metadata
            title = soup.title.string if soup.title else ''
            meta_description = ''
            meta_tag = soup.find('meta', attrs={'name': 'description'})
            if meta_tag:
                meta_description = meta_tag.get('content', '')

            return {
                'file_name': file_path.name,
                'file_type': 'HTML',
                'text': text,
                'title': title,
                'meta_description': meta_description,
                'extraction_method': 'beautifulsoup4'
            }
        except ImportError:
            logger.warning("beautifulsoup4 not installed. Install: pip install beautifulsoup4 lxml")
            return {
                'file_name': file_path.name,
                'file_type': 'HTML',
                'error': 'beautifulsoup4 not installed',
                'status': 'skipped'
            }
        except Exception as e:
            logger.error(f"HTML parsing error {file_path.name}: {str(e)}")
            return {
                'file_name': file_path.name,
                'file_type': 'HTML',
                'error': str(e),
                'status': 'error'
            }

    def parse_markdown(self, file_path: Path) -> Dict[str, Any]:
        """Parse Markdown files and convert to text"""
        try:
            import markdown
            from bs4 import BeautifulSoup

            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                md_content = f.read()

            # Convert markdown to HTML
            html = markdown.markdown(md_content)

            # Extract plain text from HTML
            soup = BeautifulSoup(html, 'html.parser')
            text = soup.get_text(separator='\n', strip=True)

            return {
                'file_name': file_path.name,
                'file_type': 'MARKDOWN',
                'text': text,
                'raw_markdown': md_content[:5000],  # Keep first 5000 chars of original
                'extraction_method': 'markdown+beautifulsoup4'
            }
        except ImportError:
            logger.warning("markdown not installed. Install: pip install markdown beautifulsoup4")
            return {
                'file_name': file_path.name,
                'file_type': 'MARKDOWN',
                'error': 'markdown not installed',
                'status': 'skipped'
            }
        except Exception as e:
            logger.error(f"Markdown parsing error {file_path.name}: {str(e)}")
            return {
                'file_name': file_path.name,
                'file_type': 'MARKDOWN',
                'error': str(e),
                'status': 'error'
            }

    def parse_zip(self, file_path: Path) -> Dict[str, Any]:
        """Parse ZIP archive and extract text-like content from supported inner files."""
        try:
            supported_text_exts = {'.txt', '.csv', '.json', '.md', '.html', '.htm', '.xml', '.eml'}
            entries = []
            aggregate_chunks = []

            with zipfile.ZipFile(file_path, 'r') as zf:
                for info in zf.infolist():
                    if info.is_dir():
                        continue

                    inner_name = info.filename
                    inner_ext = Path(inner_name).suffix.lower()
                    if inner_ext not in supported_text_exts:
                        continue

                    try:
                        raw = zf.read(info)
                        text = raw.decode('utf-8', errors='ignore')
                    except Exception:
                        continue

                    if not text.strip():
                        continue

                    # Special handling for CSV summary
                    row_count = None
                    if inner_ext == '.csv':
                        try:
                            row_count = sum(1 for _ in csv.reader(io.StringIO(text)))
                        except Exception:
                            row_count = None

                    preview = text[:5000]
                    entries.append({
                        'entry_name': inner_name,
                        'entry_size': info.file_size,
                        'row_count': row_count,
                        'preview': preview,
                    })

                    aggregate_chunks.append(f"\n--- {inner_name} ---\n")
                    aggregate_chunks.append(preview)

            if not entries:
                return {
                    'file_name': file_path.name,
                    'file_type': 'ZIP',
                    'status': 'skipped',
                    'error': 'No supported text-like entries found in archive'
                }

            return {
                'file_name': file_path.name,
                'file_type': 'ZIP',
                'entry_count': len(entries),
                'entries': entries[:200],
                'text': '\n'.join(aggregate_chunks)[:50000],
                'extraction_method': 'zipfile',
            }
        except zipfile.BadZipFile:
            return {
                'file_name': file_path.name,
                'file_type': 'ZIP',
                'status': 'error',
                'error': 'Invalid ZIP archive'
            }
        except Exception as e:
            logger.error(f"ZIP parsing error {file_path.name}: {str(e)}")
            return {
                'file_name': file_path.name,
                'file_type': 'ZIP',
                'error': str(e),
                'status': 'error'
            }
    # HELPER METHODS

    def compute_hash(self, file_path: Path) -> str:
        """Compute MD5 hash of file for deduplication"""
        md5 = hashlib.md5()
        with open(file_path, 'rb') as f:
            md5.update(f.read())
        return md5.hexdigest()

    def save_as_json(self, file_path: Path, extracted_data: Dict[str, Any]) -> Path:
        """
        Save extracted data as JSON to ai_data_final/
        Filename: {original_name}_{timestamp}.json
        """
        timestamp = int(datetime.now().timestamp() * 1000)
        safe_stem = re.sub(r'[\\/:*?"<>|\r\n\t]+', '_', file_path.stem)
        safe_stem = re.sub(r'\s+', ' ', safe_stem).strip(' ._') or 'unnamed'
        safe_stem = safe_stem[:120]
        output_name = f"{safe_stem}_{timestamp}.json"
        output_file = self.output_root / output_name

        source_path = str(file_path.relative_to(self.parser_root))
        json_data = {
            'source_file': file_path.name,
            'source_path': source_path,
            'file_type': extracted_data.get('file_type'),
            'file_hash': self.compute_hash(file_path) if file_path.exists() else None,
            'extracted_data': extracted_data,
            'processed_at': datetime.now().isoformat()
        }

        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, indent=2, ensure_ascii=False, default=str)

        self.processed_sources.add(source_path)
        self._append_processed_source(source_path)

        return output_file

    # MAIN EXECUTION

    def run(self, max_files: Optional[int] = None) -> Dict[str, Any]:
        """
        Execute complete parsing pipeline

        Args:
            max_files: Limit number of files to process (for testing)
        """
        logger.info("=" * 60)
        logger.info("PARSER ENGINE STARTING")
        logger.info("=" * 60)

        # Step 1: Discover files
        source_files = self.discover_source_files()

        if self.results['total_files'] == 0:
            logger.warning("No parseable files found!")
            self.generate_ingestion_status_manifest()
            return self.results

        # Step 2: Process each file
        file_count = 0
        for file_type, files in source_files.items():
            parser_method_name = f"parse_{file_type.lower()}"
            parser_method = getattr(self, parser_method_name, None)

            if not parser_method:
                logger.warning(f"No parser for {file_type}")
                self.results['skipped'] += len(files)
                continue

            for file_path in files:
                file_count += 1

                if max_files and file_count > max_files:
                    logger.info(f"Reached max_files limit: {max_files}")
                    break

                try:
                    logger.info(f"[{file_count}/{self.results['total_files']}] Processing: {file_path.name}")

                    if not file_path.exists():
                        logger.warning("  SKIPPED: File no longer exists")
                        self.results['skipped'] += 1
                        continue

                    # Parse file
                    extracted = parser_method(file_path)

                    # Check for errors
                    if extracted.get('status') == 'skipped':
                        logger.warning(f"  SKIPPED: {extracted.get('error', 'Unknown reason')}")
                        source_rel_path = str(file_path.relative_to(self.parser_root))
                        self._mark_failed_source(
                            source_rel_path,
                            status='skipped',
                            reason=str(extracted.get('error', 'Unknown reason')),
                        )
                        self.results['skipped'] += 1
                        continue

                    if extracted.get('status') == 'error':
                        logger.error(f"  ERROR: {extracted.get('error')}")
                        source_rel_path = str(file_path.relative_to(self.parser_root))
                        self._mark_failed_source(
                            source_rel_path,
                            status='error',
                            reason=str(extracted.get('error') or ''),
                        )
                        self.results['errors'].append({
                            'file': file_path.name,
                            'error': extracted.get('error')
                        })
                        continue

                    # Save as JSON
                    output_file = self.save_as_json(file_path, extracted)
                    source_rel_path = str(file_path.relative_to(self.parser_root))
                    self._clear_failed_source(source_rel_path)
                    self.results['processed'] += 1
                    logger.info(f"  âœ“ Saved to {output_file.name}")

                except Exception as e:
                    logger.error(f"  FATAL: {str(e)}")
                    source_rel_path = str(file_path.relative_to(self.parser_root))
                    self._mark_failed_source(source_rel_path, status='error', reason=str(e))
                    self.results['errors'].append({
                        'file': file_path.name,
                        'error': str(e)
                    })

        # Step 3: Generate report
        self._save_failed_sources()
        self.generate_report()
        self.generate_ingestion_status_manifest()

        return self.results

    def generate_report(self) -> Path:
        """Create parsing summary report"""
        success_count = self.results['processed']
        total = self.results['total_files']
        success_rate = (success_count / max(total, 1)) * 100

        report = {
            'summary': {
                'total_files': total,
                'successfully_parsed': success_count,
                'skipped': self.results['skipped'],
                'failed': len(self.results['errors']),
                'success_rate': f"{success_rate:.1f}%"
            },
            'file_types': self.results['file_types'],
            'errors': self.results['errors'][:100],  # Limit to first 100 errors
            'timestamp': self.results['created_at']
        }

        report_file = self.output_root / 'parsing_report.json'
        with open(report_file, 'w', encoding='utf-8') as f:
            json.dump(report, f, indent=2, ensure_ascii=False)

        # Print summary
        logger.info("=" * 60)
        logger.info("PARSING COMPLETE")
        logger.info("=" * 60)
        logger.info(f"Total files: {total}")
        logger.info(f"Successfully parsed: {success_count}")
        logger.info(f"Skipped: {self.results['skipped']}")
        logger.info(f"Skipped existing: {self.results.get('skipped_existing', 0)}")
        logger.info(f"Failed: {len(self.results['errors'])}")
        logger.info(f"Success rate: {success_rate:.1f}%")
        logger.info(f"Report saved: {report_file}")
        logger.info("=" * 60)

        return report_file


def main():
    """Main execution"""
    import argparse

    parser = argparse.ArgumentParser(description='Automated Parser Engine')
    parser.add_argument('--max-files', type=int, help='Limit files to process (for testing)')
    parser.add_argument('--root', default='automated_parser', help='Parser root directory')

    args = parser.parse_args()

    # Run parser
    engine = AutomatedParserEngine(args.root)
    results = engine.run(max_files=args.max_files)

    # Exit with appropriate code
    exit_code = 0 if len(results['errors']) == 0 else 1
    sys.exit(exit_code)


if __name__ == '__main__':
    main()

